"""_match_ds260_key phải nhận diện đủ 6 câu Section C (Previous U.S Travel) + 44 câu Section F
(Security and Background) + 3 câu Section G (Social Security Number).

Báo lỗi thực tế 2026-08-05: các mục này hiển thị/lưu đúng trên web (đọc trực tiếp từ
ds260_mapping qua resolve_ds260_form), nhưng KHÔNG được điền vào file Word xuất ra —
export_ds260.py dò field bằng cách so khớp regex với text thật trong template
(TRAVEL_LABEL_PATTERNS/SECURITY_LABEL_PATTERNS), và các pattern này trước đó hoàn toàn CHƯA
TỒN TẠI cho 2 mục này — field bị bỏ qua âm thầm, không lỗi/không cảnh báo.
"""

import json

from app.services.export_ds260 import _match_ds260_key


def _all_travel_and_security_fields() -> list[dict]:
    with open("data/doc_schemas/ds260_mapping.json", encoding="utf-8") as f:
        data = json.load(f)
    travel_keys = {
        "been_in_us", "us_travel_history", "issued_us_visa",
        "us_visa_history", "refused_us_visa", "us_visa_refusal_history",
    }
    out: list[dict] = []
    for sec in data.get("sections", []):
        if sec.get("id") in ("section_security", "section_ssn"):
            out.extend(sec["fields"])
    def walk(obj):
        if isinstance(obj, dict):
            yield obj
            for v in obj.values():
                yield from walk(v)
        elif isinstance(obj, list):
            for v in obj:
                yield from walk(v)
    for o in walk(data):
        if isinstance(o, dict) and o.get("key") in travel_keys:
            out.append(o)
    return out


def test_all_travel_and_security_fields_match_own_label():
    """Mỗi field trong 2 mục phải tự khớp đúng CHÍNH NÓ bằng label thật trong mapping.json —
    không khớp nhầm sang field khác, không bị bỏ trống (chuỗi rỗng)."""
    fields = _all_travel_and_security_fields()
    assert len(fields) >= 50, "Danh sách field Section C+F+G phải đủ ~53 câu"

    mismatches = []
    for f in fields:
        got = _match_ds260_key(f["label"])
        if got != f["key"]:
            mismatches.append((f["key"], f["label"], got))

    assert not mismatches, (
        f"{len(mismatches)} field không khớp đúng chính nó: "
        + "; ".join(f"{k!r} (label={lbl!r}) => matched {got!r}" for k, lbl, got in mismatches)
    )


def test_real_template_wording_from_screenshot():
    """Đúng nguyên văn text trong file Word thật (chụp màn hình báo lỗi 2026-08-05) — kể cả
    câu có lỗi chính tả "druc" thay vì "drug" trong template gốc."""
    cases = [
        ("Have you ever been in the U.S?", "been_in_us"),
        ("Have you ever been issued a U.S visa?", "issued_us_visa"),
        ("Refused U.S visa / admission / withdrawn?", "refused_us_visa"),
        (
            "Do you have a communication disease of public health significance such as tuberculosis?",
            "communicable_disease",
        ),
        (
            "Do you have documentation to establish that you have received vaccinations "
            "in accordance with U.S law?",
            "has_vaccination_docs",
        ),
        (
            "Do you have a mental or physical disorder that poses or is likely to pose a "
            "threat to the safety or welfare of yourself or others?",
            "mental_physical_disorder",
        ),
        ("Are you or have ever been a druc abuser or addict?", "drug_abuser"),
    ]
    for text, expected in cases:
        assert _match_ds260_key(text) == expected, f"Text {text!r} phải khớp {expected!r}"


def test_family_variant_not_confused_with_base_question():
    """Câu hỏi "người thân"/family — dùng cụm tiếng Việt CHUNG với câu gốc của chính đương đơn
    (vd. "buôn người", "khủng bố", "lạm quyền chiếm đoạt tài sản") — phải khớp đúng field
    family_*, KHÔNG bị pattern của câu gốc nuốt mất do đứng trước trong danh sách (bug đã fix
    2026-08-05: family_* phải được kiểm tra TRƯỚC bản gốc)."""
    cases = [
        (
            "Family human trafficking — benefited (5y)? / Hưởng lợi buôn người của người thân (5 năm)?",
            "family_human_trafficking_benefited",
        ),
        (
            "Family terrorist activity (5y)? / Người thân hoạt động khủng bố (5 năm)?",
            "family_terrorist_activity",
        ),
        (
            "Family abuse of position — property? / Người thân lạm quyền chiếm đoạt tài sản?",
            "family_abuse_position_property",
        ),
        (
            "Family CWC disclosure? / Người thân tiết lộ bí mật CWC?",
            "family_chemical_weapons_disclosure",
        ),
        # Câu gốc (không có "người thân") vẫn phải khớp đúng field của chính nó.
        ("Committed human trafficking? / Buôn người?", "human_trafficking_committed"),
        ("Engaged / seek terrorist activities? / Hoạt động khủng bố?", "terrorist_activities"),
    ]
    for text, expected in cases:
        assert _match_ds260_key(text) == expected, f"Text {text!r} phải khớp {expected!r}"


def test_smart_fill_actually_inserts_value_not_just_recognizes_key():
    """Lớp bug THỨ HAI (báo lỗi thực tế 2026-08-05, lần 2): _match_ds260_key() nhận diện đúng
    field_key KHÔNG có nghĩa là giá trị được CHÈN vào dòng — dòng không có dấu ':' (đa số câu
    Yes/No mẫu DS-260) chỉ được điền khi field nằm trong whitelist của _smart_fill_ds260_line
    (_is_question_fill_key / _APPEND_NO_COLON_KEYS / _OTHER_USED_HISTORY). Test này dùng
    NGUYÊN VĂN text từ file export thật (ds260_TRIEU_THI_DUYEN.docx) đã báo lỗi."""
    from app.services.export_ds260 import _smart_fill_ds260_line

    values = {
        "been_in_us": "No",
        "issued_us_visa": "No",
        "refused_us_visa": "No",
        "communicable_disease": "No",
        "has_vaccination_docs": "Yes",
        "mental_physical_disorder": "No",
        "drug_abuser": "No",
        "applied_ssn_before": "No",
        "want_ssn_issued": "Yes",
        "authorize_ssn_disclosure": "Yes",
    }

    lines = {
        "Have you ever been in the U.S? (Bạn đã từng đến nước MỸ chưa?)": "No",
        "Have you ever been issued a U.S visa? (Bạn đã từng được cấp Visa MỸ chưa?)": "No",
        (
            "Have you ever been refused a U.S visa, been refused admission to the U.S or "
            "withdrawn your application for admission at the port of entry? (Bạn đã từng bị "
            "từ chối Visa MỸ, bị từ chối Nhập Cảnh vào MỸ, bị rút hồ sơ xin Nhập Cảnh vào MỸ chưa?)"
        ): "No",
        (
            "Do you have a communication disease of public health significance such as "
            "tuberculosis? (Bạn có bị bệnh liên quan đến lây lan trong môi trường công cộng "
            "như Bệnh Lao?)"
        ): "No",
        (
            "Do you have documentation to establish that you have received vaccinations in "
            "accordance with U.S law? (Bạn có giấy tờ chứng tỏ Bạn đã chích ngừa theo Luật "
            "của nước MỸ)"
        ): "Yes",
        "Are you or have ever been a druc abuser or addict? (Bạn đang hoặc đã từng nghiện ngập Ma túy)": "No",
    }

    for text, expected_answer in lines.items():
        out = _smart_fill_ds260_line(text, values, "applicant", set())
        assert out != text, f"Dòng KHÔNG được điền (bug cũ): {text!r}"
        assert out.rstrip().endswith(expected_answer), (
            f"Điền sai đáp án — kỳ vọng kết thúc bằng {expected_answer!r}, ra: {out!r}"
        )


def _walk_template_with_context(doc):
    """Đi qua template theo ĐÚNG cách fill_ds260_docx_template chạy — trả về (context, text) cho
    mỗi paragraph, có cả context tracking. Đây là điều các test dựa-trên-label không mô phỏng
    được, và là chỗ 3 bug 2026-08-05 ẩn nấp (context bị kẹt → cổng section nuốt câu hỏi)."""
    from app.services.export_ds260 import _update_section_context

    out = []
    ctx = "applicant"
    for p in doc.paragraphs:
        ctx = _update_section_context(p.text, ctx)
        out.append((ctx, p.text))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    ctx = _update_section_context(p.text, ctx)
                    out.append((ctx, p.text))
    return out


# Probe (một cụm đặc trưng) → key đúng cho từng câu Section C/F/G trong TEMPLATE THẬT.
_TEMPLATE_PROBES = {
    "been_in_us": "Have you ever been in the U.S?",
    "issued_us_visa": "Have you ever been issued a U.S visa?",
    "refused_us_visa": "Have you ever been refused a U.S visa",
    "communicable_disease": "communication disease of public health",
    "has_vaccination_docs": "received vaccinations in accordance",
    "mental_physical_disorder": "mental or physical disorder",
    "drug_abuser": "druc abuser or addict",
    "arrested_convicted": "ever been arrested or convicted",
    "violated_controlled_substance": "engaged in a conspiracy to violate, any law relating to controlled",
    "family_drug_trafficking_benefited": "violated any controlled substance trafficking law",
    "prostitution": "engage in prostitution",
    "money_laundering": "engage in, money laundering",
    "human_trafficking_committed": "committed or conspired to commit a human trafficking offense in the U.S or outside the U.S?",
    "human_trafficking_aided": "aided, abetted, assisted, or colluded",
    "family_human_trafficking_benefited": "commit a human trafficking offense in the U.S or outside the U.S and have you",
    "espionage_sabotage": "espionage, sabotage, export control",
    "terrorist_activities": "seek to engage in terrorist activities",
    "terrorist_financial_support": "financial assistance or other support to terrorists",
    "terrorist_org_member": "member or representative of a terrorist organization",
    "family_terrorist_activity": "individual who has engaged in terrorist activity",
    "genocide": "participated in genocide",
    "torture": "participated in torture",
    "extrajudicial_killings": "extrajudicial killings",
    "child_soldiers": "recruitment of or the use of child soldiers",
    "religious_freedom_violations": "violations of religious freedom",
    "communist_party_member": "Communist or other totalitarian party",
    "colombia_groups_support": "Revolutionary Armed Forces of Columbia",
    "abuse_position_confiscate_property": "Have you ever, through abuse of governmental",
    "family_abuse_position_property": "minor child, or agent of an individual who has through abuse",
    "population_control_abortion": "population controls forcing a woman",
    "chemical_weapons_disclosure": "disclosed or trafficked in confidential U.S business",
    "family_chemical_weapons_disclosure": "agent of an individual who has disclosed or trafficked",
    "visa_fraud": "fraud or willful misrepresentation",
    "removed_deported": "removed or deported from any country",
    "withheld_custody_us_child": "Have you ever withheld custody of a U.S citizen child",
    "assisted_withhold_custody": "intentionally assisted another person in withholding custody",
    "voted_illegally": "voted in the U.S in violation",
    "renounced_citizenship_tax": "renounced U.S citizenship for the purpose of avoiding taxation",
    "attended_public_school_f_status": "attended a public elementary school",
    "skilled_labor_no_certification": "skilled or unskilled labor",
    "foreign_medical_grad_no_exam": "graduate of a foreign medical school",
    "health_care_worker_no_cert": "health care worker seeking to perform",
    "permanently_ineligible_citizenship": "permanently ineligible for U.S citizenship",
    "departed_us_evade_military": "evade military service during a time of war",
    "polygamy": "practice polygamy",
    "exchange_visitor_j_unfulfilled": "former exchange visitor",
    "frivolous_asylum": "frivolous application for asylum",
    "public_charge": "likely to become a public charge",
    "applied_ssn_before": "applied for a Social Security number",
    "want_ssn_issued": "Administration to issue a Social Security number and a card",
    "authorize_ssn_disclosure": "authorize disclosure of information from this form",
}


def _load_real_template():
    """Template DS-260 thật trong repo; skip nếu môi trường không có file."""
    import glob

    import pytest
    from docx import Document

    matches = sorted(glob.glob("templates/forms/*.docx"))
    if not matches:
        pytest.skip("Không tìm thấy template DS-260 .docx trong templates/forms/")
    return Document(matches[0])


def test_real_template_each_question_matches_exactly_its_own_key():
    """Chạy trên FILE WORD THẬT + context tracking (không phải label tổng hợp). Bắt cả 3 bug
    2026-08-05: (1) context kẹt child_/edu_/military nuốt câu hỏi Travel/Security/SSN; (2) regex
    lệch chữ với câu gốc; (3) câu người-thân bị câu-gốc cướp key (hoặc ngược lại). Mỗi câu phải
    tự khớp CHÍNH NÓ, và không key nào bị khớp bởi >1 paragraph."""
    from collections import Counter

    from app.services.export_ds260 import _match_ds260_key

    para_ctx = _walk_template_with_context(_load_real_template())

    wrong = []
    for key, probe in _TEMPLATE_PROBES.items():
        hits = [(ctx, t) for ctx, t in para_ctx if probe.lower() in t.lower()]
        assert hits, f"Không tìm thấy câu hỏi cho {key!r} (probe={probe!r}) — template đổi chữ?"
        ctx, text = hits[0]
        got = _match_ds260_key(text, ctx)
        if got != key:
            wrong.append((key, got, ctx, text[:90]))
    assert not wrong, "Câu hỏi khớp SAI key:\n" + "\n".join(
        f"  kỳ vọng {k!r} nhưng ra {g!r} (ctx={c}) | {t!r}" for k, g, c, t in wrong
    )

    matched = Counter(
        _match_ds260_key(t, ctx) for ctx, t in para_ctx
    )
    multi = {k: n for k, n in matched.items() if k in _TEMPLATE_PROBES and n > 1}
    assert not multi, f"Key bị nhiều paragraph cùng khớp (mất field khác): {multi}"


def test_real_template_end_to_end_fill_no_blank_answers():
    """Xuất file thật với đáp án No/Yes cho toàn bộ 49 câu Travel/Security/SSN — không câu nào
    được để trống. Đây là kịch bản báo lỗi gốc: web fill oke, xuất Word bị trống."""
    import tempfile
    from pathlib import Path

    from docx import Document

    from app.services.export_ds260 import (
        _EXPORT_DEBUG_WATCH_KEYS,
        _match_ds260_key,
        fill_ds260_docx_template,
    )
    import glob

    import pytest

    matches = sorted(glob.glob("templates/forms/*.docx"))
    if not matches:
        pytest.skip("Không tìm thấy template DS-260 .docx trong templates/forms/")
    tpl = Path(matches[0])

    # 3 ô "chi tiết lịch sử" là free-text (điền khi trả lời Yes), đúng ra để TRỐNG khi không có
    # lịch sử — không phải câu Yes/No, nên loại khỏi khẳng định "phải kết thúc bằng đáp án".
    history_detail = {"us_travel_history", "us_visa_history", "us_visa_refusal_history"}
    # Yes/No của 3 câu Section C nằm ở DÒNG GẠCH NGANG ĐẦU TIÊN riêng (không nối vào cuối dòng
    # câu hỏi nữa, khách yêu cầu 2026-08-14, xem _fill_blank_line_history_block) — kiểm tra riêng
    # bên dưới bằng cách khớp GIÁ TRỊ đoạn nguyên (không phải endswith trên dòng câu hỏi).
    blank_line_yesno = {"been_in_us", "issued_us_visa", "refused_us_visa"}
    yesno_keys = [k for k in _EXPORT_DEBUG_WATCH_KEYS if k not in history_detail and k not in blank_line_yesno]
    values = {k: ("Yes" if k in ("has_vaccination_docs", "want_ssn_issued", "authorize_ssn_disclosure") else "No")
              for k in yesno_keys}
    values.update({k: "No" for k in blank_line_yesno})

    with tempfile.TemporaryDirectory() as d:
        out = Path(d) / "out.docx"
        fill_ds260_docx_template(tpl, out, values)
        doc = Document(str(out))
        para_ctx = _walk_template_with_context(doc)
        blanks = []
        for ctx, text in para_ctx:
            key = _match_ds260_key(text, ctx)
            if key in yesno_keys:
                answer = values[key]
                # endswith (KHÔNG dùng "in") — đáp án "No" hay lọt vào giữa câu ("No" ⊂ "November
                # 30, 1996") khiến kiểm tra kiểu "in" báo pass oan trong khi câu vẫn trống.
                if not text.rstrip().endswith(answer):
                    blanks.append((key, answer, text[:90]))
        assert not blanks, "Câu bị bỏ trống khi xuất Word:\n" + "\n".join(
            f"  {k} (đáp án {a!r}) | {t!r}" for k, a, t in blanks
        )

        # 3 câu Section C: đáp án phải xuất hiện ở 1 đoạn RIÊNG (dòng gạch ngang đầu tiên) là
        # đúng bằng "No" — KHÔNG được nối vào cuối dòng câu hỏi/dòng chú thích.
        exact_no_paragraphs = {t.strip() for _ctx, t in para_ctx if t.strip() == "No"}
        assert len(exact_no_paragraphs) >= 1, "Không thấy đoạn riêng nào chỉ chứa 'No' cho Section C"
        question_lines = {
            t
            for _ctx, t in para_ctx
            if "Have you ever been in the U.S" in t
            or "Have you ever been issued a U.S visa" in t
            or "Have you ever been refused a U.S visa" in t
        }
        for line in question_lines:
            assert not line.rstrip().endswith("No"), (
                f"Câu hỏi Section C KHÔNG được nối 'No' vào cuối dòng nữa: {line!r}"
            )


def test_no_children_answers_question_not_the_count_field():
    """Khi không có con: 'Do you have any children?' phải có 'No' ở cuối dòng; 'Number of
    Children (...):' phải để TRỐNG — không phải ngược lại. Bug thật (2026-08-05): children_count
    bị ds260_mapping gán sẵn 'No' (để hiển thị nhất quán trên web, xem Case C-2 trong
    normalize_ds260_na_conventions) — giá trị đó rò vào ô SỐ trong Word, còn câu hỏi Yes/No thì bị
    bỏ trống vì children_used chưa từng nằm trong whitelist điền dòng."""
    from app.services.export_ds260 import _prepare_display_values, _smart_fill_ds260_line

    values = _prepare_display_values({"children_used": "No", "children_count": "No"})

    question = _smart_fill_ds260_line(
        "Do you have any children? (Bạn có CON không?)", values, "applicant", set()
    )
    count_line = _smart_fill_ds260_line(
        "Number of Children (bao nhiêu CON?):", values, "applicant", set()
    )
    assert question.rstrip().endswith("No"), f"Câu hỏi phải có đáp án 'No': {question!r}"
    assert count_line.rstrip().endswith(":"), f"Ô số con phải để trống: {count_line!r}"


def test_children_count_still_displays_real_number():
    """Khi CÓ con, children_count phải hiện đúng số — fix cho case 'không con' không được làm
    mất giá trị thật khi có con."""
    from app.services.export_ds260 import _prepare_display_values, _smart_fill_ds260_line

    values = _prepare_display_values({"children_used": "Yes", "children_count": "2"})
    count_line = _smart_fill_ds260_line(
        "Number of Children (bao nhiêu CON?):", values, "applicant", set()
    )
    assert count_line.rstrip().endswith("2"), f"Phải hiện số con thật: {count_line!r}"


def test_all_security_and_ssn_keys_are_fillable_without_colon():
    """Toàn bộ 44 field Section F + 3 field SSN phải nằm trong whitelist cho phép điền dòng
    không có dấu ':' — nếu thiếu field nào, dòng đó sẽ bị bỏ trống âm thầm khi xuất Word."""
    from app.services.export_ds260 import (
        _APPEND_NO_COLON_KEYS,
        _OTHER_USED_HISTORY,
        _is_question_fill_key,
    )

    fields = [f for f in _all_travel_and_security_fields() if f["key"] not in (
        "us_travel_history", "us_visa_history", "us_visa_refusal_history",
    )]
    missing = [
        f["key"] for f in fields
        if not (
            _is_question_fill_key(f["key"])
            or f["key"] in _APPEND_NO_COLON_KEYS
            or f["key"] in _OTHER_USED_HISTORY
        )
    ]
    assert not missing, f"Field chưa đăng ký vào whitelist điền dòng: {missing}"


class _FakePara:
    def __init__(self, text: str):
        self.text = text


def test_fill_blank_line_history_block_yes_on_first_blank_line_details_below():
    """Khách yêu cầu 2026-08-14: đáp án Yes/No nằm ở DÒNG GẠCH NGANG ĐẦU TIÊN sau câu hỏi, KHÔNG
    nối vào cuối câu hỏi/dòng chú thích nữa — chi tiết (nếu có) xuống các dòng kế tiếp."""
    from app.services.export_ds260 import _fill_blank_line_history_block

    paras = [
        _FakePara("Have you ever been in the U.S? (Bạn đã từng đến nước MỸ chưa?)"),
        _FakePara("(Yes or No, if 'Yes' write details below) (...)"),
        _FakePara("                    "),
        _FakePara("                    "),
        _FakePara("                    "),
        _FakePara("                    "),
    ]
    values = {"been_in_us": "Yes", "us_travel_history": "Da den My ngay 01/01/2018, o 30 ngay"}
    _fill_blank_line_history_block(paras, 0, "been_in_us", values)

    assert paras[0].text == "Have you ever been in the U.S? (Bạn đã từng đến nước MỸ chưa?)"
    assert paras[1].text == "(Yes or No, if 'Yes' write details below) (...)"
    assert paras[2].text == "Yes"
    assert paras[3].text == "Da den My ngay 01/01/2018, o 30 ngay"
    assert paras[4].text.strip() == ""
    assert paras[5].text.strip() == ""


def test_fill_blank_line_history_block_no_on_first_blank_line_rest_stays_empty():
    from app.services.export_ds260 import _fill_blank_line_history_block

    paras = [
        _FakePara("Have you ever been in the U.S? (Bạn đã từng đến nước MỸ chưa?)"),
        _FakePara("(Yes or No, if 'Yes' write details below) (...)"),
        _FakePara("                    "),
        _FakePara("                    "),
        _FakePara("                    "),
        _FakePara("                    "),
    ]
    values = {"been_in_us": "No"}
    _fill_blank_line_history_block(paras, 0, "been_in_us", values)

    assert paras[0].text == "Have you ever been in the U.S? (Bạn đã từng đến nước MỸ chưa?)"
    assert paras[1].text == "(Yes or No, if 'Yes' write details below) (...)"
    assert paras[2].text == "No"
    assert paras[3].text.strip() == ""
    assert paras[4].text.strip() == ""
    assert paras[5].text.strip() == ""


def test_address_history_semicolon_joined_reflows_to_one_entry_per_line():
    """Báo lỗi thực tế 2026-08-17: other_addresses_history đến từ nguồn dữ liệu bị dồn thành 1
    dòng, phân cách bằng "; " thay vì \\n thật — mỗi mốc thời gian phải nằm 1 dòng riêng khi xuất
    Word, không được dồn chung thành 1 khối text dài."""
    from app.services.export_ds260 import _fill_blank_line_history_block, _prepare_display_values

    raw_history = (
        "Jan 2016–Apr 2016: 22/2A Ngo Duc Ke, Ward 12, Tan Binh Dist, Ho Chi Minh City, "
        "Vietnam 700000; Aug 2004–Jan 2016: 367/12 Le Duc Tho, Ward 16, Go Vap Dist, "
        "Ho Chi Minh City, Vietnam 700000; Jun 1980–Aug 2004: 151 Phan Dinh Phung, "
        "Nguyen Nghiem Ward, Quang Ngai, Vietnam 570000"
    )
    values = {"other_addresses_used": "Yes", "other_addresses_history": raw_history}
    display = _prepare_display_values(values)
    assert display["other_addresses_history"].count("\n") == 2

    paras = [
        _FakePara("Have you lived anywhere other than this address since the age of sixteen?"),
        _FakePara("(Yes or No, if 'Yes' write details below) (...)"),
        _FakePara("                    "),
        _FakePara("                    "),
        _FakePara("                    "),
        _FakePara("                    "),
    ]
    _fill_blank_line_history_block(paras, 0, "other_addresses_used", display)

    assert paras[2].text == "Yes"
    assert paras[3].text.startswith("Jan 2016–Apr 2016:")
    assert paras[4].text.startswith("Aug 2004–Jan 2016:")
    assert paras[5].text.startswith("Jun 1980–Aug 2004:")
