"""
Xuất tài liệu họp báo cáo AI Fill — phòng Docs / quản lý.
Chạy: cd backend && .venv\\Scripts\\python.exe scripts/export_meeting_brief_docx.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "Bao_cao_hop_AI_Fill_Phong_Docs.docx"


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, *, bold: bool = False) -> None:
    r = doc.add_paragraph().add_run(text)
    r.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Bìa ---
    t = doc.add_heading("BÁO CÁO HỌP", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("AI Fill tài liệu DS-260 — Tình trạng triển khai & đề xuất phối hợp Phòng Docs")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    meta = doc.add_paragraph(f"Ngày chuẩn bị: {date.today().strftime('%d/%m/%Y')} · ImmiFill / Edupath IT")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_page_break()

    # --- Mục lục ---
    _h(doc, "Mục lục", 1)
    _bullets(
        doc,
        [
            "1. Tóm tắt cho lãnh đạo (30 giây)",
            "2. Tình trạng triển khai hiện tại",
            "3. Cấu trúc hệ thống ImmiFill",
            "4. Luồng nghiệp vụ & mapping DS-260",
            "5. Phạm vi đã làm (hồ sơ đơn + hồ sơ gia đình)",
            "6. Vướng mắc phối hợp với Phòng Docs",
            "7. Đề xuất quy trình Phòng Docs cần có",
            "8. Yêu cầu cụ thể gửi Phòng Docs trước khi IT xử lý",
            "9. Kịch bản trình bày trong cuộc họp",
            "10. Việc cần chốt sau họp",
            "Phụ lục A — Bảng đặt tên file chuẩn",
            "Phụ lục B — Các trường hợp hồ sơ DS-260",
        ],
    )
    doc.add_page_break()

    # --- 1. Tóm tắt ---
    _h(doc, "1. Tóm tắt cho lãnh đạo (30 giây)", 1)
    _p(
        doc,
        "IT đã đưa code ImmiFill lên môi trường test. Hệ thống upload giấy tờ → OCR (AI) → "
        "mapping DS-260 → review → xuất Word. Bảng mapping 133+ trường khớp yêu cầu ban đầu. "
        "Đã mở rộng từ hồ sơ 1 người sang hồ sơ gia đình (chủ hồ sơ + vợ/chồng + tối đa 4 con).",
    )
    _p(doc, "Điểm nghẽn hiện tại không nằm ở code mà ở phối hợp đầu vào:", bold=True)
    _bullets(
        doc,
        [
            "Phòng Docs chưa có quy trình nhập liệu / đặt tên file thống nhất.",
            "Hồ sơ gửi sang IT thiếu giải thích từng bộ, từng trường hợp.",
            "IT phải tự đoán, tự test, tự làm tài liệu hướng dẫn cho Docs.",
            "Phối hợp họp mang tính hình thức — thiếu phản hồi vòng lặp khi gặp case thực tế.",
        ],
    )
    _p(
        doc,
        "Đề xuất: Phòng Docs cam kết quy chuẩn đặt tên file + checklist từng bộ hồ sơ; "
        "IT tiếp tục hoàn thiện kỹ thuật. Hai bên họp định kỳ 30 phút/tuần trong giai đoạn pilot.",
    )

    # --- 2. Tình trạng ---
    _h(doc, "2. Tình trạng triển khai hiện tại", 1)
    _table(
        doc,
        ["Hạng mục", "Trạng thái", "Ghi chú"],
        [
            ["Code & deploy test", "Đã có", "Docker production — immifill.immipath.org.vn"],
            ["OCR giấy tờ (OpenAI)", "Đã có", "Passport, GKS, lý lịch, kết hôn, ly hôn, GKS con, DS-260 worksheet"],
            ["Mapping DS-260", "Khớp yêu cầu", "133+ trường — file DS260_Mapping_Reference_new.xlsx"],
            ["Hồ sơ 1 người", "Chạy được", "Luồng gốc — test với nhiều case"],
            ["Hồ sơ gia đình", "Mới bổ sung", "01–06: chủ, vợ, 4 con — DS-260 riêng từng người"],
            ["Xung đột dữ liệu", "Đã có", "Luồng 1 vs _new vs worksheet — chọn trên Review"],
            ["Chỉnh sửa & Lưu", "Đã có", "Sửa từng ô DS-260 trước khi xuất Word"],
            ["Tài liệu hướng dẫn user", "Đã có", "Huong_dan_su_dung_ImmiFill_DS260.docx"],
            ["Quy trình Phòng Docs", "Chưa có", "Đây là vướng mắc chính"],
        ],
    )

    # --- 3. Cấu trúc ---
    _h(doc, "3. Cấu trúc hệ thống ImmiFill", 1)
    _p(doc, "Kiến trúc 3 tầng:", bold=True)
    flow = (
        "┌─────────────────────────────────────────┐\n"
        "│  FRONTEND (Next.js)                     │\n"
        "│  Dashboard · Upload · Review · Export   │\n"
        "└──────────────────┬──────────────────────┘\n"
        "                   │ API /api/v1 (JWT)\n"
        "┌──────────────────▼──────────────────────┐\n"
        "│  BACKEND (FastAPI / Python)             │\n"
        "│  OCR → DocRecord → DS-260 → Word        │\n"
        "└──────────────────┬──────────────────────┘\n"
        "                   │\n"
        "┌──────────────────▼──────────────────────┐\n"
        "│  Database + File storage                │\n"
        "│  PostgreSQL · uploads/ · exports/       │\n"
        "│  Mapping: ds260_mapping.json            │\n"
        "└─────────────────────────────────────────┘"
    )
    r = doc.add_paragraph().add_run(flow)
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    _h(doc, "3.1. Module chính", 2)
    _table(
        doc,
        ["Module", "Chức năng", "Người dùng"],
        [
            ["Dashboard", "Tạo hồ sơ, danh sách, xóa, thùng rác", "Docs / IT"],
            ["Upload", "Kéo thả PDF, gán mã file 01_1…", "Docs"],
            ["Review — DS-260", "Xem mapping, sửa, giải quyết xung đột, xuất Word", "Docs"],
            ["Review — OCR", "Sửa dữ liệu OCR từng file", "Docs / IT"],
            ["Thành viên gia đình", "Khai báo chủ / vợ / con", "Docs"],
            ["Export Word", "DS-260 theo mẫu form chuẩn", "Docs"],
        ],
    )

    _h(doc, "3.2. Luồng dữ liệu kỹ thuật", 2)
    _numbered(
        doc,
        [
            "Upload PDF → phân loại giấy tờ (passport, birth_certificate, …).",
            "OpenAI OCR trích xuất field (full_name, date_of_birth, father_name…).",
            "Lưu ApplicantDocRecord — mỗi loại giấy tách riêng, không gộp.",
            "resolve_ds260_form() — mapping field DS-260 ← document + OCR field.",
            "Review: kiểm tra, giải quyết xung đột, chỉnh tay.",
            "export_ds260 → điền file Word template.",
        ],
    )

    # --- 4. Mapping ---
    _h(doc, "4. Luồng nghiệp vụ & mapping DS-260", 1)
    _p(
        doc,
        "Mỗi ô form DS-260 được định nghĩa rõ: field key → loại giấy tờ → field OCR. "
        "Ví dụ: applicant_name ← passport ← full_name; father_full_name ← birth_certificate ← father_name.",
    )
    _table(
        doc,
        ["Loại giấy (Luồng 1)", "Mã file", "Mục DS-260"],
        [
            ["PASSPORT", "01_2", "A.1 Cá nhân, A.2 Hộ chiếu"],
            ["BIRTH CERTIFICATE", "01_1", "Cha, Mẹ, thông tin GKS"],
            ["JUDICIAL CERTIFICATE", "01_4", "Lý lịch tư pháp"],
            ["MARRIAGE CERTIFICATE", "01_5", "Phối ngẫu"],
            ["DIVORCE DECREE", "01_3", "Ly hôn"],
            ["BIRTH CERTIFICATE CHILD", "03_1…", "Con + DS-260 con"],
            ["DS260 worksheet", "01_6", "Địa chỉ, SĐT, email, MXH"],
        ],
    )
    _p(doc, "Hai luồng đối chiếu:", bold=True)
    _bullets(
        doc,
        [
            "Luồng 1 (file chính): Passport, GKS… — không hậu tố _new.",
            "Đối chiếu (_new / worksheet): so sánh với Luồng 1 → báo xung đột nếu khác.",
            "User chọn giá trị đúng trên Review trước khi xuất Word.",
        ],
    )

    # --- 5. Phạm vi ---
    _h(doc, "5. Phạm vi đã làm", 1)
    _h(doc, "5.1. Hồ sơ đơn (1 người)", 2)
    _bullets(
        doc,
        [
            "Tất cả file gán mã 01_x.",
            "Một file DS-260 Word cho chủ hồ sơ.",
            "Đủ mục: cá nhân, HC, cha/mẹ, địa chỉ, phối ngẫu, con (nếu có).",
        ],
    )
    _h(doc, "5.2. Hồ sơ gia đình (mới — phát sinh sau phase 1)", 2)
    _bullets(
        doc,
        [
            "Khai báo thành viên: 01 chủ · 02 vợ/chồng · 03–06 con (tối đa 4 con).",
            "Đặt tên file theo từng người: 01_2 PASSPORT - TEN CHU.pdf, 03_1 BIRTH CERTIFICATE CHILD - TEN CON.pdf.",
            "Xuất DS-260 Word riêng cho từng người.",
            "Cha/mẹ trên DS-260 con: từ GKS con hoặc fallback chủ hồ sơ + phối ngẫu.",
            "Case test: HỒ CÔNG BẢO LONG — vợ VĂN THỊ HƯỜNG — 3 con.",
        ],
    )
    _p(
        doc,
        "Lưu ý cho họp: Phase 1 chỉ thiết kế 1 người/hồ sơ. Phase 2 IT tự mở rộng gia đình "
        "vì nghiệp vụ thực tế yêu cầu — cần Phòng Docs cung cấp case mẫu đầy đủ để validate.",
        bold=True,
    )

    # --- 6. Vướng mắc ---
    _h(doc, "6. Vướng mắc phối hợp với Phòng Docs", 1)
    _p(doc, "Diễn giải chuyên nghiệp các điểm đã ghi nhận (để trình bày trong họp):", bold=True)
    issues = [
        (
            "Thiếu quy trình nhập liệu",
            "Chưa có SOP: ai upload, ai review, ai xuất Word, ai gửi khách. IT không thể tự định nghĩa quy trình nghiệp vụ.",
        ),
        (
            "Không có quy chuẩn đặt tên file",
            "Mỗi người đặt tên khác nhau → OCR gán sai người, mapping sai. IT đã viết chuẩn 01_1… nhưng Docs chưa áp dụng.",
        ),
        (
            "Thiếu mô tả từng bộ hồ sơ",
            "Gửi file không kèm: chủ hồ sơ là ai, có vợ/con không, case đặc biệt gì. IT mất thời gian đoán.",
        ),
        (
            "Không có danh mục trường hợp",
            "Chưa liệt kê: đơn / vợ chồng / con đi cùng / ly hôn / tái hôn… IT phải tự suy ra từ file lỗi.",
        ),
        (
            "Tham gia họp hình thức",
            "Thiếu hỏi–đáp, thiếu feedback khi test. Khó biết Docs vướng chỗ nào trên hồ sơ thật.",
        ),
        (
            "IT tự làm tài liệu",
            "Hướng dẫn user, mapping Excel, quy tắc đặt tên — IT soạn vì Docs chưa có bản chuẩn.",
        ),
        (
            "Phát sinh hồ sơ gia đình",
            "Yêu cầu mới so với thiết kế ban đầu (1 người/1 hồ sơ). Cần Docs xác nhận case và file mẫu.",
        ),
    ]
    rows = [[a, b] for a, b in issues]
    _table(doc, ["Vấn đề", "Ảnh hưởng / Cách nói trong họp"], rows)

    # --- 7. Đề xuất quy trình ---
    _h(doc, "7. Đề xuất quy trình Phòng Docs cần có", 1)
    _numbered(
        doc,
        [
            "Nhận hồ sơ khách → lập checklist giấy tờ (theo loại case).",
            "Đổi tên file theo chuẩn 01_1 … 06_4 (xem Phụ lục A) trước khi gửi IT/upload.",
            "Điền phiếu mô tả bộ hồ sơ (mẫu IT đề xuất bên dưới) — 1 trang/bộ.",
            "Upload lên ImmiFill → chờ OCR xong.",
            "Review DS-260: giải quyết xung đột → sửa ô sai → Lưu từng trường.",
            "Xuất Word từng người (gia đình) → kiểm tra lần cuối → gửi khách/LSQ.",
            "Báo lỗi case đặc biệt cho IT (form + file mẫu) — không chỉ «không được».",
        ],
    )

    _h(doc, "7.1. Mẫu phiếu mô tả bộ hồ sơ (Docs điền)", 2)
    _table(
        doc,
        ["Mục", "Nội dung Docs điền"],
        [
            ["Mã hồ sơ / Tên khách", ""],
            ["Loại case", "Đơn / Gia đình / EB-3 / …"],
            ["Chủ hồ sơ (01)", "Họ tên"],
            ["Phối ngẫu (02)", "Có/Không — Họ tên"],
            ["Con (03–06)", "Danh sách tên con"],
            ["Đặc biệt", "Ly hôn, tái hôn, mất cha/mẹ, …"],
            ["File đã đủ?", "GKS, HC, lý lịch, kết hôn, DS-260 worksheet…"],
            ["Người xử lý Docs", ""],
        ],
    )

    # --- 8. Yêu cầu ---
    _h(doc, "8. Yêu cầu cụ thể gửi Phòng Docs", 1)
    _bullets(
        doc,
        [
            "Cam kết đặt tên file theo chuẩn (tài liệu Huong_dan_su_dung_ImmiFill_DS260.docx).",
            "Cung cấp tối thiểu 3 bộ hồ sơ mẫu đầy đủ: (1) đơn, (2) vợ chồng, (3) gia đình có con.",
            "Mỗi bộ kèm phiếu mô tả — không gửi «file zip không chú thích».",
            "Chỉ định 1–2 nhân sự Docs pilot — họp 30 phút/tuần với IT.",
            "Phản hồi bằng văn bản: field nào sai, file nào thiếu, case nào chưa hỗ trợ.",
            "Không kỳ vọng IT tự biết nghiệp vụ hồ sơ khi Docs chưa mô tả.",
        ],
    )

    # --- 9. Kịch bản họp ---
    _h(doc, "9. Kịch bản trình bày trong cuộc họp (~15–20 phút)", 1)
    _table(
        doc,
        ["Thời gian", "Nội dung", "Người nói"],
        [
            ["2 phút", "Mục tiêu: AI fill DS-260 — giảm nhập tay, không thay thế chuyên môn Docs", "IT / Anh"],
            ["3 phút", "Demo nhanh: Upload → Review → Xuất Word (1 hồ sơ đơn)", "IT"],
            ["3 phút", "Demo gia đình: 5 người — tab từng người — xuất 5 file Word", "IT"],
            ["2 phút", "Mapping khớp yêu cầu — show Excel / ô Review", "IT"],
            ["5 phút", "7 vướng mắc phối hợp — cần Docs làm gì", "IT"],
            ["3 phút", "Chốt: quy chuẩn file + phiếu mô tả + lịch pilot", "Anh / Trưởng Docs"],
            ["2 phút", "Hỏi đáp", "Tất cả"],
        ],
    )

    _h(doc, "9.1. Câu nói gợi ý (tự nhiên, lịch sự)", 2)
    _bullets(
        doc,
        [
            "«Code test đã chạy, mapping DS-260 khớp spec ban đầu. Phần còn lại là đầu vào từ Phòng Docs.»",
            "«IT không nắm nghiệp vụ từng bộ hồ sơ — cần Docs mô tả case trước khi gửi file.»",
            "«Đã bổ sung hồ sơ gia đình vì thực tế có — cần Docs confirm bằng file mẫu thật.»",
            "«Đề nghị 3 bộ mẫu + quy chuẩn đặt tên — IT hỗ trợ training 1 buổi.»",
            "«Hợp tác cần hai chiều: Docs test và báo lỗi cụ thể, không chỉ dùng kết quả.»",
        ],
    )

    # --- 10. Chốt ---
    _h(doc, "10. Việc cần chốt sau họp", 1)
    _table(
        doc,
        ["#", "Việc", "Bên", "Thời hạn"],
        [
            ["1", "Ban hành quy chuẩn đặt tên file (áp dụng chung)", "Docs + IT", "1 tuần"],
            ["2", "Mẫu phiếu mô tả bộ hồ sơ", "IT soạn — Docs duyệt", "3 ngày"],
            ["3", "Gửi 3 bộ hồ sơ mẫu (đơn / vợ chồng / gia đình)", "Docs", "2 tuần"],
            ["4", "Chỉ định nhân sự pilot + lịch họp tuần", "Docs", "Trong họp"],
            ["5", "IT fix bug từ feedback pilot (cha/mẹ con, worksheet…)", "IT", "Liên tục"],
            ["6", "Báo cáo tiến độ lần 2", "IT", "Sau 4 tuần pilot"],
        ],
    )

    doc.add_page_break()

    # --- Phụ lục A ---
    _h(doc, "Phụ lục A — Quy chuẩn đặt tên file (tóm tắt)", 1)
    _p(doc, "Format: {mã}_{số} {LOẠI GIẤY} - {HỌ TÊN}.pdf")
    _table(
        doc,
        ["Mã", "Người", "File chuẩn"],
        [
            ["01", "Chủ hồ sơ", "01_1 GKS · 01_2 HC · 01_3 ly hôn · 01_4 lý lịch · 01_5 kết hôn · 01_6 DS260"],
            ["02", "Vợ/chồng", "02_1 … 02_4"],
            ["03–06", "Con 1–4", "03_1 BIRTH CERTIFICATE CHILD · 03_2 PASSPORT …"],
        ],
    )
    _p(doc, "Chi tiết: file Huong_dan_su_dung_ImmiFill_DS260.docx và DS260_Mapping_Reference_new.xlsx")

    # --- Phụ lục B ---
    _h(doc, "Phụ lục B — Các trường hợp hồ sơ DS-260", 1)
    _table(
        doc,
        ["Case", "Thành viên", "File cần", "Xuất Word"],
        [
            ["Hồ sơ đơn", "01", "01_1–01_4 (+05,06 nếu có)", "1 file"],
            ["Vợ chồng đi cùng", "01 + 02", "01_x + 02_x + giấy kết hôn", "2 file"],
            ["Gia đình + con", "01–06", "GKS con = BIRTH CERTIFICATE CHILD", "1 file/người"],
            ["Đã ly hôn", "01", "01_3 DIVORCE DECREE", "1 file (+ con nếu có)"],
            ["Worksheet thiếu địa chỉ", "01", "01_6 DS260", "Bổ sung trên Review"],
        ],
    )

    doc.add_paragraph()
    foot = doc.add_paragraph("— Hết tài liệu họp — IT Edupath / ImmiFill —")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
