"""
Xuất mẫu Phiếu mô tả bộ hồ sơ — Docs điền mỗi lần gửi hồ sơ cho IT.
Chạy: cd backend && .venv\\Scripts\\python.exe scripts/export_phieu_mo_ta_docx.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "Phieu_mo_ta_bo_ho_so_mau.docx"


def _blank_line(doc: Document, label: str, lines: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run("_" * 72)
    for _ in range(lines - 1):
        doc.add_paragraph("_" * 80)


def _table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PHIẾU MÔ TẢ BỘ HỒ SƠ")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("ImmiFill — AI fill DS-260 | Phòng Docs điền mỗi lần gửi hồ sơ").italic = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "In 01 tờ / bộ hồ sơ — kèm mail hoặc đặt trong folder ZIP cùng các file PDF đã đặt tên chuẩn."
    )

    doc.add_paragraph()

    # --- Thông tin chung ---
    p = doc.add_paragraph()
    p.add_run("A. THÔNG TIN CHUNG").bold = True

    _table(
        doc,
        ["Mục", "Nội dung (Docs điền)"],
        [
            ["Ngày gửi", ""],
            ["Người gửi (Phòng Docs)", ""],
            ["Mã / tên hồ sơ trên ImmiFill", ""],
            ["Email / SĐT liên hệ Docs", ""],
        ],
        col_widths=[2.2, 4.3],
    )

    # --- Loại case ---
    p = doc.add_paragraph()
    p.add_run("B. LOẠI BỘ HỒ SƠ (chọn 01)").bold = True
    doc.add_paragraph(
        "☐ Case A — Đơn (1 người)     "
        "☐ Case B — Vợ chồng (không con)     "
        "☐ Case C — Gia đình (chủ + vợ + con)"
    )
    doc.add_paragraph(
        "☐ Case D — Chủ + con (không vợ đi cùng)     "
        "☐ Case E — Đã ly hôn / tái hôn     "
        "☐ Case F — Case đặc biệt"
    )

    # --- Thành viên ---
    p = doc.add_paragraph()
    p.add_run("C. DANH SÁCH THÀNH VIÊN TRONG BỘ").bold = True
    _table(
        doc,
        ["Mã", "Vai trò", "Họ tên (viết hoa, không dấu trên tên file)", "Ghi chú"],
        [
            ["01", "Chủ hồ sơ", "", ""],
            ["02", "Phối ngẫu", "☐ Có   ☐ Không", "Nếu không có: ghi Không, bỏ qua file 02_x"],
            ["03", "Con 1", "☐ Có   ☐ Không", ""],
            ["04", "Con 2", "☐ Có   ☐ Không", ""],
            ["05", "Con 3", "☐ Có   ☐ Không", ""],
            ["06", "Con 4", "☐ Có   ☐ Không", ""],
        ],
        col_widths=[0.6, 1.1, 2.8, 1.8],
    )

    # --- Danh sách file ---
    p = doc.add_paragraph()
    p.add_run("D. DANH SÁCH FILE PDF GỬI KÈM (đã đặt tên chuẩn)").bold = True
    doc.add_paragraph(
        "Format: {mã}_{số} {LOẠI GIẤY} - {HỌ TÊN}.pdf   "
        "Ví dụ: 01_2 PASSPORT - HO CONG BAO LONG.pdf"
    )
    _table(
        doc,
        ["STT", "Tên file (điền đúng tên trên máy)", "Người", "Loại giấy"],
        [[str(i), "", "", ""] for i in range(1, 16)],
        col_widths=[0.4, 3.2, 1.0, 1.7],
    )
    doc.add_paragraph("Nếu nhiều hơn 15 file: thêm dòng hoặc kèm sheet Excel.")

    # --- Case đặc biệt ---
    p = doc.add_paragraph()
    p.add_run("E. GIẤY TỜ ĐẶC BIỆT / CASE LẠ (nếu có)").bold = True
    doc.add_paragraph(
        "Ghi rõ bằng chữ — IT cần biết TRƯỚC khi xử lý. Ví dụ: cha/mẹ mất, tên khác HC/GKS, "
        "con nuôi, nhiều lần ly hôn…"
    )
    for _ in range(4):
        doc.add_paragraph("_" * 80)

    # --- Xác nhận ---
    p = doc.add_paragraph()
    p.add_run("F. XÁC NHẬN TRƯỚC KHI GỬI").bold = True
    checks = [
        "☐ Tất cả PDF đã đổi tên đúng format (không dùng scan001.pdf, giay to.pdf…)",
        "☐ Giấy khai sinh CON có từ BIRTH CERTIFICATE CHILD trong tên file",
        "☐ Tên trên file khớp tên trên phiếu (mục C và D)",
        "☐ Đủ giấy từng người: GKS, HC, lý lịch (theo loại case)",
        "☐ Có worksheet DS260 (01_6) cho địa chỉ, SĐT, email",
        "☐ Case đặc biệt đã ghi chú mục E",
    ]
    for c in checks:
        doc.add_paragraph(c)

    _table(
        doc,
        ["Mục", "Nội dung"],
        [
            ["Số file PDF gửi kèm", ""],
            ["Đã đặt tên file chuẩn?", "☐ Có   ☐ Chưa — IT từ chối xử lý nếu chưa"],
            ["Người xử lý Docs sau khi IT upload", ""],
            ["Ghi chú thêm cho IT", ""],
        ],
        col_widths=[2.5, 4.0],
    )

    # Chữ ký
    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.style = "Table Grid"
    sig.rows[0].cells[0].text = "Người lập phiếu (Docs)\n\n\nKý / Họ tên: _______________"
    sig.rows[0].cells[1].text = "Ngày: ___ / ___ / ______"

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.add_run(f"Mẫu ImmiFill — cập nhật {date.today().strftime('%d/%m/%Y')}").italic = True

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
