"""
Xuất hướng dẫn sử dụng ImmiFill cho người dùng (Word .docx).
Chạy: cd backend && .venv\\Scripts\\python.exe scripts/export_user_guide_docx.py
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASE = Path(__file__).resolve().parents[1]
DATA = BASE / "data" / "doc_schemas"
OUT = BASE.parent / "Huong_dan_su_dung_ImmiFill_DS260.docx"

DOC_LABELS: dict[str, str] = {
    "passport": "Hộ chiếu (PASSPORT)",
    "birth_certificate": "Giấy khai sinh (BIRTH CERTIFICATE)",
    "birth_certificate_child": "Giấy khai sinh con (BIRTH CERTIFICATE CHILD)",
    "marriage_certificate": "Giấy kết hôn (MARRIAGE CERTIFICATE)",
    "divorce": "Quyết định ly hôn (DIVORCE DECREE)",
    "judicial_certificate": "Lý lịch tư pháp (JUDICIAL CERTIFICATE)",
    "death_certificate": "Giấy báo tử",
    "military_discharge": "Giấy xuất ngũ",
    "ds260_customer_form": "DS-260 khách khai (DS260)",
    "address_document": "Giấy chứng minh địa chỉ",
    "spouse_applicant_profile": "Hồ sơ phối ngẫu",
}

UPLOAD_TO_DS260: list[list[str]] = [
    ["01_2 PASSPORT", "Hộ chiếu", "A.1 Cá nhân, A.2 Hộ chiếu"],
    ["01_1 BIRTH CERTIFICATE", "GKS chủ hồ sơ", "Cha, Mẹ, thông tin GKS"],
    ["01_4 JUDICIAL CERTIFICATE", "Lý lịch tư pháp", "Mục lý lịch tư pháp"],
    ["01_5 MARRIAGE CERTIFICATE", "Giấy kết hôn", "Phối ngẫu (B.3)"],
    ["01_3 DIVORCE DECREE", "Ly hôn", "Ly hôn, phối ngẫu cũ"],
    ["01_6 DS260", "Worksheet khách", "Địa chỉ, SĐT, email, MXH"],
    ["03_1 BIRTH CERTIFICATE CHILD", "GKS con", "DS-260 con + danh sách con (chủ hồ sơ)"],
]


def _load_ds260_mapping() -> dict:
    with (DATA / "ds260_mapping.json").open(encoding="utf-8") as f:
        return json.load(f)


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def _p(doc: Document, text: str, *, bold: bool = False) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Lưu ý: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB7, 0x5E, 0x00)
    p.add_run(text)


def _tip(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run("Mẹo: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.add_run(text)


def _add_section_ds260_mapping(doc: Document, mapping: dict) -> None:
    """Mục 10 — Cách mapping DS-260: giấy tờ → OCR → form Word."""
    doc.add_page_break()
    _h1(doc, "10. Mapping DS-260 — giấy tờ điền vào form")
    _p(
        doc,
        "Mapping là quy tắc hệ thống dùng để biết mỗi ô trên form DS-260 lấy dữ liệu từ "
        "file giấy tờ nào và trường OCR nào. Nhân viên không cần cấu hình mapping — "
        "chỉ cần upload đúng loại giấy, đặt tên file chuẩn, và kiểm tra kết quả trên Review.",
    )

    _h2(doc, "10.1. Ba bước mapping")
    _numbered(
        doc,
        [
            "Upload PDF → hệ thống OCR đọc giấy tờ (OpenAI) → lưu các trường như full_name, date_of_birth, father_name…",
            "Bảng mapping quy định: ô DS-260 «Họ tên» ← document passport ← field full_name.",
            "Trang Review hiển thị giá trị + nguồn; xuất Word điền vào đúng nhãn trên mẫu form.",
        ],
    )

    _h2(doc, "10.2. Sơ đồ luồng dữ liệu")
    flow = (
        "File PDF (01_2 PASSPORT - TEN.pdf)\n"
        "    ↓ OCR\n"
        "Trường OCR: full_name, passport_number, date_of_birth…\n"
        "    ↓ Mapping DS-260\n"
        "Ô form: Applicant Name, Passport Number, Date of Birth…\n"
        "    ↓ Xuất Word\n"
        "File .docx DS-260 hoàn chỉnh"
    )
    p = doc.add_paragraph()
    run = p.add_run(flow)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    _h2(doc, "10.3. Công thức mapping")
    _p(doc, "Mỗi trường DS-260 được định nghĩa bởi:", bold=True)
    _bullets(
        doc,
        [
            "key — mã trường nội bộ (vd. applicant_name, father_full_name)",
            "document — loại giấy tờ nguồn chính (passport, birth_certificate, ds260_customer_form…)",
            "field — tên trường OCR trên giấy đó (vd. full_name, father_name)",
            "aliases — tên thay thế OCR có thể dùng (vd. name, dob, father_full_name)",
            "derive — quy tắc tách/suy ra (vd. country_from_location từ nơi sinh)",
        ],
    )
    _p(doc, "Ví dụ:", bold=True)
    _p(doc, "applicant_name ← passport ← full_name (hoặc alias: name)")
    _p(doc, "father_full_name ← birth_certificate ← father_name (GKS của chính người đang xuất DS-260)")
    _p(doc, "current_address ← ds260_customer_form ← current_address (file DS260 worksheet)")

    _h2(doc, "10.4. Giấy tờ upload → mục DS-260")
    _table(doc, ["File upload (ví dụ)", "Loại giấy", "Mục DS-260 được điền"], UPLOAD_TO_DS260)

    _h2(doc, "10.5. Mapping theo từng mục form")
    sections = mapping.get("sections", [])
    sec_rows: list[list[str]] = []
    for sec in sections:
        fields = sec.get("fields", [])
        if not fields:
            continue
        docs: set[str] = set()
        for f in fields:
            d = f.get("document", "")
            if d and d != "spouse_applicant_profile":
                docs.add(DOC_LABELS.get(d, d))
        sec_rows.append(
            [
                sec.get("title", sec.get("id", "")),
                sec.get("subtitle", ""),
                str(len(fields)),
                "; ".join(sorted(docs)) if docs else "—",
            ]
        )
    _table(
        doc,
        ["Mục DS-260", "Nguồn chính", "Số trường", "Giấy tờ liên quan"],
        sec_rows,
    )

    _h2(doc, "10.6. Ví dụ mapping chi tiết (một số trường quan trọng)")
    examples = [
        ("applicant_name", "Họ và tên", "passport", "full_name", "01_2 PASSPORT"),
        ("date_of_birth", "Ngày sinh", "passport", "date_of_birth", "01_2 PASSPORT"),
        ("passport_number", "Số hộ chiếu", "passport", "passport_number", "01_2 PASSPORT"),
        ("father_full_name", "Họ tên cha", "birth_certificate", "father_name", "01_1 BIRTH CERTIFICATE"),
        ("mother_full_name", "Họ tên mẹ", "birth_certificate", "mother_name", "01_1 BIRTH CERTIFICATE"),
        ("current_address", "Địa chỉ hiện tại", "ds260_customer_form", "current_address", "01_6 DS260"),
        ("primary_phone", "SĐT chính", "ds260_customer_form", "primary_phone_number", "01_6 DS260"),
        ("spouse_full_name", "Họ tên phối ngẫu", "marriage_certificate", "spouse_full_name", "01_5 MARRIAGE CERTIFICATE"),
        ("child_1_full_name", "Tên con 1", "birth_certificate_child", "child_full_name", "03_1 BIRTH CERTIFICATE CHILD"),
        ("judicial_certificate_number", "Số lý lịch", "judicial_certificate", "document_number", "01_4 JUDICIAL CERTIFICATE"),
    ]
    _table(
        doc,
        ["Field DS-260", "Nhãn", "Giấy tờ", "Field OCR", "File mẫu"],
        [[e[0], e[1], DOC_LABELS.get(e[2], e[2]), e[3], e[4]] for e in examples],
    )

    _h2(doc, "10.7. Mapping hồ sơ gia đình")
    _table(
        doc,
        ["Người", "Giấy tờ dùng mapping", "Ghi chú"],
        [
            ["Chủ hồ sơ (01)", "01_x — passport, GKS, lý lịch của chủ hộ", "Cha/mẹ lấy từ 01_1, không lấy GKS con"],
            ["Phối ngẫu (02)", "02_x — GKS, HC phối ngẫu", "Xuất DS-260 riêng — mục B.3"],
            ["Con (03–06)", "03_x … — BIRTH CERTIFICATE CHILD + HC con", "Cha/mẹ trên DS-260 con = tên trên GKS con"],
            ["Giấy chung", "Giấy kết hôn, ly hôn — gán chủ hồ sơ", "01_5 hoặc không prefix"],
        ],
    )

    _h2(doc, "10.8. Đọc nguồn mapping trên Review")
    _bullets(
        doc,
        [
            "Mỗi ô DS-260 có cột «Nguồn» — cho biết file và loại giấy.",
            "derived=doc_scan_fill — điền từ OCR giấy tờ được phép.",
            "derived=reference_cross_fill — bổ sung từ file _new đối chiếu.",
            "derived=ds260_worksheet_fill — từ worksheet DS-260 khách khai.",
            "derived=conflict_resolution — giá trị user đã chọn khi giải quyết xung đột.",
            "Chỉnh tay trên Review — ưu tiên cao nhất khi xuất Word.",
        ],
    )

    _h2(doc, "10.9. Thứ tự ưu tiên khi nhiều nguồn")
    _table(
        doc,
        ["Bước", "Nguồn", "Mô tả"],
        [
            ["1", "Luồng 1 (file chính)", "01_2 PASSPORT, 01_1 GKS… — không có _new"],
            ["2", "Đối chiếu (_new)", "Bổ sung trường Luồng 1 thiếu"],
            ["3", "Worksheet DS-260", "Địa chỉ, liên lạc — sau giấy tờ chính"],
            ["4", "Xung đột đã chọn", "User chọn A hoặc B trên Review"],
            ["5", "Chỉnh tay", "User sửa trực tiếp — ghi đè tất cả"],
        ],
    )

    _note(
        doc,
        "Bảng mapping đầy đủ xem Phụ lục C trong tài liệu này, hoặc file Excel "
        "DS260_Mapping_Reference_new.xlsx (sheet «DS260 Mapping»).",
    )


def _add_appendix_mapping_detail(doc: Document, mapping: dict) -> None:
    """Phụ lục C — bảng mapping đầy đủ từ ds260_mapping.json."""
    doc.add_page_break()
    _h1(doc, "Phụ lục C — Bảng mapping DS-260 đầy đủ")
    _p(
        doc,
        "Bảng dưới liệt kê toàn bộ trường DS-260, giấy tờ nguồn và field OCR. "
        "Dùng để đối chiếu khi kiểm tra Review hoặc khi OCR thiếu/sai trường.",
    )
    rows: list[list[str]] = []
    for sec in mapping.get("sections", []):
        sec_title = sec.get("title", "")
        for field in sec.get("fields", []):
            doc_type = field.get("document", "")
            if doc_type == "spouse_applicant_profile":
                continue
            aliases = field.get("aliases") or []
            alias_txt = ", ".join(aliases[:4])
            if len(aliases) > 4:
                alias_txt += "…"
            derive = field.get("derive", "")
            note = f"derive={derive}" if derive else ""
            rows.append(
                [
                    sec_title,
                    field.get("label", field.get("key", "")),
                    field.get("key", ""),
                    DOC_LABELS.get(doc_type, doc_type),
                    field.get("field", ""),
                    alias_txt,
                    note,
                ]
            )
    _table(
        doc,
        ["Mục DS-260", "Nhãn", "Field key", "Giấy tờ", "Field OCR", "Alias OCR", "Ghi chú"],
        rows,
    )
    _p(doc, f"Tổng số trường mapping: {len(rows)}", bold=True)


def build_guide() -> Document:
    mapping = _load_ds260_mapping()
    doc = Document()
    _set_doc_defaults(doc)

    # --- Trang bìa ---
    _title(doc, "HƯỚNG DẪN SỬ DỤNG IMMIFILL")
    sub = doc.add_paragraph("Đặt tên hồ sơ · Upload · Mapping DS-260 · Xung đột · Xuất Word")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()
    ver = doc.add_paragraph(f"Phiên bản tài liệu: {date.today().strftime('%d/%m/%Y')}")
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Mục lục thủ công ---
    _h1(doc, "Mục lục")
    toc = [
        "1. Giới thiệu",
        "2. Quy trình làm việc tổng quan",
        "3. Tạo hồ sơ trên hệ thống",
        "4. Đặt tên file giấy tờ (chuẩn 01–06)",
        "5. Đưa hồ sơ lên hệ thống (Upload)",
        "6. Khai báo thành viên gia đình",
        "7. Trang Review — kiểm tra dữ liệu",
        "8. Giải quyết xung đột dữ liệu",
        "9. Chỉnh sửa thủ công & xuất DS-260 Word",
        "10. Mapping DS-260 — giấy tờ điền vào form",
        "11. Các trường hợp DS-260 theo loại hồ sơ",
        "12. Lỗi thường gặp & cách khắc phục",
        "Phụ lục A — Bảng đặt tên file mẫu",
        "Phụ lục B — Hai luồng dữ liệu (Luồng 1 & đối chiếu)",
        "Phụ lục C — Bảng mapping DS-260 đầy đủ",
    ]
    _bullets(doc, toc)
    doc.add_page_break()

    # --- 1. Giới thiệu ---
    _h1(doc, "1. Giới thiệu")
    _p(
        doc,
        "ImmiFill là hệ thống hỗ trợ xử lý hồ sơ định cư: upload giấy tờ, nhận dạng chữ (OCR), "
        "điền tự động form DS-260 và xuất file Word theo mẫu chuẩn. Tài liệu này hướng dẫn "
        "nhân viên và khách hàng cách đặt tên file, đưa hồ sơ lên hệ thống, xử lý xung đột "
        "và xuất DS-260 cho từng trường hợp.",
    )

    # --- 2. Quy trình ---
    _h1(doc, "2. Quy trình làm việc tổng quan")
    _numbered(
        doc,
        [
            "Đăng nhập ImmiFill → Dashboard.",
            "Tạo hồ sơ mới (nhập tên chủ hồ sơ).",
            "Vào trang Upload → đặt tên file đúng chuẩn → kéo thả hoặc chọn file PDF.",
            "Chờ hệ thống OCR xử lý (trạng thái chuyển sang hoàn tất).",
            "Vào trang Review → khai báo vợ/chồng/con (nếu hồ sơ gia đình).",
            "Kiểm tra bảng DS-260 → giải quyết xung đột (nếu có).",
            "Chỉnh sửa thủ công các trường còn thiếu/sai.",
            "Xuất file Word DS-260 cho từng người (chủ hồ sơ, vợ/chồng, từng con).",
        ],
    )
    _tip(doc, "Luôn đặt tên file trước khi upload — hệ thống dùng tên file để gán đúng người và loại giấy tờ.")

    # --- 3. Tạo hồ sơ ---
    _h1(doc, "3. Tạo hồ sơ trên hệ thống")
    _h2(doc, "3.1. Hồ sơ đơn (một người)")
    _bullets(
        doc,
        [
            "Trên Dashboard, bấm Tạo hồ sơ mới.",
            "Nhập họ tên chủ hồ sơ (viết hoa, không dấu hoặc có dấu đều được — nên thống nhất).",
            "Không cần khai báo thành viên — mọi file sẽ tự gán mã 01.",
        ],
    )
    _h2(doc, "3.2. Hồ sơ gia đình")
    _bullets(
        doc,
        [
            "Tạo hồ sơ với tên chủ hồ sơ (người nộp chính).",
            "Sau khi upload, vào Review → mục Thành viên → thêm Phối ngẫu và các Con.",
            "Lưu tên từng người — hệ thống gán mã 01 (chủ), 02 (vợ/chồng), 03–06 (con 1–4).",
        ],
    )
    _note(
        doc,
        "Tên trên hệ thống phải khớp tên trên giấy tờ (hộ chiếu, giấy khai sinh). "
        "Nếu sai tên, OCR có thể gán nhầm file cho người khác.",
    )

    # --- 4. Đặt tên file ---
    _h1(doc, "4. Đặt tên file giấy tờ (chuẩn 01–06)")
    _h2(doc, "4.1. Công thức đặt tên")
    _p(doc, "Format chuẩn:", bold=True)
    _p(doc, "{mã người}_{số file} {LOẠI GIẤY} - {HỌ TÊN}.pdf")
    _p(doc, "Ví dụ: 01_1 BIRTH CERTIFICATE - DANG VAN HUNG.pdf")

    _h2(doc, "4.2. Mã người")
    _table(
        doc,
        ["Mã", "Vai trò", "Ghi chú"],
        [
            ["01", "Chủ hồ sơ", "Người nộp chính"],
            ["02", "Vợ / chồng", "Chỉ khi có phối ngẫu đi cùng"],
            ["03", "Con 1", ""],
            ["04", "Con 2", ""],
            ["05", "Con 3", ""],
            ["06", "Con 4", "Tối đa 4 con trên hệ thống"],
        ],
    )

    _h2(doc, "4.3. Số file chuẩn (_1 đến _4)")
    _table(
        doc,
        ["Số", "Loại giấy", "Từ khóa trong tên file", "Người lớn", "Con"],
        [
            ["_1", "Giấy khai sinh", "BIRTH CERTIFICATE / BIRTH CERTIFICATE CHILD", "GKS người lớn", "GKS con (bắt buộc có CHILD)"],
            ["_2", "Hộ chiếu", "PASSPORT", "Có", "Có"],
            ["_3", "Quyết định ly hôn", "DIVORCE DECREE", "Nếu có", "Nếu có"],
            ["_4", "Lý lịch tư pháp", "JUDICIAL CERTIFICATE", "Có", "Có"],
        ],
    )

    _h2(doc, "4.4. File bổ sung (_5, _6, _7…)")
    _table(
        doc,
        ["Số", "Loại giấy", "Từ khóa"],
        [
            ["_5", "Giấy kết hôn", "MARRIAGE CERTIFICATE"],
            ["_6", "DS-260 khách khai (worksheet)", "DS260"],
            ["_7+", "Giấy tờ khác", "DOCUMENT hoặc tên mô tả"],
        ],
    )
    _note(
        doc,
        "Giấy khai sinh con phải dùng BIRTH CERTIFICATE CHILD (có từ CHILD), "
        "không dùng BIRTH CERTIFICATE của người lớn. Nếu upload nhầm, thông tin cha/mẹ "
        "trên DS-260 chủ hồ sơ có thể bị sai.",
    )

    # --- 5. Upload ---
    _h1(doc, "5. Đưa hồ sơ lên hệ thống (Upload)")
    _numbered(
        doc,
        [
            "Từ Dashboard, mở hồ sơ → bấm Upload.",
            "Xem bảng gợi ý đặt tên file trên trang Upload (theo từng thành viên).",
            "Đổi tên file trên máy tính theo chuẩn trước khi upload.",
            "Kéo thả file PDF vào vùng upload hoặc bấm chọn file.",
            "Chờ cột trạng thái chuyển sang xử lý xong.",
            "Kiểm tra cột Mã file (01_1, 02_2…) — nếu hiển thị «—» cần kiểm tra lại tên file hoặc khai báo thành viên.",
        ],
    )
    _h2(doc, "5.1. Định dạng file chấp nhận")
    _bullets(doc, ["PDF (khuyến nghị)", "Ảnh: JPG, PNG, WEBP", "Word/Excel: DOCX, XLSX (hạn chế — nên dùng PDF)"])
    _h2(doc, "5.2. Bộ giấy tờ Luồng 1 (bắt buộc cho DS-260)")
    _bullets(
        doc,
        [
            "Hộ chiếu (PASSPORT)",
            "Giấy khai sinh (BIRTH CERTIFICATE)",
            "Lý lịch tư pháp (JUDICIAL CERTIFICATE)",
            "Giấy kết hôn (MARRIAGE CERTIFICATE) — nếu đã kết hôn",
            "Quyết định ly hôn (DIVORCE DECREE) — nếu đã ly hôn",
        ],
    )
    _h2(doc, "5.3. File đối chiếu (_new)")
    _p(
        doc,
        "Khách có thể upload thêm bản đối chiếu (tên file có hậu tố _new hoặc worksheet DS-260). "
        "Hệ thống so sánh với Luồng 1; nếu khác nhau sẽ báo xung đột trên trang Review.",
    )

    # --- 6. Thành viên ---
    _h1(doc, "6. Khai báo thành viên gia đình")
    _numbered(
        doc,
        [
            "Vào Review → mục Thành viên trong hồ sơ.",
            "Nhập tên Phối ngẫu → Lưu tên (nếu có).",
            "Thêm từng Con → Lưu tên.",
            "Quay lại Upload — cột Mã file sẽ cập nhật theo vai trò (01, 02, 03…).",
        ],
    )
    _note(
        doc,
        "Nếu không có vợ/chồng, không khai báo mã 02. Con được đánh số 03, 04… "
        "(bỏ qua 02 khi không có phối ngẫu).",
    )

    # --- 7. Review ---
    _h1(doc, "7. Trang Review — kiểm tra dữ liệu")
    _p(doc, "Trang Review gồm các phần chính:")
    _bullets(
        doc,
        [
            "Danh sách giấy tờ đã upload và kết quả OCR.",
            "Thành viên hồ sơ (chủ, vợ/chồng, con).",
            "Bảng mapping DS-260 theo từng người (tab Chủ hồ sơ / Vợ / Con…).",
            "Panel xung đột dữ liệu (màu vàng).",
            "Kết quả kiểm tra (validation) — trường thiếu hoặc cần xem lại.",
            "Nút Xuất DS-260 Word cho từng thành viên.",
        ],
    )
    _tip(doc, "Mỗi trường DS-260 hiển thị nguồn dữ liệu (file nào, OCR hay worksheet) — bấm để kiểm tra khi nghi ngờ sai.")

    # --- 8. Xung đột ---
    _h1(doc, "8. Giải quyết xung đột dữ liệu")
    _h2(doc, "8.1. Xung đột là gì?")
    _p(
        doc,
        "Xung đột xảy ra khi hai nguồn dữ liệu cho cùng một trường DS-260 nhưng giá trị khác nhau. "
        "Thường gặp nhất:",
    )
    _bullets(
        doc,
        [
            "Luồng 1 (giấy tờ chính) vs bản đối chiếu (_new).",
            "Giấy tờ chính vs DS-260 worksheet khách khai (file DS260).",
        ],
    )

    _h2(doc, "8.2. Cách chọn giá trị đúng")
    _numbered(
        doc,
        [
            "Mở panel «Xung đột dữ liệu DS-260» trên Review.",
            "Đọc tên trường và hai giá trị Nguồn A / Nguồn B.",
            "Nguồn A — Luồng 1: lấy từ giấy tờ chính (hộ chiếu, GKS, lý lịch…).",
            "Nguồn B — Đối chiếu: lấy từ file _new hoặc worksheet DS-260 khách khai.",
            "Bấm chọn nguồn đúng — giá trị tự điền vào bảng DS-260 và file Word khi xuất.",
            "Hoặc nhập giá trị khác ở ô «Giá trị tùy chỉnh» nếu cả A và B đều sai.",
        ],
    )

    _h2(doc, "8.3. Nguyên tắc ưu tiên (mặc định)")
    _table(
        doc,
        ["Thứ tự", "Nguồn", "Khi nào dùng"],
        [
            ["1", "Luồng 1 — giấy tờ chính", "Mặc định khi chưa có xung đột"],
            ["2", "Bản đối chiếu (_new)", "Khi Luồng 1 thiếu trường — hệ thống tự bổ sung"],
            ["3", "Giá trị đã giải quyết xung đột", "Sau khi user chọn A hoặc B"],
            ["4", "Chỉnh sửa thủ công", "User sửa trực tiếp trên Review — ưu tiên cao nhất"],
            ["5", "Worksheet DS-260", "Địa chỉ, điện thoại, email — thường từ worksheet"],
        ],
    )

    _h2(doc, "8.4. Các trường hay xung đột")
    _bullets(
        doc,
        [
            "Họ tên, ngày sinh, giới tính, quốc tịch",
            "Số hộ chiếu, ngày cấp, ngày hết hạn",
            "Tình trạng hôn nhân",
            "Địa chỉ hiện tại, số điện thoại, email (giấy tờ vs worksheet)",
        ],
    )

    # --- 9. Xuất ---
    _h1(doc, "9. Chỉnh sửa thủ công & xuất DS-260 Word")
    _h2(doc, "9.1. Chỉnh sửa trên Review")
    _bullets(
        doc,
        [
            "Bấm vào ô giá trị trong bảng DS-260 → sửa → Lưu.",
            "Giá trị chỉnh tay ghi đè OCR và xung đột đã chọn.",
            "Nên sửa xong tất cả trước khi xuất Word.",
        ],
    )
    _h2(doc, "9.2. Xuất file Word")
    _numbered(
        doc,
        [
            "Chọn tab thành viên cần xuất (Chủ hồ sơ / Vợ / Con…).",
            "Bấm Xuất DS-260 Word.",
            "File tải về theo mẫu form đã cấu hình trên hệ thống.",
            "Mở file Word → kiểm tra lần cuối trước khi gửi khách/LSQ.",
        ],
    )
    _note(doc, "Mỗi người trong hồ sơ gia đình có một file DS-260 riêng — cần xuất từng người.")

    _add_section_ds260_mapping(doc, mapping)

    # --- 11. Các trường hợp ---
    _h1(doc, "11. Các trường hợp DS-260 theo loại hồ sơ")
    doc.add_page_break()

    _h2(doc, "11.1. Hồ sơ đơn — một người nộp")
    _table(
        doc,
        ["Hạng mục", "Chi tiết"],
        [
            ["Thành viên", "Chỉ chủ hồ sơ — mã 01"],
            ["File upload", "01_1 GKS · 01_2 HC · 01_3 ly hôn · 01_4 lý lịch · 01_5 giấy kết hôn (nếu có)"],
            ["DS-260 xuất", "Một file Word cho chủ hồ sơ"],
            ["Mục điền", "Đầy đủ: cá nhân, HC, cha/mẹ, địa chỉ, phối ngẫu (nếu có), con (nếu có)"],
        ],
    )

    _h2(doc, "11.2. Hồ sơ gia đình — chủ hồ sơ (mã 01)")
    _table(
        doc,
        ["Hạng mục", "Chi tiết"],
        [
            ["File cần có", "01_1 … 01_4 (+ giấy kết hôn 01_5 nếu có)"],
            ["Cha / Mẹ trên DS-260", "Lấy từ GKS của chính chủ hồ sơ (01_1) — không lấy từ GKS con"],
            ["Phối ngẫu", "Từ giấy kết hôn + HC chủ hồ sơ"],
            ["Danh sách con", "Từ GKS con (03_1, 04_1…) + worksheet"],
            ["Xuất", "Tab Chủ hồ sơ → Xuất DS-260 Word"],
        ],
    )

    _h2(doc, "11.3. Phối ngẫu (mã 02)")
    _table(
        doc,
        ["Hạng mục", "Chi tiết"],
        [
            ["Khai báo", "Review → Thêm Phối ngẫu → Lưu tên"],
            ["File", "02_1 GKS · 02_2 HC · 02_4 lý lịch (theo bảng chuẩn)"],
            ["DS-260", "Mục B.3 Spouse / Phối ngẫu — xuất riêng một file Word"],
            ["Lưu ý", "GKS phối ngẫu dùng BIRTH CERTIFICATE (không có CHILD)"],
        ],
    )

    _h2(doc, "11.4. Con đi cùng (mã 03–06)")
    _table(
        doc,
        ["Hạng mục", "Chi tiết"],
        [
            ["File GKS", "Bắt buộc: BIRTH CERTIFICATE CHILD — ví dụ 03_1 BIRTH CERTIFICATE CHILD - TEN CON.pdf"],
            ["Cha / Mẹ trên DS-260 con", "Lấy từ GKS con (tên cha/mẹ trên giấy khai sinh)"],
            ["Không điền", "Phối ngẫu, ly hôn, danh sách con — các mục chỉ dành người lớn"],
            ["Xuất", "Chọn tab từng con → Xuất DS-260 Word riêng"],
        ],
    )

    _h2(doc, "11.5. Ví dụ thực tế — Hồ sơ HỒ CÔNG BẢO LONG")
    _p(doc, "Thành viên:", bold=True)
    _bullets(
        doc,
        [
            "01 — HỒ CÔNG BẢO LONG (chủ hồ sơ)",
            "02 — VĂN THỊ HƯỜNG (phối ngẫu)",
            "03 — HỒ BẢO CHÂU (con)",
            "04 — HỒ BẢO HÂN (con)",
            "05 — HỒ CÔNG BẢO TRÍ (con)",
        ],
    )
    _p(doc, "File mẫu:", bold=True)
    _bullets(
        doc,
        [
            "01_1 BIRTH CERTIFICATE - HO CONG BAO LONG.pdf",
            "01_2 PASSPORT - HO CONG BAO LONG.pdf",
            "02_2 PASSPORT - VAN THI HUONG.pdf",
            "03_1 BIRTH CERTIFICATE CHILD - HO BAO CHAU.pdf",
            "03_2 PASSPORT - HO BAO CHAU.pdf",
            "(tương tự cho con 04, 05…)",
        ],
    )
    _note(
        doc,
        "Nếu cha/mẹ trên DS-260 chủ hồ sơ hiển thị tên vợ/chồng thay vì ông bà — "
        "kiểm tra file 01_1 có đúng GKS chủ hồ sơ và GKS con có dùng BIRTH CERTIFICATE CHILD.",
    )

    # --- 12. Lỗi thường gặp ---
    _h1(doc, "12. Lỗi thường gặp & cách khắc phục")
    _table(
        doc,
        ["Triệu chứng", "Nguyên nhân", "Cách xử lý"],
        [
            ["Cột Mã file hiển thị «—»", "Chưa khai báo thành viên hoặc tên file sai format", "Đổi tên file theo chuẩn; khai báo thành viên trên Review"],
            ["Cha/mẹ chủ hồ sơ sai (hiện tên vợ/con)", "GKS con bị upload nhầm loại birth_certificate", "Đổi tên có CHILD; xóa upload lại; kiểm tra 01_1"],
            ["Xung đột nhiều trường", "Có cả Luồng 1 và file _new/worksheet", "Giải quyết từng trường — ưu tiên giấy tờ chính"],
            ["DS-260 thiếu địa chỉ/SĐT", "Chưa có worksheet DS-260", "Upload 01_6 DS260 - TEN.pdf hoặc điền thủ công"],
            ["OCR sai họ tên", "Ảnh mờ, scan chất lượng thấp", "Upload PDF rõ; sửa thủ công trên Review"],
            ["Trường DS-260 trống dù đã có giấy", "Sai loại giấy hoặc OCR thiếu field", "Đối chiếu Phụ lục C — kiểm tra field OCR; upload file _new"],
            ["Không xuất được Word", "Chưa chọn mẫu form hoặc thiếu quyền", "Liên hệ admin kiểm tra mẫu DS-260"],
        ],
    )

    doc.add_page_break()

    # --- Phụ lục A ---
    _h1(doc, "Phụ lục A — Bảng đặt tên file mẫu (gia đình đủ 6 người)")
    members = [
        ("01", "Chủ hồ sơ", "DANG VAN HUNG", False),
        ("02", "Vợ/chồng", "MAI THI HUONG", False),
        ("03", "Con 1", "DANG MAI PHUONG THAO", True),
        ("04", "Con 2", "DANG KHOI NGUYEN", True),
        ("05", "Con 3", "TEN CON 3", True),
        ("06", "Con 4", "TEN CON 4", True),
    ]
    slots = [
        (1, "GKS", "BIRTH CERTIFICATE", "BIRTH CERTIFICATE CHILD"),
        (2, "HC", "PASSPORT", "PASSPORT"),
        (3, "Ly hôn", "DIVORCE DECREE", "DIVORCE DECREE"),
        (4, "Lý lịch", "JUDICIAL CERTIFICATE", "JUDICIAL CERTIFICATE"),
    ]
    rows: list[list[str]] = []
    for code, role, name, is_child in members:
        for seq, label, adult_part, child_part in slots:
            part = child_part if is_child else adult_part
            rows.append([code, role, name, f"{code}_{seq}", label, f"{code}_{seq} {part} - {name}.pdf"])
    _table(doc, ["Mã", "Vai trò", "Họ tên", "Mã file", "Loại", "Tên file gợi ý"], rows)

    # --- Phụ lục B ---
    _h1(doc, "Phụ lục B — Hai luồng dữ liệu")
    _table(
        doc,
        ["Khái niệm", "Mô tả", "Ví dụ tên file"],
        [
            ["Luồng 1 (standard)", "Giấy tờ chính — nguồn ưu tiên", "01_2 PASSPORT - TEN.pdf"],
            ["Đối chiếu (exception / _new)", "Bản khách upload để so sánh", "01_2 PASSPORT - TEN_new.pdf"],
            ["Worksheet DS-260", "Form khách khai thêm thông tin", "01_6 DS260 - TEN.pdf"],
        ],
    )
    _p(
        doc,
        "Khi Luồng 1 và đối chiếu khác nhau → hệ thống báo xung đột. "
        "Người dùng chọn giá trị đúng trên Review trước khi xuất Word.",
    )

    _add_appendix_mapping_detail(doc, mapping)

    # --- Chân trang ---
    doc.add_paragraph()
    footer = doc.add_paragraph("— Hết tài liệu — ImmiFill / Edupath —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer.runs[0].font.size = Pt(9)

    return doc


def main() -> None:
    doc = build_guide()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
