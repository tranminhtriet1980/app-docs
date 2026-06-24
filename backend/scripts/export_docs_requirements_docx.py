"""
Yêu cầu gửi tài liệu Phòng Docs + kịch bản từng bộ hồ sơ.
Chạy: cd backend && .venv\\Scripts\\python.exe scripts/export_docs_requirements_docx.py
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[2] / "Yeu_cau_Phong_Docs_gui_ho_so.docx"
OUT_EMAIL = Path(__file__).resolve().parents[2] / "Email_mau_gui_Phong_Docs.txt"
OUT_EMAIL_DOCX = Path(__file__).resolve().parents[2] / "Email_mau_gui_Phong_Docs.docx"
OUT_PHIEU = Path(__file__).resolve().parents[2] / "Phieu_mo_ta_bo_ho_so_mau.docx"


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, *, bold: bool = False) -> None:
    r = doc.add_paragraph().add_run(text)
    r.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


EMAIL_BODY = """Kính gửi Anh/Chị Phòng Docs,

IT đã triển khai hệ thống ImmiFill (AI fill DS-260) lên môi trường test. Để hệ thống điền form chính xác, kính đề nghị Phòng Docs gửi hồ sơ theo đúng quy chuẩn dưới đây — không chỉ gửi file PDF mà phải kèm mô tả từng bộ.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
I. NGUYÊN TẮC CHUNG (BẮT BUỘC)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. MỖI BỘ HỒ SƠ GỬI KÈM 01 PHIẾU MÔ TẢ (mẫu: Phieu_mo_ta_bo_ho_so_mau.docx đính kèm).
2. ĐẶT TÊN FILE ĐÚNG CHUẨN trước khi upload — không upload file tên tùy ý.
3. MỖI FILE = 01 PDF (ưu tiên scan rõ, không photo mờ).
4. GỬI ĐỦ GIẤY TỜ THEO TỪNG NGƯỜI trong bộ (chủ / vợ / từng con).
5. GHI RÕ LOẠI HỒ SƠ: đơn / vợ chồng / gia đình có con / đã ly hôn / case đặc biệt.

Format tên file:
  {mã người}_{số file} {LOẠI GIẤY} - {HỌ TÊN}.pdf

Ví dụ:
  01_2 PASSPORT - HO CONG BAO LONG.pdf
  03_1 BIRTH CERTIFICATE CHILD - HO BAO CHAU.pdf

Mã người:
  01 = Chủ hồ sơ | 02 = Vợ/chồng | 03–06 = Con 1–4

Số file chuẩn:
  _1 Giấy khai sinh | _2 Hộ chiếu | _3 Ly hôn | _4 Lý lịch tư pháp
  _5 Giấy kết hôn | _6 DS-260 worksheet khách khai | _7+ Giấy khác

Lưu ý: Giấy khai sinh CON phải có từ CHILD trong tên:
  BIRTH CERTIFICATE CHILD (không dùng BIRTH CERTIFICATE của người lớn).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
II. NỘI DUNG GỬI THEO TỪNG LOẠI BỘ HỒ SƠ
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

【Case A】HỒ SƠ ĐƠN — 1 người nộp, không vợ/con đi cùng
  Phiếu mô tả: Loại = Đơn | Chỉ có mã 01
  File tối thiểu:
    01_1 BIRTH CERTIFICATE - [TÊN CHỦ].pdf
    01_2 PASSPORT - [TÊN CHỦ].pdf
    01_4 JUDICIAL CERTIFICATE - [TÊN CHỦ].pdf
    01_6 DS260 - [TÊN CHỦ].pdf (worksheet địa chỉ, SĐT, email)
  Nếu đã kết hôn (phối ngẫu không đi cùng): thêm 01_5 MARRIAGE CERTIFICATE
  Nếu đã ly hôn: thêm 01_3 DIVORCE DECREE
  Kết quả: 01 file DS-260 Word cho chủ hồ sơ

【Case B】VỢ CHỒNG — 2 người đi cùng, không con
  Phiếu mô tả: Loại = Vợ chồng | 01 + 02
  File chủ hồ sơ (01_x):
    01_1, 01_2, 01_4, 01_5 (giấy kết hôn), 01_6 DS260
  File phối ngẫu (02_x):
    02_1 BIRTH CERTIFICATE - [TÊN VỢ/CHỒNG].pdf
    02_2 PASSPORT - [TÊN VỢ/CHỒNG].pdf
    02_4 JUDICIAL CERTIFICATE - [TÊN VỢ/CHỒNG].pdf
  Kết quả: 02 file DS-260 Word (chủ + phối ngẫu)

【Case C】GIA ĐÌNH — chủ + vợ + con đi cùng (phổ biến nhất)
  Phiếu mô tả: Loại = Gia đình | Liệt kê đủ tên 01, 02, 03, 04…
  Ví dụ: HỒ CÔNG BẢO LONG (01), VĂN THỊ HƯỜNG (02), 3 con (03, 04, 05)

  Chủ hồ sơ 01_x: GKS, HC, lý lịch, kết hôn, DS260 worksheet
  Phối ngẫu 02_x: GKS, HC, lý lịch (tối thiểu GKS + HC)
  Mỗi con 03_x, 04_x…:
    03_1 BIRTH CERTIFICATE CHILD - [TÊN CON].pdf  ← BẮT BUỘC
    03_2 PASSPORT - [TÊN CON].pdf (nếu có)

  Kết quả: 1 file DS-260 Word / người (5 người = 5 file)

【Case D】GIA ĐÌNH — chủ + con, KHÔNG có vợ/chồng đi cùng
  Phiếu mô tả: Loại = Gia đình (đơn thân + con) | Không khai 02
  Chỉ mã 01 + 03, 04… (bỏ qua 02)
  File con: BIRTH CERTIFICATE CHILD bắt buộc

【Case E】ĐÃ LY HÔN — chủ hồ sơ tái hôn hoặc đơn thân
  Phiếu mô tả: Ghi rõ lịch sử hôn nhân
  Bắt buộc: 01_3 DIVORCE DECREE - [TÊN CHỦ].pdf
  Nếu tái hôn: 01_5 MARRIAGE CERTIFICATE (hôn nhân hiện tại)

【Case F】CASE ĐẶC BIỆT — IT cần biết TRƯỚC
  Ghi rõ trên phiếu mô tả, ví dụ:
  - Cha/mẹ mất, không có thông tin trên GKS
  - Tên khác nhau giữa HC và GKS
  - Con nhận nuôi, con riêng từng vợ/chồng
  - Nhiều lần ly hôn / nhiều đợt kết hôn
  → Kèm ghi chú bằng chữ, không chỉ gửi file im lặng

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
III. CÁCH GỬI CHO IT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Cách 1 (khuyến nghị): Upload trực tiếp lên ImmiFill + gửi mail thông báo mã hồ sơ
Cách 2: Gửi mail kèm:
  - 01 file Word/PDF: Phiếu mô tả bộ hồ sơ (điền đủ)
  - Folder ZIP: tất cả PDF đã đổi tên đúng chuẩn
  - Dòng tiêu đề mail: [ImmiFill] {Loại case} - {Tên chủ hồ sơ} - {Ngày}

Không gửi:
  - File zip không có phiếu mô tả
  - File tên "scan001.pdf", "giay to.pdf"
  - GKS con không có từ CHILD
  - Chỉ gửi worksheet DS-260 mà thiếu HC, GKS

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
IV. SAU KHI IT XỬ LÝ — PHÒNG DOCS LÀM GÌ
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. Vào ImmiFill → Review → tab DS-260
2. Kiểm tra từng người (chủ / vợ / con)
3. Giải quyết xung đột (nếu có) — chọn nguồn đúng
4. Sửa ô sai → bấm Lưu từng trường
5. Xuất DS-260 Word từng người
6. Báo IT nếu: field sai, thiếu giấy, case chưa hỗ trợ (kèm tên field + ảnh chụp)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
V. YÊU CẦU GIAI ĐOẠN PILOT (2 TUẦN ĐẦU)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Kính đề nghị Phòng Docs gửi 3 bộ mẫu hoàn chỉnh:
  1. Hồ sơ đơn (Case A)
  2. Vợ chồng không con (Case B)
  3. Gia đình có con (Case C) — ví dụ bộ HỒ CÔNG BẢO LONG

Mỗi bộ kèm phiếu mô tả + file đã đặt tên chuẩn.

Tài liệu tham khảo đính kèm mail:
  - Phieu_mo_ta_bo_ho_so_mau.docx (mẫu phiếu mô tả — điền mỗi bộ hồ sơ)
  - Yeu_cau_Phong_Docs_gui_ho_so.docx
  - Huong_dan_su_dung_ImmiFill_DS260.docx
  - DS260_Mapping_Reference_new.xlsx

Trân trọng,
[Tên bạn] — IT Edupath
"""


def build_word() -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_heading("YÊU CẦU GỬI HỒ SƠ — PHÒNG DOCS", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("ImmiFill · AI fill DS-260 · Kèm kịch bản từng bộ hồ sơ")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Ngày: {date.today().strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    _h(doc, "Phần 1 — Email gửi Phòng Docs (copy vào mail)")
    _p(doc, "Tiêu đề mail gợi ý:")
    _p(doc, "[ImmiFill] Quy chuẩn gửi hồ sơ AI fill DS-260 — Yêu cầu Phòng Docs", bold=True)
    _p(doc, "Nội dung: xem file Email_mau_gui_Phong_Docs.txt hoặc copy từ Phần 1 trong tài liệu họp.")
    doc.add_page_break()

    _h(doc, "Phần 2 — Nguyên tắc chung (bắt buộc)", 1)
    _numbered(
        doc,
        [
            "Mỗi bộ hồ sơ kèm 01 Phiếu mô tả (Phụ lục 1 cuối file).",
            "Đặt tên file đúng chuẩn trước khi upload.",
            "Mỗi file = 01 PDF, scan rõ.",
            "Gửi đủ giấy theo từng người (01, 02, 03…).",
            "Ghi rõ loại hồ sơ trên phiếu mô tả.",
        ],
    )

    _h(doc, "Quy tắc đặt tên file", 2)
    _p(doc, "Format: {mã}_{số} {LOẠI GIẤY} - {HỌ TÊN}.pdf", bold=True)
    _table(
        doc,
        ["Mã", "Vai trò", "_1", "_2", "_3", "_4", "_5", "_6"],
        [
            ["01", "Chủ hồ sơ", "GKS", "HC", "Ly hôn", "Lý lịch", "Kết hôn", "DS260"],
            ["02", "Vợ/chồng", "GKS", "HC", "Ly hôn", "Lý lịch", "—", "—"],
            ["03–06", "Con 1–4", "GKS CHILD", "HC", "—", "—", "—", "—"],
        ],
    )
    _p(doc, "Con: BIRTH CERTIFICATE CHILD — bắt buộc có từ CHILD.", bold=True)

    doc.add_page_break()
    _h(doc, "Phần 3 — Kịch bản chi tiết từng bộ hồ sơ", 1)

    cases = [
        (
            "Case A — Hồ sơ đơn (1 người)",
            "Một người nộp DS-260, không vợ/con đi cùng.",
            [
                "Phiếu mô tả: Loại = Đơn | Chỉ mã 01",
                "Không khai báo vợ/con trên ImmiFill",
            ],
            [
                ["01_1", "BIRTH CERTIFICATE", "Chủ hồ sơ", "Bắt buộc"],
                ["01_2", "PASSPORT", "Chủ hồ sơ", "Bắt buộc"],
                ["01_4", "JUDICIAL CERTIFICATE", "Chủ hồ sơ", "Bắt buộc"],
                ["01_6", "DS260", "Chủ hồ sơ", "Địa chỉ, SĐT, email"],
                ["01_5", "MARRIAGE CERTIFICATE", "Chủ hồ sơ", "Nếu đã kết hôn"],
                ["01_3", "DIVORCE DECREE", "Chủ hồ sơ", "Nếu đã ly hôn"],
            ],
            [
                "Upload tất cả file 01_x",
                "Review → chỉ tab Chủ hồ sơ",
                "Giải quyết xung đột (nếu có)",
                "Sửa ô sai → Lưu",
                "Xuất 01 file DS-260 Word",
            ],
            "1 file Word",
        ),
        (
            "Case B — Vợ chồng (2 người, không con)",
            "Hai người đi cùng, xuất DS-260 riêng cho mỗi người.",
            [
                "Phiếu mô tả: Loại = Vợ chồng | 01 + 02",
                "Khai báo 02 = tên phối ngẫu trên Review",
            ],
            [
                ["01_1–01_4", "GKS, HC, lý lịch", "Chủ", "Bắt buộc"],
                ["01_5", "MARRIAGE CERTIFICATE", "Chủ", "Bắt buộc"],
                ["01_6", "DS260", "Chủ", "Worksheet chung hộ gia đình"],
                ["02_1", "BIRTH CERTIFICATE", "Phối ngẫu", "Bắt buộc"],
                ["02_2", "PASSPORT", "Phối ngẫu", "Bắt buộc"],
                ["02_4", "JUDICIAL CERTIFICATE", "Phối ngẫu", "Khuyến nghị"],
            ],
            [
                "Upload file 01_x và 02_x",
                "Review → khai báo Phối ngẫu → Lưu tên",
                "Tab Chủ hồ sơ: kiểm tra mục Phối ngẫu từ giấy kết hôn",
                "Tab Phối ngẫu: kiểm tra A.1, A.2, cha/mẹ từ GKS 02_1",
                "Xuất Word: chủ + phối ngẫu",
            ],
            "2 file Word",
        ),
        (
            "Case C — Gia đình (chủ + vợ + con)",
            "Phổ biến nhất — ví dụ: HỒ CÔNG BẢO LONG + VĂN THỊ HƯỜNG + 3 con.",
            [
                "Phiếu mô tả: Loại = Gia đình | Liệt kê đủ 01, 02, 03, 04, 05",
                "Mỗi con một dòng tên trên phiếu",
            ],
            [
                ["01_x", "GKS, HC, lý lịch, kết hôn, DS260", "Chủ", "Đủ bộ 01"],
                ["02_x", "GKS, HC, lý lịch", "Phối ngẫu", "Tối thiểu 02_1, 02_2"],
                ["03_1", "BIRTH CERTIFICATE CHILD", "Con 1", "Bắt buộc CHILD"],
                ["03_2", "PASSPORT", "Con 1", "Nếu có"],
                ["04_1", "BIRTH CERTIFICATE CHILD", "Con 2", "Tương tự"],
                ["05_1", "BIRTH CERTIFICATE CHILD", "Con 3", "Tương tự"],
            ],
            [
                "Đặt tên file đủ 01–05 (hoặc 01–06 nếu 4 con)",
                "Upload → Review → Lưu thành viên (nếu chưa có)",
                "Tab từng người: chủ / vợ / từng con",
                "Con: cha = chủ hồ sơ, mẹ = phối ngẫu (từ GKS con hoặc hệ thống fallback)",
                "Chủ: mục Con cái từ GKS con + worksheet",
                "Xuất Word từng tab → N người = N file",
            ],
            "N file Word (1/người)",
        ),
        (
            "Case D — Chủ hồ sơ + con (không vợ đi cùng)",
            "Đơn thân hoặc vợ không nhập cư cùng — không khai 02.",
            [
                "Phiếu mô tả: Loại = Gia đình (chủ + con) | Không có 02",
                "Con bắt đầu từ mã 03 (bỏ qua 02)",
            ],
            [
                ["01_x", "Bộ chủ hồ sơ", "Chủ", "Đủ"],
                ["03_1", "BIRTH CERTIFICATE CHILD", "Con 1", "Bắt buộc"],
                ["04_1", "BIRTH CERTIFICATE CHILD", "Con 2", "Nếu có thêm con"],
            ],
            [
                "Không khai báo phối ngẫu trên Review",
                "File con chỉ 03_x, 04_x…",
                "Xuất Word: chủ + từng con",
            ],
            "1 + số con file Word",
        ),
        (
            "Case E — Đã ly hôn / tái hôn",
            "Cần giấy ly hôn; nếu tái hôn thêm giấy kết hôn mới.",
            [
                "Phiếu mô tả: Ghi rõ lịch sử hôn nhân",
                "Case đặc biệt: ghi chú trên phiếu",
            ],
            [
                ["01_3", "DIVORCE DECREE", "Chủ", "Bắt buộc nếu ly hôn"],
                ["01_5", "MARRIAGE CERTIFICATE", "Chủ", "Nếu tái hôn hiện tại"],
                ["01_2", "PASSPORT", "Chủ", "Bắt buộc"],
            ],
            [
                "Upload 01_3 trước hoặc cùng bộ",
                "Review: mục Ly hôn / Previous spouse từ giấy ly hôn",
                "Không để worksheet điền nhầm mục ly hôn — sửa tay + Lưu",
            ],
            "Tùy số người trong bộ",
        ),
        (
            "Case F — Case đặc biệt (báo IT trước)",
            "IT cần mô tả bằng chữ — không chỉ gửi file.",
            [
                "Ghi trên phiếu: tên khác HC/GKS, cha mẹ mất, con nuôi…",
                "Đính kèm email riêng cho IT nếu phức tạp",
            ],
            [
                ["—", "Mô tả case", "Phiếu", "Bắt buộc"],
                ["01_x", "File chuẩn", "Theo case", "Càng đủ càng tốt"],
            ],
            [
                "Gửi phiếu + file + ghi chú đặc biệt",
                "Chờ IT xác nhận hoặc xử lý thủ công một phần",
                "Không kỳ vọng auto 100% nếu chưa từng gặp case",
            ],
            "Theo thỏa thuận",
        ),
    ]

    for title, desc, notes, files, steps, output in cases:
        _h(doc, title, 2)
        _p(doc, desc)
        _p(doc, "Ghi chú:", bold=True)
        _bullets(doc, notes)
        _p(doc, "Danh sách file:", bold=True)
        _table(doc, ["Mã file", "Loại giấy", "Người", "Ghi chú"], files)
        _p(doc, "Kịch bản xử lý trên ImmiFill:", bold=True)
        _numbered(doc, steps)
        _p(doc, f"Kết quả xuất: {output}", bold=True)
        doc.add_paragraph()

    doc.add_page_break()
    _h(doc, "Phụ lục 1 — PHIẾU MÔ TẢ BỘ HỒ SƠ (Docs điền mỗi lần gửi)", 1)
    _p(
        doc,
        "Dùng file mẫu riêng: Phieu_mo_ta_bo_ho_so_mau.docx — in 01 tờ/bộ hồ sơ, "
        "kèm mail hoặc đặt trong folder ZIP.",
        bold=True,
    )
    fields = [
        ("Ngày gửi", ""),
        ("Người gửi (Docs)", ""),
        ("Mã / tên hồ sơ trên ImmiFill", ""),
        ("Loại case", "☐ Đơn  ☐ Vợ chồng  ☐ Gia đình  ☐ Ly hôn  ☐ Đặc biệt"),
        ("Chủ hồ sơ (01) — Họ tên", ""),
        ("Phối ngẫu (02) — Có / Không — Họ tên", ""),
        ("Con — liệt kê tên + mã dự kiến", "03: …  04: …  05: …  06: …"),
        ("Số file PDF gửi kèm", ""),
        ("Đã đặt tên file chuẩn?", "☐ Có  ☐ Chưa — IT từ chối xử lý nếu chưa"),
        ("Giấy tờ đặc biệt / case lạ", ""),
        ("Người xử lý Docs sau khi IT upload", ""),
        ("Ghi chú thêm cho IT", ""),
    ]
    _table(doc, ["Mục", "Nội dung (Docs điền)"], [[a, b] for a, b in fields])

    _h(doc, "Phụ lục 2 — Checklist trước khi bấm Gửi mail", 1)
    _bullets(
        doc,
        [
            "☐ Tất cả PDF đã đổi tên đúng format",
            "☐ GKS con có BIRTH CERTIFICATE CHILD",
            "☐ Phiếu mô tả đã điền đủ",
            "☐ Loại case đã chọn đúng",
            "☐ Tên trên file khớp tên trên phiếu",
            "☐ Case đặc biệt đã ghi chú",
        ],
    )

    return doc


def build_email_word() -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EMAIL MẪU GỬI PHÒNG DOCS — ImmiFill DS-260")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Ngày cập nhật: {date.today().strftime('%d/%m/%Y')}").italic = True
    doc.add_paragraph()

    _p(doc, "Tiêu đề mail gợi ý:", bold=True)
    _p(doc, "[ImmiFill] {Loại case} - {Tên chủ hồ sơ} - {Ngày}")
    doc.add_paragraph()

    _p(doc, "Kính gửi Anh/Chị Phòng Docs,", bold=True)
    _p(
        doc,
        "IT đã triển khai hệ thống ImmiFill (AI fill DS-260) lên môi trường test. "
        "Để hệ thống điền form chính xác, kính đề nghị Phòng Docs gửi hồ sơ theo đúng "
        "quy chuẩn dưới đây — không chỉ gửi file PDF mà phải kèm mô tả từng bộ.",
    )

    _h(doc, "I. Nguyên tắc chung (bắt buộc)", 1)
    _numbered(
        doc,
        [
            "MỖI BỘ HỒ SƠ GỬI KÈM 01 PHIẾU MÔ TẢ (mẫu: Phieu_mo_ta_bo_ho_so_mau.docx đính kèm).",
            "ĐẶT TÊN FILE ĐÚNG CHUẨN trước khi upload — không upload file tên tùy ý.",
            "MỖI FILE = 01 PDF (ưu tiên scan rõ, không photo mờ).",
            "GỬI ĐỦ GIẤY TỜ THEO TỪNG NGƯỜI trong bộ (chủ / vợ / từng con).",
            "GHI RÕ LOẠI HỒ SƠ: đơn / vợ chồng / gia đình có con / đã ly hôn / case đặc biệt.",
        ],
    )
    _p(doc, "Format tên file:", bold=True)
    _p(doc, "{mã người}_{số file} {LOẠI GIẤY} - {HỌ TÊN}.pdf")
    _p(doc, "Ví dụ:", bold=True)
    _bullets(
        doc,
        [
            "01_2 PASSPORT - HO CONG BAO LONG.pdf",
            "03_1 BIRTH CERTIFICATE CHILD - HO BAO CHAU.pdf",
        ],
    )
    _p(doc, "Mã người: 01 = Chủ hồ sơ | 02 = Vợ/chồng | 03–06 = Con 1–4")
    _p(doc, "Số file chuẩn:", bold=True)
    _p(
        doc,
        "_1 Giấy khai sinh | _2 Hộ chiếu | _3 Ly hôn | _4 Lý lịch tư pháp | "
        "_5 Giấy kết hôn | _6 DS-260 worksheet | _7+ Giấy khác",
    )
    _p(
        doc,
        "Lưu ý: Giấy khai sinh CON phải có từ CHILD trong tên "
        "(BIRTH CERTIFICATE CHILD — không dùng BIRTH CERTIFICATE của người lớn).",
        bold=True,
    )

    _h(doc, "II. Nội dung gửi theo từng loại bộ hồ sơ", 1)

    _h(doc, "Case A — Hồ sơ đơn (1 người)", 2)
    _bullets(
        doc,
        [
            "Phiếu mô tả: Loại = Đơn | Chỉ có mã 01",
            "File tối thiểu: 01_1 GKS, 01_2 HC, 01_4 lý lịch, 01_6 DS260",
            "Nếu đã kết hôn (phối ngẫu không đi cùng): thêm 01_5 MARRIAGE CERTIFICATE",
            "Nếu đã ly hôn: thêm 01_3 DIVORCE DECREE",
            "Kết quả: 01 file DS-260 Word cho chủ hồ sơ",
        ],
    )

    _h(doc, "Case B — Vợ chồng (2 người, không con)", 2)
    _bullets(
        doc,
        [
            "Phiếu mô tả: Loại = Vợ chồng | 01 + 02",
            "Chủ (01_x): 01_1, 01_2, 01_4, 01_5 kết hôn, 01_6 DS260",
            "Phối ngẫu (02_x): 02_1 GKS, 02_2 HC, 02_4 lý lịch",
            "Kết quả: 02 file DS-260 Word (chủ + phối ngẫu)",
        ],
    )

    _h(doc, "Case C — Gia đình (chủ + vợ + con)", 2)
    _bullets(
        doc,
        [
            "Phiếu mô tả: Loại = Gia đình | Liệt kê đủ tên 01, 02, 03, 04…",
            "Ví dụ: HỒ CÔNG BẢO LONG (01), VĂN THỊ HƯỜNG (02), 3 con (03, 04, 05)",
            "Chủ 01_x: GKS, HC, lý lịch, kết hôn, DS260 worksheet",
            "Phối ngẫu 02_x: GKS, HC, lý lịch (tối thiểu GKS + HC)",
            "Mỗi con: 03_1 BIRTH CERTIFICATE CHILD (bắt buộc), 03_2 PASSPORT (nếu có)",
            "Kết quả: 1 file DS-260 Word / người",
        ],
    )

    _h(doc, "Case D — Chủ + con (không vợ đi cùng)", 2)
    _bullets(
        doc,
        [
            "Phiếu mô tả: Loại = Gia đình (đơn thân + con) | Không khai 02",
            "Chỉ mã 01 + 03, 04… (bỏ qua 02)",
            "File con: BIRTH CERTIFICATE CHILD bắt buộc",
        ],
    )

    _h(doc, "Case E — Đã ly hôn / tái hôn", 2)
    _bullets(
        doc,
        [
            "Phiếu mô tả: Ghi rõ lịch sử hôn nhân",
            "Bắt buộc: 01_3 DIVORCE DECREE",
            "Nếu tái hôn: 01_5 MARRIAGE CERTIFICATE (hôn nhân hiện tại)",
        ],
    )

    _h(doc, "Case F — Case đặc biệt (báo IT trước)", 2)
    _bullets(
        doc,
        [
            "Cha/mẹ mất, không có thông tin trên GKS",
            "Tên khác nhau giữa HC và GKS",
            "Con nhận nuôi, con riêng từng vợ/chồng",
            "Nhiều lần ly hôn / nhiều đợt kết hôn",
            "→ Kèm ghi chú bằng chữ, không chỉ gửi file im lặng",
        ],
    )

    _h(doc, "III. Cách gửi cho IT", 1)
    _p(doc, "Cách 1 (khuyến nghị): Upload trực tiếp lên ImmiFill + gửi mail thông báo mã hồ sơ")
    _p(doc, "Cách 2: Gửi mail kèm:", bold=True)
    _bullets(
        doc,
        [
            "01 file Word/PDF: Phiếu mô tả bộ hồ sơ (điền đủ)",
            "Folder ZIP: tất cả PDF đã đổi tên đúng chuẩn",
            "Tiêu đề mail: [ImmiFill] {Loại case} - {Tên chủ hồ sơ} - {Ngày}",
        ],
    )
    _p(doc, "Không gửi:", bold=True)
    _bullets(
        doc,
        [
            "File zip không có phiếu mô tả",
            'File tên "scan001.pdf", "giay to.pdf"',
            "GKS con không có từ CHILD",
            "Chỉ gửi worksheet DS-260 mà thiếu HC, GKS",
        ],
    )

    _h(doc, "IV. Sau khi IT xử lý — Phòng Docs làm gì", 1)
    _numbered(
        doc,
        [
            "Vào ImmiFill → Review → tab DS-260",
            "Kiểm tra từng người (chủ / vợ / con)",
            "Giải quyết xung đột (nếu có) — chọn nguồn đúng",
            "Sửa ô sai → bấm Lưu từng trường",
            "Xuất DS-260 Word từng người",
            "Báo IT nếu: field sai, thiếu giấy, case chưa hỗ trợ (kèm tên field + ảnh chụp)",
        ],
    )

    _h(doc, "V. Yêu cầu giai đoạn pilot (2 tuần đầu)", 1)
    _p(doc, "Kính đề nghị Phòng Docs gửi 3 bộ mẫu hoàn chỉnh:")
    _numbered(
        doc,
        [
            "Hồ sơ đơn (Case A)",
            "Vợ chồng không con (Case B)",
            "Gia đình có con (Case C) — ví dụ bộ HỒ CÔNG BẢO LONG",
        ],
    )
    _p(doc, "Mỗi bộ kèm phiếu mô tả + file đã đặt tên chuẩn.")
    _p(doc, "Tài liệu tham khảo đính kèm mail:", bold=True)
    _bullets(
        doc,
        [
            "Phieu_mo_ta_bo_ho_so_mau.docx (mẫu phiếu mô tả — điền mỗi bộ hồ sơ)",
            "Yeu_cau_Phong_Docs_gui_ho_so.docx",
            "Huong_dan_su_dung_ImmiFill_DS260.docx",
            "DS260_Mapping_Reference_new.xlsx",
        ],
    )

    doc.add_paragraph()
    _p(doc, "Trân trọng,")
    _p(doc, "[Tên bạn] — IT Edupath", bold=True)

    return doc


def main() -> None:
    import sys

    scripts_dir = Path(__file__).resolve().parent
    if str(scripts_dir) not in sys.path:
        sys.path.insert(0, str(scripts_dir))
    from export_phieu_mo_ta_docx import OUT as PHIEU_OUT, build as build_phieu

    doc = build_word()
    doc.save(OUT)
    OUT_EMAIL.write_text(EMAIL_BODY, encoding="utf-8")
    email_doc = build_email_word()
    email_doc.save(OUT_EMAIL_DOCX)
    build_phieu().save(PHIEU_OUT)
    print(f"Created: {OUT}")
    print(f"Created: {OUT_EMAIL}")
    print(f"Created: {OUT_EMAIL_DOCX}")
    print(f"Created: {PHIEU_OUT}")


if __name__ == "__main__":
    main()
