"""Xuất DS260 từ doc records + mapping — điền template Word (label: value)."""

from __future__ import annotations

import copy
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.entities import Applicant, ApplicantStatus, Export, FormTemplate
from app.services.ds260_mapping import flatten_ds260_mappings, resolve_ds260_form
from app.services.ds260_dates import format_ds260_export_date, is_date_field_key
from app.services.ds260_normalize import normalize_gender
from app.services.ds260_validate import flatten_ds260_values, validate_ds260
from app.services.postal_code import derive_postal_code_from_location
from app.services.birth_location import (
    canonical_vn_city,
    find_vn_locality,
    format_address_english,
    format_birth_city_display,
    format_nationality_country,
    format_person_name_ascii,
    format_place_name_title,
    normalize_location,
    normalize_present_marker,
)

DS260_TEMPLATE_CODE = "ds260_final"
DS260_DEFAULT_TEMPLATE_CODE = "6_eb3_tt_-___n_ds260_-_h_ng_1"
DEFAULT_TEMPLATE_FALLBACKS = (
    DS260_DEFAULT_TEMPLATE_CODE,
    "ds260_custom",
    "ds260_new",
    DS260_TEMPLATE_CODE,
)

_COLON_FILLER = re.compile(r"[\s\t_\.·…\-\u00a0\u2013\u2014]+")
_TIME_LIKE = re.compile(r":\d{1,2}\b")
_URL_LIKE = re.compile(r"://")

# Một dòng "trống dành để viết câu trả lời" trong template: chỉ chứa dấu ':' đầu dòng (tùy chọn)
# + khoảng trắng/gạch dưới/gạch ngang — không có nội dung câu hỏi nào khác.
_CONTINUATION_BLANK_RE = re.compile(r"^[\s:\._·…\-\u00a0\u2013\u2014]*$")
_BARE_COLON_REST_MAX_LEN = 6

DS260_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"full name in native language", re.I), "applicant_name_native"),
    (re.compile(r"name\s*\(last/middle/first\)", re.I), "applicant_name"),
    (re.compile(r"current marital status|tình trạng hôn nhân hiện tại", re.I), "current_marital_status"),
    (re.compile(r"father.*surnames?|father.*surname|họ.*cha(?!ng)|họ của cha", re.I), "father_surname"),
    (re.compile(r"father.*given names?|father.*given|tên.*cha(?!ng)", re.I), "father_given_names"),
    (re.compile(r"father.*date of birth|ngày.*sinh.*cha|date of birth.*father", re.I), "father_date_of_birth"),
    (re.compile(r"father.*city of birth|thành phố.*sinh.*cha|city of birth.*father", re.I), "father_birth_city"),
    (re.compile(r"father.*state.*birth|father.*province.*birth|tỉnh.*sinh.*cha|bang.*sinh.*cha", re.I), "father_birth_state"),
    (re.compile(r"father.*country.*birth|quốc gia.*sinh.*cha|country.*birth.*father", re.I), "father_birth_country"),
    (re.compile(r"father still living", re.I), "father_is_living"),
    (re.compile(r"mother still living", re.I), "mother_is_living"),
    (re.compile(r"spouse.?s surnames|họ của chồng /vợ|họ của vợ|họ của chồng", re.I), "spouse_surname"),
    (re.compile(r"spouse.?s given names|tên của chồng/vợ|tên của vợ|tên của chồng", re.I), "spouse_given_names"),
    (re.compile(r"date of previous marriage", re.I), "previous_marriage_date"),
    (re.compile(r"date of marriage|ngày tháng năm kết hôn|ngày kết hôn", re.I), "spouse_marriage_date"),
    (re.compile(r"marriage city|tại thành phố nào", re.I), "spouse_marriage_city"),
    (re.compile(r"marriage state|marriage province|tại tỉnh / bang nào", re.I), "spouse_marriage_state"),
    (re.compile(r"marriage country|tại quốc gia nào", re.I), "spouse_marriage_country"),
    (re.compile(r"^occupation|nghề nghiệp", re.I), "spouse_occupation"),
    (re.compile(r"specify other|ngành nghề gì", re.I), "spouse_occupation_other"),
    (re.compile(r"immigrating.*u\.s.*with you|nhập cư sang mỹ cùng", re.I), "spouse_immigrating"),
    (re.compile(r"do you have any previous spouses|bạn có.*chồng.*vợ trước", re.I), "previous_spouses_used"),
    (re.compile(r"previous spouse.?s name|họ tên.*vợ.*chồng cũ", re.I), "previous_spouse_full_name"),
    (re.compile(r"date of divorce|ngày.*ly hôn", re.I), "previous_divorce_date"),
    (re.compile(r"number of children|bao nhiêu con", re.I), "children_count"),
    (re.compile(r"do you have any children|bạn có con", re.I), "children_used"),
    (re.compile(r"mother.*surnames?|mother.*surname|họ.*mẹ|họ của mẹ", re.I), "mother_surname"),
    (re.compile(r"mother.*given names?|mother.*given|tên.*mẹ", re.I), "mother_given_names"),
    (re.compile(r"mother.*date of birth|ngày.*sinh.*mẹ|date of birth.*mother", re.I), "mother_date_of_birth"),
    (re.compile(r"mother.*city of birth|thành phố.*sinh.*mẹ|city of birth.*mother", re.I), "mother_birth_city"),
    (re.compile(r"mother.*state.*birth|mother.*province.*birth|tỉnh.*sinh.*mẹ|bang.*sinh.*mẹ", re.I), "mother_birth_state"),
    (re.compile(r"mother.*country.*birth|quốc gia.*sinh.*mẹ|country.*birth.*mother", re.I), "mother_birth_country"),
]

# DS-260 C — Previous U.S Travel (mục "Đã từng đến Mỹ") — trước đây KHÔNG có pattern nào nên
# export không bao giờ điền được, dù web hiển thị đúng (báo lỗi thực tế 2026-08-05: web fill
# oke, xuất file Word bị trống — export_ds260.py dò text theo regex trong template, field
# không có pattern thì bị bỏ qua âm thầm, không lỗi/không cảnh báo).
TRAVEL_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"ever been in the u\.?s|đã từng đến mỹ", re.I), "been_in_us"),
    (re.compile(r"travel details.*i-?94|chi tiết.*ngày đến.*ở bao lâu", re.I), "us_travel_history"),
    (re.compile(r"ever been issued a u\.?s visa|đã từng được cấp visa", re.I), "issued_us_visa"),
    (re.compile(r"visa details.*type/date/number|chi tiết visa.*loại.*ngày cấp", re.I), "us_visa_history"),
    (re.compile(r"refused u\.?s visa|admission.*withdrawn|bị từ chối visa", re.I), "refused_us_visa"),
    (re.compile(r"refusal details.*visa type|chi tiết.*loại visa.*ngày.*lý do", re.I), "us_visa_refusal_history"),
]

# DS-260 F — Security and Background Information (44 câu Yes/No) + G — Social Security Number
# (3 câu) — CÙNG LỖI với TRAVEL_LABEL_PATTERNS ở trên: chưa từng có pattern nào, export bỏ
# trống toàn bộ mục An ninh & SSN dù web hiển thị đúng (báo lỗi thực tế 2026-08-05).
SECURITY_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"communicable disease|tuberculosis|bệnh lây nhiễm", re.I), "communicable_disease"),
    (re.compile(r"vaccinations? in accordance|giấy chích ngừa", re.I), "has_vaccination_docs"),
    (re.compile(r"mental or physical disorder|rối loạn tâm thần", re.I), "mental_physical_disorder"),
    # "abuser or addict" (không ràng buộc "drug") — template thật gặp lỗi chính tả "druc abuser"
    # thay vì "drug abuser" (báo lỗi thực tế 2026-08-05); "abuser or addict" đủ đặc trưng để
    # không trùng field nào khác trong danh sách.
    (re.compile(r"abuser or addict|nghiện ma túy", re.I), "drug_abuser"),
    # "ever been arrested" (template thật) — "been" bắt buộc trong câu gốc, để optional cho an toàn.
    (re.compile(r"ever (been )?arrested or convicted|bị bắt giữ", re.I), "arrested_convicted"),
    # family_*/aided phải kiểm tra TRƯỚC bản gốc — câu "người thân" bắt đầu bằng "Are you the
    # spouse, son, or daughter..." NHƯNG thân câu vẫn chứa nguyên cụm của câu gốc ("violated any
    # controlled substance", "committed a human trafficking", "financial assistance to
    # terrorists", "abuse of position...confiscated") nên nếu bản gốc đứng trước sẽ nuốt luôn câu
    # người thân (báo lỗi thực tế 2026-08-05: family_* bị bản gốc cướp → mất field). Word-boundary
    # \b để "son" KHÔNG khớp bên trong "per\bson\bal gain" của chính câu gốc.
    (re.compile(r"\b(spouse|son|daughter|minor child|family member)\b.*controlled substance.*benefited|hưởng lợi buôn ma túy", re.I), "family_drug_trafficking_benefited"),
    (re.compile(r"violat(ed|ion).*controlled substance|vi phạm luật chất cấm", re.I), "violated_controlled_substance"),
    (re.compile(r"prostitution|mại dâm", re.I), "prostitution"),
    (re.compile(r"money laundering|rửa tiền", re.I), "money_laundering"),
    (re.compile(r"aided, abetted|aided.*severe form of trafficking|giúp đỡ.*buôn người", re.I), "human_trafficking_aided"),
    (re.compile(r"\b(spouse|son|daughter|minor child|family member)\b.*human trafficking.*benefited|hưởng lợi buôn người", re.I), "family_human_trafficking_benefited"),
    (re.compile(r"committed.*human trafficking|thực hiện.*buôn người", re.I), "human_trafficking_committed"),
    (re.compile(r"espionage|sabotage|export.*violat|gián điệp", re.I), "espionage_sabotage"),
    # Thứ tự bắt buộc: family_terrorist (thân câu có "financial assistance...terrorist") → financial
    # → activities. financial phải TRƯỚC activities vì câu financial "provide financial assistance
    # or other support to terrorists" từng bị activities cướp (báo lỗi thực tế 2026-08-05).
    (re.compile(r"\b(spouse|son|daughter|minor child|family member)\b.*terrorist activity|người thân hoạt động khủng bố", re.I), "family_terrorist_activity"),
    (re.compile(r"financial (assistance|support).*terrorist|hỗ trợ tài chính khủng bố", re.I), "terrorist_financial_support"),
    (re.compile(r"engage.*terrorist activit|seek to engage.*terrorist|tham gia vào các hoạt động khủng bố", re.I), "terrorist_activities"),
    (re.compile(r"member (of|or representative of) (a )?terrorist organization|thành viên.*tổ chức khủng bố", re.I), "terrorist_org_member"),
    (re.compile(r"participated in genocide|tham gia diệt chủng", re.I), "genocide"),
    (re.compile(r"participated in torture|tham gia tra tấn", re.I), "torture"),
    (re.compile(r"extrajudicial killings|giết người phi pháp", re.I), "extrajudicial_killings"),
    (re.compile(r"recruitment.*use of child soldiers|tuyển dụng binh lính trẻ em", re.I), "child_soldiers"),
    (re.compile(r"religious freedom|vi phạm tự do tôn giáo", re.I), "religious_freedom_violations"),
    (re.compile(r"communist.*totalitarian party|thành viên đảng cộng sản", re.I), "communist_party_member"),
    (re.compile(r"farc|eln|auc|hỗ trợ farc", re.I), "colombia_groups_support"),
    (re.compile(r"\b(spouse|son|daughter|minor child|family member)\b.*confiscat|người thân lạm quyền", re.I), "family_abuse_position_property"),
    (re.compile(r"abuse.*position.*confiscat|lạm quyền chiếm đoạt tài sản", re.I), "abuse_position_confiscate_property"),
    (re.compile(r"population controls?.*abortion|coercive population control|forced abortion|ép phá thai", re.I), "population_control_abortion"),
    (re.compile(r"\b(spouse|son|daughter|minor child|family member)\b.*chemical weapons|người thân tiết lộ bí mật cwc", re.I), "family_chemical_weapons_disclosure"),
    (re.compile(r"disclosed.*confidential.*chemical weapons|tiết lộ bí mật.*vũ khí hóa học", re.I), "chemical_weapons_disclosure"),
    (re.compile(r"willful misrepresentation|visa fraud|gian lận visa", re.I), "visa_fraud"),
    (re.compile(r"removed or deported|bị trục xuất", re.I), "removed_deported"),
    (re.compile(r"withheld.*custody.*u\.?s\.? citizen child|giữ trái phép con là công dân", re.I), "withheld_custody_us_child"),
    (re.compile(r"assist(ed)?.*withhold(ing)?.*custody|giúp giữ trái phép con", re.I), "assisted_withhold_custody"),
    (re.compile(r"voted in.*u\.?s\.? in violation|bỏ phiếu ở mỹ trái luật", re.I), "voted_illegally"),
    (re.compile(r"renounced.*u\.?s\.? citizenship.*tax|từ bỏ quốc tịch mỹ", re.I), "renounced_citizenship_tax"),
    (re.compile(r"attended a public (elementary )?school|attended.*u\.?s\.? public (elementary )?school.*f.*status|học trường công.*visa f", re.I), "attended_public_school_f_status"),
    (re.compile(r"skilled or unskilled labor|lao động chưa được bộ lao động", re.I), "skilled_labor_no_certification"),
    (re.compile(r"graduate of a foreign medical school|tốt nghiệp y khoa nước ngoài", re.I), "foreign_medical_grad_no_exam"),
    (re.compile(r"health.?care worker.*certif|nhân viên y tế chưa có chứng nhận", re.I), "health_care_worker_no_cert"),
    (re.compile(r"permanently ineligible.*citizenship|vĩnh viễn không đủ điều kiện quốc tịch", re.I), "permanently_ineligible_citizenship"),
    (re.compile(r"depart.*u\.?s.*(evade|avoid).*military|rời khỏi mỹ.*(tránh|né).*(quân sự|nghĩa vụ)", re.I), "departed_us_evade_military"),
    (re.compile(r"practic(e|ing) polygamy|đến mỹ để đa thê", re.I), "polygamy"),
    (re.compile(r"former exchange visitor|exchange visitor.*\(j\)|exchange visitor.*(j-?1|two.?year)|visa j", re.I), "exchange_visitor_j_unfulfilled"),
    (re.compile(r"frivolous application for asylum|tị nạn gian dối", re.I), "frivolous_asylum"),
    (re.compile(r"likely to become a public charge|xin trợ cấp chính phủ mỹ", re.I), "public_charge"),
    (re.compile(r"applied for a (u\.?s\.? )?social security number|đã từng xin ssn", re.I), "applied_ssn_before"),
    (re.compile(r"want the (ssa|social security administration).*issue|muốn cấp ssn", re.I), "want_ssn_issued"),
    (re.compile(r"authorize.*disclosure.*(dhs|social security)|cho phép chia sẻ ssn", re.I), "authorize_ssn_disclosure"),
]

# Nhãn nơi sinh / ngày sinh không ghi "father/mother" — map theo section (cha/mẹ/cá nhân).
GENERIC_BIRTH_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"date of birth|ngày tháng năm sinh", re.I), "date_of_birth"),
    (re.compile(r"city of birth|thành phố nơi sinh", re.I), "birth_city"),
    (re.compile(r"state.*province.*birth|tỉnh.*bang nơi sinh", re.I), "birth_state"),
    (re.compile(r"country.*region.*birth|quốc gia nơi sinh", re.I), "birth_country"),
]

APPLICANT_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"name\s*\(last/middle/first\)", re.I), "applicant_name"),
    (re.compile(r"^sex|male or female|nam hay n", re.I), "gender"),
    (re.compile(r"country.*region.*origin|nationality|quốc tịch", re.I), "nationality"),
    (re.compile(r"passport id|số hộ chiếu", re.I), "passport_number"),
    (re.compile(r"country.*authority.*issued|quốc gia cấp.*hộ chiếu", re.I), "passport_issuing_country"),
    (re.compile(r"issuance date|ngày cấp.*hộ chiếu", re.I), "passport_issue_date"),
    (re.compile(r"expiration date|ngày hết hạn.*hộ chiếu", re.I), "passport_expiration_date"),
    (re.compile(r"year of death|năm mất", re.I), "death_date"),
    (re.compile(r"deceased|người mất", re.I), "death_deceased_name"),
    (re.compile(r"husband|chồng", re.I), "marriage_husband_name"),
    (re.compile(r"wife|vợ", re.I), "marriage_wife_name"),
    (re.compile(r"marriage date|ngày cưới", re.I), "marriage_date"),
    (re.compile(r"military.*country", re.I), "military_country"),
    (re.compile(r"military.*branch", re.I), "military_branch"),
    (re.compile(r"military.*rank|cấp bậc", re.I), "military_rank"),
    (re.compile(r"service.*from|phục vụ.*từ", re.I), "military_service_start"),
    (re.compile(r"service.*to|phục vụ.*đến", re.I), "military_service_end"),
    # E.2 Thông tin bổ sung — ngôn ngữ khác & du lịch 5 năm (Yes/No + chi tiết).
    (re.compile(r"other languages beside|ngôn ngữ nào khác", re.I), "other_languages_used"),
    (re.compile(r"traveled to any countries|đã từng du lịch|du lịch đến các nước", re.I), "traveled_countries_5yr_used"),
]

# Mục NGHĨA VỤ QUÂN SỰ — nhãn song ngữ. Gắn context "military" để "Country/Region" KHÔNG rơi vào
# địa chỉ hiện tại. Thứ tự: nhãn đặc thù trước, "country/nước" cuối cùng.
MILITARY_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"have you ever served|đã từng phục vụ trong quân", re.I), "military_served"),
    (re.compile(r"branch|chiến khu", re.I), "military_branch"),
    (re.compile(r"rank|positon|position|cấp bậc|công việc gì", re.I), "military_rank"),
    (re.compile(r"specialty|chuyên ngành", re.I), "military_specialty"),
    (re.compile(r"service\s*from|phục vụ.*từ", re.I), "military_service_start"),
    (re.compile(r"service\s*to|phục vụ.*đến", re.I), "military_service_end"),
    (re.compile(r"country|region|nước|lãnh thổ", re.I), "military_country"),
]

_SECTION_CONTEXT_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"personal information|thông tin cá nhân", re.I), "applicant"),
    (re.compile(r"address\s*[-–]\s*địa chỉ", re.I), "address"),
    (re.compile(r"contact information|thông tin liên lạc", re.I), "contact"),
    (re.compile(r"social media|mạng xã hội", re.I), "social"),
    (re.compile(r"thông tin của cha|father.?s information", re.I), "father"),
    (re.compile(r"thông tin của mẹ|mother.?s information", re.I), "mother"),
    (re.compile(r"thông tin người phối ngẫu cũ|previous spouse", re.I), "previous_spouse"),
    (re.compile(r"thông tin của người phối ngẫu", re.I), "current_spouse"),
    (re.compile(r"thông tin của con cái", re.I), "children"),
    # Work / Education (Section D) — sub-context theo cấp học.
    (re.compile(r"middle school|seconday school|secondary school|cấp 2", re.I), "edu_middle"),
    (re.compile(r"highschool|high school|cấp 3", re.I), "edu_high"),
    (re.compile(r"college\s*/?\s*univer|cao đẳng|đại học", re.I), "edu_college"),
    (re.compile(r"work\s*/?\s*education|primary occupation|nghề nghiệp chính", re.I), "work"),
    # Reset context khi sang phần khác để không kẹt ở work/edu.
    (re.compile(r"additional information|security and background|previous u\.?s.*travel", re.I), "applicant"),
    (re.compile(r"military service|nghĩa vụ quân sự", re.I), "military"),
    # THÔNG TIN KHÁC (E.2) — thoát context military để ngôn ngữ/du lịch khớp ở context applicant.
    (re.compile(r"thông tin khác", re.I), "applicant"),
]

# Section D — Work
WORK_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"primary occupation|nghề nghiệp chính", re.I), "work_primary_occupation"),
    (re.compile(r"specific other|specify other|ngành nghề gì ghi rõ", re.I), "work_occupation_other_specify"),
    (re.compile(r"present employer|công ty hiện tại", re.I), "work_present_employer"),
    (re.compile(r"do you have other occupation|công việc nào khác", re.I), "work_other_occupation_used"),
    (re.compile(r"previously employed|trong vòng 10 năm", re.I), "work_prior_jobs_used"),
    (re.compile(r"postal zone|zip code|mã bưu điện", re.I), "work_employer_postal_code"),
    (re.compile(r"state/province|tỉnh\s*/?\s*bang", re.I), "work_employer_state"),
    (re.compile(r"country/region|quốc gia", re.I), "work_employer_country"),
    (re.compile(r"^\s*city|thành phố", re.I), "work_employer_city"),
    (re.compile(r"address|địa chỉ", re.I), "work_employer_address"),
]

# Section D — Education. Key middle/high là edu_{level}_school_*, riêng college là edu_college_*.
EDU_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"school name|tên trường", re.I), "name"),
    (re.compile(r"major|ngành học", re.I), "major"),
    (re.compile(r"address|địa chỉ", re.I), "address"),
    (re.compile(r"period|thời gian học", re.I), "period"),
]
_EDU_CONTEXT_KEYS: dict[str, dict[str, str]] = {
    "edu_middle": {
        "name": "edu_middle_school_name",
        "address": "edu_middle_school_address",
        "period": "edu_middle_school_period",
    },
    "edu_high": {
        "name": "edu_high_school_name",
        "address": "edu_high_school_address",
        "period": "edu_high_school_period",
    },
    "edu_college": {
        "name": "edu_college_name",
        "address": "edu_college_address",
        "major": "edu_college_major",
        "period": "edu_college_period",
    },
}

_CHILD_INDEX_CONTEXT = re.compile(r"child.?s name\s*\((\d+)\)|con thứ\s*(\d+)", re.I)

GENERIC_ADDRESS_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"postal zone|zip code|mã bưu điện", re.I), "postal"),
    (re.compile(r"from date|ở từ", re.I), "from_date"),
    (re.compile(r"lived anywhere|chỗ khác.*16|từng ở những chổ khác", re.I), "other_addresses"),
    (re.compile(r"state.*province.*birth|tỉnh.*bang.*sinh|state/province of birth", re.I), "birth_state"),
    (re.compile(r"state.*province|tỉnh\s*/?\s*bang", re.I), "state"),
    (re.compile(r"city of birth|thành phố nơi sinh", re.I), "birth_city"),
    (re.compile(r"^city\s*\(|city\s*\(\s*thành phố", re.I), "city"),
    (re.compile(r"current address|địa chỉ hiện tại", re.I), "address_line"),
    (re.compile(r"country/region(?!.*origin)(?!.*nationality)", re.I), "country"),
]

CONTACT_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"primary phone|điện thoại chính", re.I), "primary_phone"),
    (re.compile(r"secondary phone|điện thoại phụ", re.I), "secondary_phone"),
    (re.compile(r"work phone|điện thoại.*làm việc|nơi làm việc", re.I), "work_phone"),
    (re.compile(r"other.*telephone|điện thoại khác.*5|other phone", re.I), "other_phones_used"),
    # "Other Email used..." phải khớp TRƯỚC nhãn email chung (dòng đó chứa "địa chỉ EMAIL KHÁC").
    (re.compile(r"other.*email.*5|email khác.*5", re.I), "other_emails_used"),
    (re.compile(r"email address|địa chỉ email", re.I), "email"),
]

SOCIAL_LABEL_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    # "Other Social Media used..." phải khớp TRƯỚC (dòng đó chứa "TÊN MẠNG XÃ HỘI KHÁC" → đụng identifier).
    (re.compile(r"other social media.*5|mạng xã hội khác.*5", re.I), "other_social_media_used"),
    (re.compile(r"social media provider|social media platform|mạng xã hội nào", re.I), "social_media_platform"),
    (re.compile(r"social media identifier|tên mạng xã hội|link trang", re.I), "social_media_identifier"),
]

_CONTEXT_FIELD_MAP: dict[str, dict[str, str]] = {
    "applicant": {
        "postal": "postal_code",
        "state": "current_state",
        "city": "current_city",
        "country": "current_country",
        "address_line": "current_address",
    },
    "address": {
        "postal": "postal_code",
        "state": "current_state",
        "city": "current_city",
        "country": "current_country",
        "address_line": "current_address",
        "from_date": "address_from_date",
        "other_addresses": "other_addresses_used",
    },
    "contact": {
        "primary_phone": "primary_phone",
        "secondary_phone": "secondary_phone",
        "work_phone": "work_phone",
        "other_phones_used": "other_phones_used",
        "email": "email",
        "other_emails_used": "other_emails_used",
    },
    "social": {
        "social_media_platform": "social_media_platform",
        "social_media_identifier": "social_media_identifier",
        "other_social_media_used": "other_social_media_used",
    },
    "father": {
        "postal": "father_postal_code",
        "state": "father_state",
        "city": "father_city",
        "country": "father_country",
        "address_line": "father_address",
    },
    "mother": {
        "postal": "mother_postal_code",
        "state": "mother_state",
        "city": "mother_city",
        "country": "mother_country",
        "address_line": "mother_address",
    },
    "current_spouse": {
        "postal": "spouse_postal_code",
        "state": "spouse_state",
        "city": "spouse_city",
        "country": "spouse_country",
        "address_line": "spouse_address",
    },
}

_CONTEXT_BIRTH_KEY: dict[str, dict[str, str]] = {
    "applicant": {
        "date_of_birth": "date_of_birth",
        "birth_city": "birth_city",
        "birth_state": "birth_state",
        "birth_country": "birth_country",
    },
    "father": {
        "date_of_birth": "father_date_of_birth",
        "birth_city": "father_birth_city",
        "birth_state": "father_birth_state",
        "birth_country": "father_birth_country",
    },
    "mother": {
        "date_of_birth": "mother_date_of_birth",
        "birth_city": "mother_birth_city",
        "birth_state": "mother_birth_state",
        "birth_country": "mother_birth_country",
    },
    "current_spouse": {
        "date_of_birth": "spouse_date_of_birth",
        "birth_city": "spouse_birth_city",
        "birth_state": "spouse_birth_state",
        "birth_country": "spouse_birth_country",
    },
    "previous_spouse": {
        "date_of_birth": "previous_spouse_date_of_birth",
    },
}

DS260_SOURCE_MARKERS = [
    "PASSPORT",
    "BIRTH CERTIFICATE",
    "MARRIAGE CERTIFICATE",
    "Death certificate",
    "JUDICIAL CERTIFICATE",
    "MILITARY DISCHARGE",
]

_CHILD_BIRTH_CITY_KEY = re.compile(r"^child_\d+_birth_city$")
_CHILD_BIRTH_COUNTRY_KEY = re.compile(r"^child_\d+_birth_country$")


def _format_gender(val: str) -> str:
    return normalize_gender(val)


def _has_meaningful_current_address(out: dict[str, str]) -> bool:
    addr = (out.get("current_address") or "").strip()
    if not addr:
        return False
    low = addr.lower()
    if "immipath" in low or "sẽ điền" in low or "se dien" in low:
        return False
    return True


def _enrich_postal_codes(out: dict[str, str]) -> None:
    """Tự map mã bưu điện VN từ tỉnh/thành — chỉ khi đã có ĐỊA CHỈ HIỆN TẠI (nơi ở/nhận thư).

    CHỈ dùng field địa chỉ HIỆN TẠI của từng người (father_state/city/address,
    mother_state/city/address, ...) — KHÔNG fallback sang *_birth_state/*_birth_city. Nơi sinh và
    nơi ở hiện tại là hai khái niệm khác nhau, mỗi nơi có postal code riêng; suy luận postal code
    (vốn thuộc về ĐỊA CHỈ HIỆN TẠI) từ tỉnh SINH ra một mã sai lệch không phản ánh nơi cha/mẹ thực
    sự đang ở (báo lỗi thực tế 2026-08-05: mother_postal_code bị suy từ mother_birth_state "Nghe
    An" → 460000, trong khi DS-260 gốc của NGUYEN HOANG YEN ghi mẹ ở postal 78607). Không có địa
    chỉ hiện tại → để TRỐNG, không đoán bừa.
    """

    def _set_postal(key: str, *, state: str, city: str, country: str, address: str) -> None:
        if out.get(key):
            return
        code = derive_postal_code_from_location(state=state, city=city, country=country, address=address)
        if code:
            out[key] = code

    if _has_meaningful_current_address(out):
        _set_postal(
            "postal_code",
            state=out.get("current_state") or "",
            city=out.get("current_city") or "",
            country=out.get("current_country") or "",
            address=out.get("current_address") or "",
        )
    else:
        out["postal_code"] = ""
    _set_postal(
        "father_postal_code",
        state=out.get("father_state", ""),
        city=out.get("father_city", ""),
        country=out.get("father_country", ""),
        address=out.get("father_address") or out.get("father_address_line1", ""),
    )
    _set_postal(
        "mother_postal_code",
        state=out.get("mother_state", ""),
        city=out.get("mother_city", ""),
        country=out.get("mother_country", ""),
        address=out.get("mother_address") or out.get("mother_address_line1", ""),
    )
    _set_postal(
        "spouse_postal_code",
        state=out.get("spouse_state", ""),
        city=out.get("spouse_city", ""),
        country=out.get("spouse_country", ""),
        address=out.get("spouse_address") or "",
    )
    _set_postal(
        "work_employer_postal_code",
        state=out.get("work_employer_state", ""),
        city=out.get("work_employer_city", ""),
        country=out.get("work_employer_country", ""),
        address=out.get("work_employer_address", ""),
    )


def _normalize_vn_phone(v: str) -> str:
    """SĐT nội địa VN (bắt đầu '0', 9–11 chữ số) → định dạng quốc tế +84 (bỏ số 0 đầu).
    Giữ nguyên nếu đã có '+', là 'N/A'/trống, hoặc không giống số VN."""
    v = (v or "").strip()
    if not v or v.upper() == "N/A" or v.startswith("+"):
        return v
    digits = re.sub(r"\D", "", v)
    if digits.startswith("0") and 9 <= len(digits) <= 11:
        return "+84" + digits[1:]
    return v


def _normalize_birth_city_state(out: dict[str, str]) -> None:
    """Quy ước DS-260: sinh tại TP trực thuộc TW (Hà Nội/HCM/Hải Phòng/Đà Nẵng/Cần Thơ/Huế)
    → City = tên TP, State/Province = 'N/A'. Chặn lỗi để cả City lẫn State = tên TP."""
    for city_key, state_key in (
        ("birth_city", "birth_state"),
        ("father_birth_city", "father_birth_state"),
        ("mother_birth_city", "mother_birth_state"),
        ("spouse_birth_city", "spouse_birth_state"),
        # Địa chỉ hiện tại ở TP trực thuộc TW cũng theo quy ước State = N/A.
        ("current_city", "current_state"),
        ("father_city", "father_state"),
        ("mother_city", "mother_state"),
        ("spouse_city", "spouse_state"),
    ):
        if canonical_vn_city(out.get(city_key, "")):
            if normalize_location(out.get(state_key, "")) not in ("n/a", "na"):
                out[state_key] = "N/A"


def _cross_check_birth_country(out: dict[str, str]) -> None:
    """Guard xác định: nếu NƠI SINH là địa danh Việt Nam (thành phố/tỉnh) thì Quốc gia nơi sinh
    phải là 'Vietnam'. Chặn lỗi AI trích xuất nhầm nước (vd. birth_city='Can Tho' nhưng
    birth_country='Canada'). Chỉ ép khi nhận diện được địa danh VN — không đụng ca sinh ở nước ngoài."""
    for city_key, state_key, country_key in (
        ("birth_city", "birth_state", "birth_country"),
        ("father_birth_city", "father_birth_state", "father_birth_country"),
        ("mother_birth_city", "mother_birth_state", "mother_birth_country"),
        ("spouse_birth_city", "spouse_birth_state", "spouse_birth_country"),
    ):
        if find_vn_locality(out.get(city_key, "")) or find_vn_locality(out.get(state_key, "")):
            if normalize_location(out.get(country_key, "")) not in {"vietnam", "viet nam", "vietnamese"}:
                out[country_key] = "Vietnam"


_HISTORY_ENTRY_BOUNDARY_RE = re.compile(r";\s+(?=[A-Z][a-z]{2}\s+\d{4}\s*[–—-])")


def _reflow_history_newlines(text: str) -> str:
    """Lịch sử nhiều mốc thời gian (địa chỉ/công việc/du lịch…) đôi khi bị dồn thành 1 dòng,
    phân cách bằng "; " thay vì xuống dòng thật (vd. sau khi qua enrich/loose-match, ranh giới
    \\n gốc không được giữ) — Word cần \\n thật để mỗi mốc nằm 1 dòng riêng qua
    _fill_blank_line_history_block (khách yêu cầu 2026-08-07: "mỗi mốc thời gian thì phải xuống
    dòng"). Khôi phục dựa vào mẫu bắt đầu mốc mới "MMM YYYY–…:" nếu chưa có \\n thật; không đụng
    tới nếu đã có \\n (giữ nguyên ranh giới OCR gốc)."""
    if not text or "\n" in text:
        return text
    return _HISTORY_ENTRY_BOUNDARY_RE.sub("\n", text)


def _prepare_display_values(values: dict[str, str]) -> dict[str, str]:
    out = {k: (v or "").strip() for k, v in values.items()}
    for _hk in _OTHER_USED_HISTORY.values():
        if out.get(_hk):
            out[_hk] = _reflow_history_newlines(out[_hk])
    # children_count là ô SỐ ("Number of Children:") — khi không có con, ds260_mapping cố ý gán
    # "No" cho nó để hiển thị nhất quán trên web (Case C-2: không để trống trong khi anh chị em
    # ghi "No"). Đáp án "No" đó thuộc về câu hỏi children_used ("Do you have any children?"), KHÔNG
    # phải số lượng con — nếu giữ nguyên, dòng "Number of Children (...):" trong Word sẽ hiện chữ
    # "No" thay vì để trống (báo lỗi thực tế 2026-08-05).
    if out.get("children_count", "").strip().lower() in {"no", "không", "khong"}:
        out["children_count"] = ""
    _normalize_birth_city_state(out)
    _cross_check_birth_country(out)
    for _pk in ("primary_phone", "secondary_phone", "work_phone"):
        if out.get(_pk):
            out[_pk] = _normalize_vn_phone(out[_pk])
    fam = out.get("family_name", "")
    given = out.get("given_names", "")
    raw_name = out.get("applicant_name") or f"{fam} {given}".strip()
    native = (out.get("applicant_name_native") or "").strip()
    if not native:
        native = raw_name
    if not native and (fam or given):
        native = f"{fam} {given}".strip()
    if native:
        # Full Name in Native Language → IN HOA, GIỮ dấu tiếng Việt (vd. NGUYỄN VĂN A).
        out["applicant_name_native"] = " ".join(native.split()).upper()
    if raw_name:
        # Name (Last/Middle/First) → IN HOA, KHÔNG dấu (vd. NGUYEN VAN A).
        out["applicant_name"] = format_person_name_ascii(raw_name)
    # City/State nơi sinh đã được chuẩn hóa theo quy ước DS-260 ở resolve
    # (TP trực thuộc TW → City; tỉnh → State; cái còn lại → N/A) — không ép bằng nhau.
    for key in ("birth_city", "birth_state"):
        if out.get(key):
            out[key] = format_birth_city_display(out[key])
    if out.get("mother_birth_city"):
        out["mother_birth_city"] = format_birth_city_display(out["mother_birth_city"])
    if out.get("mother_city"):
        out["mother_city"] = format_birth_city_display(out["mother_city"])
    if out.get("mother_country"):
        out["mother_country"] = format_nationality_country(out["mother_country"])
    if out.get("mother_birth_country"):
        out["mother_birth_country"] = format_nationality_country(out["mother_birth_country"])
    if out.get("spouse_marriage_city"):
        out["spouse_marriage_city"] = format_birth_city_display(out["spouse_marriage_city"])
    if out.get("spouse_birth_city"):
        out["spouse_birth_city"] = format_birth_city_display(out["spouse_birth_city"])
    if out.get("spouse_birth_state"):
        out["spouse_birth_state"] = format_place_name_title(out["spouse_birth_state"])
    for key in list(out.keys()):
        if _CHILD_BIRTH_CITY_KEY.match(key) and out.get(key):
            out[key] = format_birth_city_display(out[key])
        if _CHILD_BIRTH_COUNTRY_KEY.match(key) and out.get(key):
            out[key] = format_nationality_country(out[key])
    if out.get("nationality"):
        out["nationality"] = format_nationality_country(out["nationality"])
    elif out.get("birth_country"):
        out["nationality"] = out["birth_country"]
    if out.get("passport_issuing_country"):
        out["passport_issuing_country"] = format_nationality_country(out["passport_issuing_country"])
    if out.get("birth_country"):
        out["birth_country"] = format_nationality_country(out["birth_country"])
    if out.get("spouse_birth_country"):
        out["spouse_birth_country"] = format_nationality_country(out["spouse_birth_country"])
    if out.get("spouse_marriage_country"):
        out["spouse_marriage_country"] = format_nationality_country(out["spouse_marriage_country"])
    _enrich_postal_codes(out)
    # Địa chỉ → tiếng Anh (Quốc lộ→Highway, Xã→Commune, Huyện/Quận→District, bỏ dấu…).
    # Sau _enrich_postal_codes để không ảnh hưởng suy luận mã bưu điện.
    _ADDRESS_ENGLISH_KEYS = (
        "current_address",
        "current_city",
        "other_addresses_history",
        "spouse_address",
        "father_address",
        "mother_address",
        "work_employer_address",
        "edu_middle_school_address",
        "edu_high_school_address",
        "edu_college_address",
    )
    for key in _ADDRESS_ENGLISH_KEYS:
        if out.get(key):
            out[key] = format_address_english(out[key])
    if out.get("other_addresses_history"):
        # Chốt an toàn cuối cùng trước khi ghi ra Word — phòng dữ liệu không qua ocr_pipeline
        # (nhập tay/nguồn khác) còn sót "hiện tại/đến nay" tiếng Việt (DS-260 chỉ dùng tiếng Anh).
        out["other_addresses_history"] = normalize_present_marker(out["other_addresses_history"])
    for key in list(out.keys()):
        if re.match(r"^child_\d+_current_address$", key) and out.get(key):
            out[key] = format_address_english(out[key])
    for key, val in list(out.items()):
        if not val:
            continue
        if key == "gender":
            out[key] = _format_gender(val)
            continue
        if is_date_field_key(key):
            # Giữ nguyên text nếu không phải 1 ngày chuẩn (vd. "Sep 2008 to Now") — đừng xoá trắng.
            # Xuất Word dùng mốc trọn kiểu Anh, tháng viết tắt ("15 Aug 1990") — dashboard vẫn dd/mm/yyyy
            # (format_sections_date_display, không đổi).
            out[key] = format_ds260_export_date(val) or val
    return out


def _build_label_mapping() -> dict[str, str]:
    mapping: dict[str, str] = {}
    for field in flatten_ds260_mappings().values():
        mapping[field.label] = field.key
        mapping[field.key] = field.key
    mapping.update(
        {
            "Full Name": "applicant_name",
            "Full Name in Native Language": "applicant_name_native",
            "Passport Number": "passport_number",
            "Date of Birth": "date_of_birth",
            "Sex": "gender",
            "Current Marital Status": "current_marital_status",
            "Nationality": "nationality",
        }
    )
    return mapping


def _format_value_after_colon(label_with_colon: str, value: str) -> str:
    label = label_with_colon.rstrip()
    if not label.endswith(":"):
        label = f"{label}:"
    val = value.strip()
    if not val:
        return label.rstrip()
    return f"{label} {val}"


def _is_placeholder_rest(rest: str) -> bool:
    r = rest.strip()
    if not r:
        return True
    if _COLON_FILLER.fullmatch(r):
        return True
    low = r.lower()
    if "immipath" in low or "sẽ điền" in low or "se dien" in low:
        return True
    for marker in DS260_SOURCE_MARKERS:
        if low == marker.lower() or low.startswith(marker.lower()):
            return True
    return False


def _rest_is_answer(rest: str, value: str) -> bool:
    if not value:
        return False
    rest_strip = rest.strip()
    if not rest_strip:
        return True
    if rest_strip == value:
        return True
    if _is_placeholder_rest(rest):
        return True
    m = re.match(r"^([\s\t_\.·…\-\u00a0\u2013\u2014]+)(.*)$", rest, re.DOTALL)
    if m and m.group(2).strip() == value:
        return True
    return False


def _collapse_colon_gap(text: str, value: str) -> str:
    if not value or ":" not in text or "{{" in text:
        return text
    if _URL_LIKE.search(text) or _TIME_LIKE.search(text):
        return text
    idx = text.rfind(":")
    if idx < 0:
        return text
    label = text[: idx + 1]
    rest = text[idx + 1 :]
    if not _rest_is_answer(rest, value):
        return text
    return _format_value_after_colon(label, value)


def _child_birth_key_map(child_n: int) -> dict[str, str]:
    prefix = f"child_{child_n}_"
    return {
        "date_of_birth": f"{prefix}date_of_birth",
        "birth_city": f"{prefix}birth_city",
        "birth_state": f"{prefix}birth_state",
        "birth_country": f"{prefix}birth_country",
    }


def _update_section_context(text: str, context: str) -> str:
    child_match = _CHILD_INDEX_CONTEXT.search(text)
    if child_match:
        n = child_match.group(1) or child_match.group(2)
        return f"child_{n}"
    # Chỉ tiêu đề section (dòng ngắn, KHÔNG phải câu hỏi) mới được đổi context. Câu hỏi Yes/No dài
    # của Section F đôi khi chứa từ khóa section ("public secondary school", "evade military
    # service") làm trôi context sang edu/military rồi nuốt các câu sau (báo lỗi thực tế
    # 2026-08-05). Dòng câu hỏi luôn có "?"; tiêu đề section thì không.
    if "?" in text:
        return context
    for pattern, section in _SECTION_CONTEXT_PATTERNS:
        if pattern.search(text):
            return section
    return context


def _match_ds260_key(text: str, context: str = "applicant") -> str:
    # Section C (Previous U.S Travel) + F (Security) + G (SSN) — dò TRƯỚC mọi cổng context.
    # Đây là các câu hỏi nguyên văn rất đặc trưng (một dòng = một câu), không trùng nhãn ngắn của
    # Work/Edu/Military/Children. Trước đây khi context bị "kẹt" ở child_/edu_/military (do thân câu
    # hỏi vô tình khớp _SECTION_CONTEXT_PATTERNS) thì các cổng đó return sớm và nuốt mất câu — khiến
    # export bỏ trống toàn bộ mục Travel/Security/SSN dù web hiển thị đúng (báo lỗi thực tế
    # 2026-08-05). Đặt ở đây khiến các mục này miễn nhiễm vĩnh viễn với lỗi trôi context.
    for pattern, key in TRAVEL_LABEL_PATTERNS:
        if pattern.search(text):
            return key
    for pattern, key in SECURITY_LABEL_PATTERNS:
        if pattern.search(text):
            return key

    # Work / Education (Section D) — gate theo context để nhãn dùng chung
    # ("Address", "City", "nghề nghiệp"…) không lọt sang field phối ngẫu/cá nhân.
    if context == "work":
        for pattern, key in WORK_LABEL_PATTERNS:
            if pattern.search(text):
                return key
        return ""
    # Nghĩa vụ quân sự — gate riêng để "Country/Region" tiếng Việt không lọt sang địa chỉ.
    if context == "military":
        for pattern, key in MILITARY_LABEL_PATTERNS:
            if pattern.search(text):
                return key
        return ""
    if context in _EDU_CONTEXT_KEYS:
        keys = _EDU_CONTEXT_KEYS[context]
        for pattern, suffix in EDU_LABEL_PATTERNS:
            if pattern.search(text):
                return keys.get(suffix, "")
        return ""

    # Children — gate TRƯỚC DS260_LABEL_PATTERNS để dòng "immigrating to the U.S" của con
    # không bị nhận nhầm là spouse_immigrating.
    if context.startswith("child_"):
        idx_match = re.match(r"child_(\d+)$", context)
        if idx_match:
            idx = int(idx_match.group(1))
            if re.search(r"child.?s name|họ và tên con", text, re.I):
                return f"child_{idx}_full_name"
            if re.search(r"live with you|đang ở với", text, re.I):
                return f"child_{idx}_lives_with"
            if re.search(r"join you in the future|định cư trong tương lai", text, re.I):
                return f"child_{idx}_immigrating_future"
            if re.search(r"immigrating to the u\.s|nhập cư sang mỹ", text, re.I):
                return f"child_{idx}_immigrating"
            if re.search(r"current address|địa chỉ hiện tại", text, re.I):
                return f"child_{idx}_current_address"
            ctx_map = _child_birth_key_map(idx)
            for pattern, generic_key in GENERIC_BIRTH_LABEL_PATTERNS:
                if pattern.search(text):
                    mapped = ctx_map.get(generic_key, "")
                    if mapped:
                        return mapped
            return ""

    # Năm mất cha/mẹ — chỉ trong section cha/mẹ. (Pattern "năm mất" ở APPLICANT_LABEL_PATTERNS map sang
    # death_date của giấy báo tử, KHÔNG áp dụng cho ô này.)
    if context in ("father", "mother") and re.search(r"year of death|năm mất", text, re.I):
        return f"{context}_death_year"

    for pattern, key in DS260_LABEL_PATTERNS:
        if pattern.search(text):
            return key

    ctx_map = _CONTEXT_BIRTH_KEY.get(context, {})
    for pattern, generic_key in GENERIC_BIRTH_LABEL_PATTERNS:
        if pattern.search(text):
            mapped = ctx_map.get(generic_key, "")
            if mapped:
                return mapped
            return ""

    addr_ctx = _CONTEXT_FIELD_MAP.get(context, {})
    if addr_ctx:
        for pattern, generic_key in GENERIC_ADDRESS_LABEL_PATTERNS:
            if not pattern.search(text):
                continue
            if generic_key.startswith("birth_"):
                birth_key = ctx_map.get(generic_key, "")
                if birth_key:
                    return birth_key
                continue
            mapped = addr_ctx.get(generic_key, "")
            if mapped:
                return mapped

    if context == "contact":
        for pattern, key in CONTACT_LABEL_PATTERNS:
            if pattern.search(text):
                return key
    if context == "social":
        for pattern, key in SOCIAL_LABEL_PATTERNS:
            if pattern.search(text):
                return key

    if context in ("applicant", "address"):
        for pattern, key in APPLICANT_LABEL_PATTERNS:
            if pattern.search(text):
                return key
    return ""


_PERIOD_FROM_TO_RE = re.compile(r"(?i)(from\s*\(t[uừ]\))(.*?)(to\s*\(đ[eế]n\)\s*:?)")


def _fill_period_from_to(text: str, value: str) -> str | None:
    """Tách 'Period: from (từ) ... to (đến):' thành 2 mốc ngày (vd. '05/09/1991 - 30/05/1994')."""
    if not _PERIOD_FROM_TO_RE.search(text):
        return None
    parts = re.split(r"\s+[-–—]\s+|\s+to\s+|\s+đến\s+|\s*->\s*|\s*→\s*", value, maxsplit=1, flags=re.I)
    parts = [p.strip() for p in parts if p.strip()]
    if len(parts) < 2:
        return None
    d_from = format_ds260_export_date(parts[0]) or parts[0]
    d_to = format_ds260_export_date(parts[1]) or parts[1]
    return _PERIOD_FROM_TO_RE.sub(
        lambda m: f"{m.group(1)} {d_from}   {m.group(3)} {d_to}", text, count=1
    )


# Field câu hỏi Yes/No thường KHÔNG có dấu ':' trên mẫu — đáp án gắn cuối dòng.
_QUESTION_FILL_KEYS = frozenset({"spouse_immigrating"})
_QUESTION_FILL_KEY_RE = re.compile(r"^child_\d+_(immigrating|immigrating_future|lives_with)$")

# Field có nhãn KHÔNG kết thúc bằng ':' trên mẫu (vd. Social Media Identifier kết thúc bằng "(LINK NGẮN))")
# — vẫn phải điền giá trị, gắn ở cuối dòng.
#
# _match_ds260_key() nhận diện ĐÚNG field_key không có nghĩa là _smart_fill_ds260_line() sẽ
# CHÈN giá trị — dòng không có dấu ':' (đa số câu Yes/No mẫu DS-260) chỉ được điền khi key nằm
# trong 1 trong 3 danh sách trắng này (_is_question_fill_key / _APPEND_NO_COLON_KEYS /
# _OTHER_USED_HISTORY); thiếu thì _smart_fill_ds260_line trả nguyên text, không báo lỗi (báo
# lỗi thực tế 2026-08-05: TRAVEL_LABEL_PATTERNS/SECURITY_LABEL_PATTERNS thêm rồi nhưng field
# xuất Word vẫn trống — do đây là lớp gate THỨ HAI, tách biệt với việc nhận diện field).
#
# 44 câu Section F (Security) + 3 câu Section G (SSN) là Yes/No đơn giản, không có field chi
# tiết đi kèm — đăng ký thẳng vào đây để được gắn giá trị cuối dòng câu hỏi.
_SECURITY_AND_SSN_SIMPLE_YES_NO_KEYS = frozenset({
    "communicable_disease", "has_vaccination_docs", "mental_physical_disorder", "drug_abuser",
    "arrested_convicted", "violated_controlled_substance", "family_drug_trafficking_benefited",
    "prostitution", "money_laundering", "human_trafficking_committed", "human_trafficking_aided",
    "family_human_trafficking_benefited", "espionage_sabotage", "terrorist_activities",
    "terrorist_financial_support", "terrorist_org_member", "family_terrorist_activity",
    "genocide", "torture", "extrajudicial_killings", "child_soldiers",
    "religious_freedom_violations", "communist_party_member", "colombia_groups_support",
    "abuse_position_confiscate_property", "family_abuse_position_property",
    "population_control_abortion", "chemical_weapons_disclosure",
    "family_chemical_weapons_disclosure", "visa_fraud", "removed_deported",
    "withheld_custody_us_child", "assisted_withhold_custody", "voted_illegally",
    "renounced_citizenship_tax", "attended_public_school_f_status",
    "skilled_labor_no_certification", "foreign_medical_grad_no_exam",
    "health_care_worker_no_cert", "permanently_ineligible_citizenship",
    "departed_us_evade_military", "polygamy", "exchange_visitor_j_unfulfilled",
    "frivolous_asylum", "public_charge",
    "applied_ssn_before", "want_ssn_issued", "authorize_ssn_disclosure",
})

_APPEND_NO_COLON_KEYS = frozenset({"social_media_identifier", "military_served", "children_used"}) | (
    _SECURITY_AND_SSN_SIMPLE_YES_NO_KEYS
)

# Câu hỏi Yes/No có phần khai chi tiết. Quy tắc: có chi tiết (history) → "Yes - <chi tiết>";
# không có → "No". Map cờ used → field chi tiết tương ứng.
_OTHER_USED_HISTORY: dict[str, str] = {
    "other_phones_used": "other_phones_history",
    "other_emails_used": "other_emails_history",
    "other_social_media_used": "other_social_history",
    # Section D — Work: nghề khác + lịch sử việc làm 10 năm (narrative phải hiện ra, không chỉ Yes/No).
    "work_other_occupation_used": "work_other_occupation_detail",
    "work_prior_jobs_used": "work_prior_jobs_history",
    # E.2 Thông tin bổ sung — ngôn ngữ khác + du lịch 5 năm.
    "other_languages_used": "other_languages",
    "traveled_countries_5yr_used": "traveled_countries_history",
    # A.3 Địa chỉ — "đã từng ở chỗ khác kể từ 16 tuổi?" → Yes + lịch sử địa chỉ / No.
    "other_addresses_used": "other_addresses_history",
    # Section C — Previous U.S Travel: mẫu ghi "(Yes or No, if 'Yes' write details below)" +
    # nhiều dòng trống bên dưới để viết chi tiết — cùng khuôn "Yes - <chi tiết>" / "No" như các
    # field _used khác, KHÔNG phải field đơn giản (báo lỗi thực tế 2026-08-05).
    "been_in_us": "us_travel_history",
    "issued_us_visa": "us_visa_history",
    "refused_us_visa": "us_visa_refusal_history",
}

_AFFIRMATIVE_TOKENS = frozenset({"yes", "y", "có", "co", "true", "1"})


def _other_used_answer(key: str, values: dict[str, str]) -> str:
    used = (values.get(key) or "").strip()
    history = re.sub(r"\s*\n\s*", "; ", (values.get(_OTHER_USED_HISTORY[key]) or "").strip())
    if history:
        return f"Yes - {history}"
    if used.lower() in _AFFIRMATIVE_TOKENS:
        return "Yes"
    return "No"


# Mọi field "Yes/No + chi tiết" trong _OTHER_USED_HISTORY đều dùng chung layout DS-260 gốc: câu
# hỏi chỉ nhận "Yes -"/"No", còn chi tiết (mỗi mốc/dòng riêng) được điền vào các dòng trống
# "_______" NGAY BÊN DƯỚI câu hỏi, không dồn chung vào cuối dòng câu hỏi (khách yêu cầu
# 2026-08-07 rồi mở rộng lại 2026-08-12: áp dụng cho TẤT CẢ field dạng này, không chỉ vài field).
# An toàn cho field không có dòng trống trong template: _find_blank_run_ahead trả None khi gặp
# '{{' trước khi thấy dòng trống, lúc đó _fill_blank_line_history_block tự rơi về cách cũ (dồn
# "Yes - <chi tiết>" vào cuối dòng câu hỏi) — không mất dữ liệu.
_BLANK_LINE_HISTORY_KEYS = frozenset(_OTHER_USED_HISTORY.keys())
_HISTORY_BLANK_LOOKAHEAD = 25


def _find_blank_run_ahead(paras: list, start_idx: int) -> int | None:
    """Tìm dòng TRỐNG (chỉ khoảng trắng/rỗng) gần nhất từ start_idx để bắt đầu điền chi tiết theo
    từng dòng. Gặp '{{...}}' (field khác chưa render) trước khi thấy dòng trống → trả None, nơi
    gọi sẽ dùng lại cách cũ (dồn hết vào dòng câu hỏi) để không mất dữ liệu."""
    limit = min(len(paras), start_idx + _HISTORY_BLANK_LOOKAHEAD)
    for i in range(start_idx, limit):
        t = paras[i].text
        if "{{" in t:
            return None
        if t.strip() == "":
            return i
    return None


def _insert_blank_line_after(paras: list, idx: int) -> int:
    """Chèn 1 dòng MỚI ngay sau paras[idx] — nhân bản định dạng (cùng font/underline dòng trống
    gốc) cả trong list `paras` LẪN cây XML thật của document (addnext). Dùng khi số mốc thời gian
    nhiều hơn số dòng trống có sẵn trong template, để mỗi mốc luôn nằm 1 dòng riêng thay vì dồn
    chung vào dòng cuối (khách yêu cầu 2026-08-07: "mỗi mốc thời gian thì phải xuống dòng")."""
    template_p = paras[idx]._p
    new_p = copy.deepcopy(template_p)
    template_p.addnext(new_p)
    new_para = Paragraph(new_p, paras[idx]._parent)
    paras.insert(idx + 1, new_para)
    return idx + 1


def _fill_history_into_blank_lines(paras: list, blank_start_idx: int, entries: list[str]) -> int:
    """Điền entries vào các dòng trống liên tiếp từ blank_start_idx — mỗi entry LUÔN có dòng
    riêng: dùng hết dòng trống sẵn có trong template rồi CHÈN THÊM dòng mới cho phần dư, không
    bao giờ dồn nhiều mốc vào chung 1 dòng."""
    i = blank_start_idx
    idx = 0
    while i < len(paras) and paras[i].text.strip() == "" and idx < len(entries):
        paras[i].text = entries[idx]
        idx += 1
        i += 1
    last_used = i - 1
    while idx < len(entries):
        last_used = _insert_blank_line_after(paras, last_used)
        paras[last_used].text = entries[idx]
        idx += 1
        i = last_used + 1
    return i


def _fill_blank_line_history_block(
    paras: list, question_idx: int, key: str, values: dict[str, str]
) -> int:
    """Đáp án Yes/No nằm ở DÒNG GẠCH NGANG ĐẦU TIÊN sau câu hỏi (không nối vào cuối dòng câu
    hỏi/dòng chú thích nữa) — chi tiết (nếu có) xuống các dòng gạch ngang kế tiếp, mỗi mốc 1
    dòng (khách yêu cầu 2026-08-14). Câu hỏi và dòng chú thích "(Yes or No, if 'Yes' write
    details below)..." giữ NGUYÊN VĂN, không chỉnh sửa gì."""
    used = (values.get(key) or "").strip()
    history_raw = (values.get(_OTHER_USED_HISTORY[key]) or "").strip()
    entries = [e.strip() for e in history_raw.split("\n") if e.strip()]
    answer = "Yes" if (entries or used.lower() in _AFFIRMATIVE_TOKENS) else "No"

    blank_idx = _find_blank_run_ahead(paras, question_idx + 1)
    if blank_idx is None:
        # Không có dòng trống nào để dùng — rơi về cách cũ (nhồi vào cuối dòng câu hỏi) để
        # không mất dữ liệu, còn hơn âm thầm bỏ qua.
        question = paras[question_idx]
        suffix = f"Yes - {'; '.join(entries)}" if entries else answer
        new_text = _fill_question_line(question.text, suffix)
        if new_text != question.text:
            question.text = new_text
        return question_idx + 1

    return _fill_history_into_blank_lines(paras, blank_idx, [answer, *entries])


def _is_bare_colon_question(text: str) -> bool:
    """True nếu dòng kết thúc bằng dấu ':' gần như trơ — KHÔNG có khoảng trống/gạch dưới riêng
    ngay trên dòng này để điền câu trả lời (khác đa số field 1-dòng đã có sẵn 1 chuỗi dài
    khoảng trắng/gạch dưới ngay sau dấu ':')."""
    idx = text.rfind(":")
    if idx < 0:
        return False
    rest = text[idx + 1 :]
    return len(rest.strip()) == 0 and len(rest) <= _BARE_COLON_REST_MAX_LEN


def _fill_bare_colon_continuation(paras: list, question_idx: int, value: str) -> bool:
    """Khi câu hỏi kết thúc bằng dấu ':' trơ nhưng dòng NGAY SAU là dòng trống dành riêng để viết
    câu trả lời (mẫu DS-260 gốc hay tách riêng — vd mục địa chỉ hiện tại của con), điền value vào
    dòng đó, KHÔNG nhồi vào cuối dòng câu hỏi. Áp dụng chung cho MỌI field có kiểu layout này
    (khách yêu cầu 2026-08-12: không riêng field nào, gặp layout này ở đâu cũng phải xử lý).
    Trả về False nếu dòng kế tiếp không phải dòng trống dành để viết — nơi gọi sẽ dùng lại cách cũ
    (nhồi vào cuối dòng câu hỏi) để không mất dữ liệu."""
    if not value:
        return False
    if question_idx + 1 >= len(paras):
        return False
    blank_para = paras[question_idx + 1]
    blank_text = blank_para.text
    if blank_text.strip() == "" or not _CONTINUATION_BLANK_RE.match(blank_text):
        return False
    prefix = ":" if blank_text.lstrip().startswith(":") else ""
    blank_para.text = f"{prefix} {value}".strip()
    return True


def _is_question_fill_key(key: str) -> bool:
    return key in _QUESTION_FILL_KEYS or bool(_QUESTION_FILL_KEY_RE.match(key))


def _fill_question_line(text: str, value: str) -> str:
    """Dòng câu hỏi Yes/No không có dấu ':' (vd. 'Does this child live with you? (...?)') — gắn đáp án ở cuối.

    Chống điền trùng bằng cách kiểm tra dòng đã KẾT THÚC bằng đáp án hay chưa — KHÔNG dùng
    "value in text" vì đáp án ngắn ("No"/"Yes") hay lọt vào giữa câu ("No" nằm trong "November",
    "November 30, 1996") khiến câu bị bỏ trống oan (báo lỗi thực tế 2026-08-05:
    attended_public_school_f_status).
    """
    if value and text.rstrip().endswith(value):
        return text
    return f"{text.rstrip()}   {value}"


def _smart_fill_ds260_line(
    text: str,
    values: dict[str, str],
    context: str = "applicant",
    filled: set[str] | None = None,
) -> str:
    if "{{" in text:
        return text
    key = _match_ds260_key(text, context)
    if not key:
        return text
    # Cờ "Other ... used last Five years?" — luôn ghi Yes+chi tiết / No, kể cả khi field trống.
    if key in _OTHER_USED_HISTORY:
        # Câu hỏi có thể trải 2 dòng (Anh + Việt) — chỉ điền đáp án một lần / tài liệu.
        if filled is not None:
            if key in filled:
                return text
            filled.add(key)
        return _fill_question_line(text, _other_used_answer(key, values))
    value = (values.get(key) or "").strip()
    if not value:
        return text
    # Câu hỏi Yes/No đơn giản → LUÔN gắn đáp án ở cuối dòng, kiểm tra TRƯỚC nhánh xử lý dấu ':'.
    # Một số câu Section F có dấu ':' lạc trong phần dịch tiếng Việt (vd. "...bởi Tổng thống MỸ):
    # cá nhân...") — nếu để nhánh ':' chạy trước, code tưởng là dòng "nhãn: giá trị" rồi bỏ qua,
    # khiến câu bị trống dù đã nhận diện đúng field (báo lỗi thực tế 2026-08-05: human_trafficking_aided).
    if _is_question_fill_key(key) or key in _APPEND_NO_COLON_KEYS:
        return _fill_question_line(text, value)
    if ":" not in text:
        return text
    if key.endswith("_period"):
        split = _fill_period_from_to(text, value)
        if split is not None:
            return split
    updated = text
    idx = updated.rfind(":")
    if idx >= 0:
        rest = updated[idx + 1 :].strip()
        if _is_placeholder_rest(rest):
            updated = _format_value_after_colon(updated[: idx + 1], value)
            return updated
        for marker in DS260_SOURCE_MARKERS:
            if rest.upper() == marker.upper() or rest.upper().startswith(marker.upper()):
                updated = _format_value_after_colon(updated[: idx + 1], value)
                return updated
    updated = _collapse_colon_gap(updated, value)
    if re.search(r":\s*$", updated):
        updated = _format_value_after_colon(updated.rstrip(), value)
    return updated


def _build_replacements(values: dict[str, str], mapping: dict[str, str]) -> dict[str, str]:
    reps: dict[str, str] = {}
    for key, val in values.items():
        reps[f"{{{{{key}}}}}"] = val or ""
    for label, field_key in mapping.items():
        reps[f"{{{{{label}}}}}"] = values.get(field_key, "")
    return reps


def _replace_in_paragraph(
    paragraph, values: dict[str, str], mapping: dict[str, str], context: str, filled: set[str] | None = None
) -> None:
    text = paragraph.text
    if "{{" in text:
        for old, new in _build_replacements(values, mapping).items():
            if old in text and new:
                text = text.replace(old, new)
    text = _smart_fill_ds260_line(text, values, context, filled)
    if text != paragraph.text:
        paragraph.text = text


def _fill_paragraphs_with_context(
    paragraphs, values: dict[str, str], mapping: dict[str, str], context: str, filled: set[str] | None = None
) -> str:
    paras = list(paragraphs)
    i = 0
    while i < len(paras):  # len(paras) tính lại mỗi vòng — paras có thể phình ra khi chèn dòng mới
        paragraph = paras[i]
        text = paragraph.text
        context = _update_section_context(text, context)
        key = None if "{{" in text else _match_ds260_key(text, context)
        if key in _BLANK_LINE_HISTORY_KEYS and (filled is None or key not in filled):
            if filled is not None:
                filled.add(key)
            i = _fill_blank_line_history_block(paras, i, key, values)
            continue
        if (
            key
            and key not in _OTHER_USED_HISTORY
            and (filled is None or key not in filled)
            and _is_bare_colon_question(text)
        ):
            value = (values.get(key) or "").strip()
            if value and _fill_bare_colon_continuation(paras, i, value):
                if filled is not None:
                    filled.add(key)
                i += 2
                continue
        _replace_in_paragraph(paragraph, values, mapping, context, filled)
        i += 1
    return context


def fill_ds260_docx_template(
    template_path: Path,
    out_path: Path,
    values: dict[str, str],
    mapping: dict[str, str] | None = None,
) -> None:
    doc = DocxDocument(str(template_path))
    mapping = mapping or _build_label_mapping()
    display = _prepare_display_values(values)

    context = "applicant"
    filled: set[str] = set()
    context = _fill_paragraphs_with_context(doc.paragraphs, display, mapping, context, filled)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                context = _fill_paragraphs_with_context(cell.paragraphs, display, mapping, context, filled)

    doc.save(str(out_path))


async def resolve_ds260_template(
    db: AsyncSession,
    template_code: str | None = None,
) -> FormTemplate:
    codes = [template_code] if template_code else list(DEFAULT_TEMPLATE_FALLBACKS)
    for code in codes:
        if not code:
            continue
        result = await db.execute(select(FormTemplate).where(FormTemplate.code == code))
        tpl = result.scalar_one_or_none()
        if tpl and tpl.template_path and Path(tpl.template_path).exists():
            if not tpl.mapping_config or tpl.mapping_config == "{}":
                tpl.mapping_config = json.dumps(_build_label_mapping())
            return tpl

    result = await db.execute(
        select(FormTemplate)
        .where(
            FormTemplate.is_active.is_(True),
            or_(
                FormTemplate.code.like("ds260%"),
                FormTemplate.code.like("%ds260%"),
                FormTemplate.name.ilike("%ds260%"),
                FormTemplate.name.ilike("%ds-260%"),
            ),
        )
        .order_by(FormTemplate.created_at.desc())
    )
    for tpl in result.scalars().all():
        if tpl.template_path and Path(tpl.template_path).exists():
            return tpl

    return await ensure_ds260_template(db)


async def ensure_ds260_template(db: AsyncSession) -> FormTemplate:
    result = await db.execute(select(FormTemplate).where(FormTemplate.code == DS260_TEMPLATE_CODE))
    template = result.scalar_one_or_none()
    disk_path = settings.templates_path / f"{DS260_TEMPLATE_CODE}.docx"
    mapping = _build_label_mapping()

    if template:
        if disk_path.exists():
            template.template_path = str(disk_path)
        if not template.mapping_config or template.mapping_config == "{}":
            template.mapping_config = json.dumps(mapping)
        return template

    template = FormTemplate(
        code=DS260_TEMPLATE_CODE,
        name="DS-260 Final",
        description="Auto-filled from document mapping table",
        mapping_config=json.dumps(mapping),
        template_path=str(disk_path) if disk_path.exists() else None,
    )
    db.add(template)
    await db.flush()
    return template


def _build_ds260_values(form: dict) -> dict[str, str]:
    values = flatten_ds260_values(form)
    for mapping in flatten_ds260_mappings().values():
        values.setdefault(mapping.key, "")
    return values


def _generate_ds260_table_export(applicant: Applicant, form: dict, out_path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("DS-260 — Immigrant Visa Application", level=0)
    doc.add_paragraph(f"Applicant: {applicant.display_name}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")
    for sec in form.get("sections", []):
        doc.add_heading(sec.get("title", ""), level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"
        hdr[2].text = "Source"
        for field in sec.get("fields", []):
            row = table.add_row().cells
            row[0].text = field.get("label") or field.get("key", "")
            row[1].text = field.get("value") or ""
            src = field.get("source") or {}
            row[2].text = f"{src.get('document_type', '')} · {src.get('source_field', '')}"
        doc.add_paragraph("")
    doc.save(str(out_path))


# Field cần soi khi báo lỗi "web fill oke, export vẫn trống" — dump ra JSON để xem CHÍNH XÁC
# giá trị export nhận được, thay vì đoán mò giữa "code sai" và "data thật sự trống" (báo lỗi
# thực tế 2026-08-05).
_EXPORT_DEBUG_WATCH_KEYS: tuple[str, ...] = (
    "been_in_us", "us_travel_history", "issued_us_visa", "us_visa_history",
    "refused_us_visa", "us_visa_refusal_history",
    "communicable_disease", "has_vaccination_docs", "mental_physical_disorder", "drug_abuser",
    "arrested_convicted", "violated_controlled_substance", "prostitution", "money_laundering",
    "human_trafficking_committed", "human_trafficking_aided", "espionage_sabotage",
    "terrorist_activities", "terrorist_financial_support", "terrorist_org_member", "genocide",
    "torture", "extrajudicial_killings", "child_soldiers", "religious_freedom_violations",
    "communist_party_member", "colombia_groups_support", "abuse_position_confiscate_property",
    "population_control_abortion", "chemical_weapons_disclosure", "visa_fraud",
    "removed_deported", "withheld_custody_us_child", "assisted_withhold_custody",
    "voted_illegally", "renounced_citizenship_tax", "attended_public_school_f_status",
    "skilled_labor_no_certification", "foreign_medical_grad_no_exam",
    "health_care_worker_no_cert", "permanently_ineligible_citizenship",
    "departed_us_evade_military", "polygamy", "exchange_visitor_j_unfulfilled",
    "frivolous_asylum", "public_charge",
    "applied_ssn_before", "want_ssn_issued", "authorize_ssn_disclosure",
)


def _dump_export_debug(applicant: Applicant, values: dict[str, str]) -> None:
    """Lưu giá trị THẬT của các field Section C/F/G tại thời điểm export — để phân biệt
    "code không điền được dù có data" (bug) với "data thật sự trống" (không phải bug, do
    worksheet/OCR của người này thiếu thông tin) mà không cần truy DB trực tiếp."""
    try:
        import json as _json
        from datetime import datetime as _dt, timezone as _tz

        debug_dir = settings.base_dir / "ocr_debug" / "export_ds260_values"
        debug_dir.mkdir(parents=True, exist_ok=True)
        ts = _dt.now(_tz.utc).strftime("%Y%m%dT%H%M%SZ")
        safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", applicant.display_name or "unknown")[:60]
        out_path = debug_dir / f"{ts}_{applicant.id}_{safe_name}.json"
        payload = {
            "applicant_id": str(applicant.id),
            "display_name": applicant.display_name,
            "watched_fields": {k: values.get(k, "<MISSING KEY>") for k in _EXPORT_DEBUG_WATCH_KEYS},
        }
        out_path.write_text(_json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def generate_ds260_export_file(
    applicant: Applicant,
    template: FormTemplate,
    form: dict,
    *,
    member_label: dict | None = None,
) -> Path:
    export_dir = settings.export_path / str(applicant.id)
    export_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    code = template.code or "ds260"
    suffix = ""
    if member_label and member_label.get("display_name"):
        safe = re.sub(r"[^\w]+", "_", member_label["display_name"], flags=re.UNICODE).strip("_")
        if safe:
            suffix = f"_{safe}"
    out_path = export_dir / f"{code}{suffix}_{timestamp}.docx"
    values = _build_ds260_values(form)
    _dump_export_debug(applicant, values)
    mapping = json.loads(template.mapping_config) if template.mapping_config else _build_label_mapping()

    if template.template_path:
        src = Path(template.template_path)
        if src.exists():
            fill_ds260_docx_template(src, out_path, values, mapping)
            return out_path

    _generate_ds260_table_export(applicant, form, out_path)
    return out_path


async def create_ds260_export(
    db: AsyncSession,
    applicant: Applicant,
    *,
    filename_map: dict[str, str] | None = None,
    skip_validation: bool = False,
    template_code: str | None = None,
    member_id=None,
) -> tuple[Export, dict]:
    validation = await validate_ds260(
        db, applicant.id, filename_map=filename_map, member_id=member_id
    )
    if not skip_validation and not validation["valid"]:
        raise ValueError(
            f"DS260 validation failed ({validation['error_count']} errors). "
            "Fix errors before export."
        )

    form = await resolve_ds260_form(
        db, applicant.id, filename_map=filename_map, member_id=member_id
    )
    template = await resolve_ds260_template(db, template_code or DS260_DEFAULT_TEMPLATE_CODE)
    out_path = generate_ds260_export_file(applicant, template, form, member_label=form.get("member"))

    from app.services.export import _finalize_export_file

    stored_path = await _finalize_export_file(out_path, applicant.id)

    export = Export(
        applicant_id=applicant.id,
        template_id=template.id,
        file_path=stored_path,
    )
    db.add(export)
    applicant.status = ApplicantStatus.exported
    await db.flush()
    return export, validation
