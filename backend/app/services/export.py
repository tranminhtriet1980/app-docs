import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.entities import Applicant, ApplicantStatus, Export, FormTemplate, ProfileField
from app.services.field_mapping import FORM_MAPPINGS

PROTECTED_TEMPLATE_CODES = frozenset({"ds160_worksheet", "i539_worksheet"})

LABEL_FIELD_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    (re.compile(r"surname|last name|họ\s*\(", re.IGNORECASE), "identity.family_name"),
    (re.compile(r"middle name|tên đệm", re.IGNORECASE), "identity.middle_name"),
    (re.compile(r"given name|first name|tên\s*\(", re.IGNORECASE), "identity.given_names"),
    (re.compile(r"native language|bản ngữ|họ tên bản ngữ", re.IGNORECASE), "identity.full_name_native"),
    (re.compile(r"other name used|tên khác", re.IGNORECASE), "identity.other_name_used"),
    (re.compile(r"marital status|hôn nhân", re.IGNORECASE), "identity.marital_status"),
    (re.compile(r"sex|male|female|nam hay n|giới tính", re.IGNORECASE), "identity.gender"),
    (re.compile(r"full name|họ và tên", re.IGNORECASE), "identity.full_name"),
    (re.compile(r"date of birth|ng[aà]y th[aá]ng n[aă]m sinh|ngày sinh", re.IGNORECASE), "identity.date_of_birth"),
    (re.compile(r"city of birth|th[aà]nh ph[oố].*n[ơo]i sinh", re.IGNORECASE), "identity.birth_city"),
    (re.compile(r"state.*birth|province.*birth|tỉnh.*sinh|bang.*sinh", re.IGNORECASE), "identity.birth_state"),
    (re.compile(r"country.*birth|qu[oố]c gia n[ơo]i sinh", re.IGNORECASE), "identity.birth_country"),
    (re.compile(r"place of birth|nơi sinh", re.IGNORECASE), "identity.place_of_birth"),
    (re.compile(r"nationality|qu[oố]c t[iị]ch|country.*origin", re.IGNORECASE), "identity.nationality"),
    (re.compile(r"passport id|passport number|s[oố].*h[oộ] chi[eế]u", re.IGNORECASE), "passport.number"),
    (re.compile(r"issued passport|issuing country|qu[oố]c gia c[aấ]p", re.IGNORECASE), "passport.issuing_country"),
    (re.compile(r"issuance date|ng[aà]y c[aấ]p.*h[oộ] chiếu", re.IGNORECASE), "passport.issue_date"),
    (re.compile(r"expiration date|ng[aà]y h[eế]t h[aạ]n", re.IGNORECASE), "passport.expiry_date"),
    (re.compile(r"current address|địa chỉ hiện tại|street", re.IGNORECASE), "contact.address_line1"),
    (re.compile(r"^city|thành phố", re.IGNORECASE), "contact.city"),
    (re.compile(r"state.*province|tỉnh|bang", re.IGNORECASE), "contact.state"),
    (re.compile(r"postal|zip|mã bưu", re.IGNORECASE), "contact.postal_code"),
    (re.compile(r"country.*region|quốc gia", re.IGNORECASE), "contact.country"),
    (re.compile(r"from date|ở từ", re.IGNORECASE), "contact.address_from_date"),
    (re.compile(r"lived anywhere|chỗ khác.*16", re.IGNORECASE), "contact.other_addresses_used"),
    (re.compile(r"prior address|địa chỉ trước|address history", re.IGNORECASE), "contact.other_addresses_history"),
    (re.compile(r"primary phone|điện thoại chính", re.IGNORECASE), "contact.phone_primary"),
    (re.compile(r"secondary phone|điện thoại phụ", re.IGNORECASE), "contact.phone_secondary"),
    (re.compile(r"work phone|điện thoại.*làm việc", re.IGNORECASE), "contact.phone_work"),
    (re.compile(r"other.*telephone|điện thoại khác.*5", re.IGNORECASE), "contact.other_phones_used"),
    (re.compile(r"other.*phone.*detail|chi tiết.*điện thoại", re.IGNORECASE), "contact.other_phones_history"),
    (re.compile(r"email address|địa chỉ email", re.IGNORECASE), "contact.email"),
    (re.compile(r"other.*email.*5|email khác.*5", re.IGNORECASE), "contact.other_emails_used"),
    (re.compile(r"other.*email.*detail|chi tiết.*email", re.IGNORECASE), "contact.other_emails_history"),
    (re.compile(r"social media provider|mạng xã hội nào|platform.*social", re.IGNORECASE), "social.platform"),
    (re.compile(r"social media identifier|link.*mạng|profile.*url", re.IGNORECASE), "social.identifier"),
    (re.compile(r"other.*social.*5|mxh khác.*5|mạng xã hội khác.*5", re.IGNORECASE), "social.other_used"),
    (re.compile(r"other.*social.*detail|chi tiết.*mxh|mạng xã hội khác", re.IGNORECASE), "social.other_history"),
    (re.compile(r"^phone|điện thoại", re.IGNORECASE), "contact.phone_primary"),
    (re.compile(r"\bi-94\b|admission number", re.IGNORECASE), "immigration.i94_number"),
    (re.compile(r"visa number|s[oố].*visa", re.IGNORECASE), "immigration.visa_number"),
    (re.compile(r"father", re.IGNORECASE), "family.father_name"),
    (re.compile(r"mother", re.IGNORECASE), "family.mother_name"),
]

SOURCE_HINT_MARKERS = [
    "PASSPORT",
    "BIRTH CERTIFICATE",
    "VISA",
    "I-94",
    "DIVORCE",
    "MARRIAGE CERTIFICATE",
]

# Gap between label ":" and answer (spaces, tabs, underscores, dots, en-dash fill-ins)
_COLON_FILLER = re.compile(r"[\s\t_\.·…\-\u00a0\u2013\u2014]+")
# Skip times / URLs (12:30, https://)
_TIME_LIKE = re.compile(r":\d{1,2}\b")
_URL_LIKE = re.compile(r"://")


def _format_value_after_colon(label_with_colon: str, value: str, tail: str = "") -> str:
    """Place answer immediately after colon — one space, no wide tab gap."""
    label = label_with_colon.rstrip()
    if not label.endswith(":"):
        label = f"{label}:"
    val = value.strip()
    if not val:
        return f"{label}{tail}".rstrip()
    return f"{label} {val}{tail}".rstrip()


def _rest_is_answer(rest: str, value: str) -> bool:
    """True only when everything after ':' is filler and/or exactly this value (not a substring)."""
    if not value:
        return False
    rest_strip = rest.strip()
    if not rest_strip:
        return True
    if rest_strip == value:
        return True
    if _COLON_FILLER.fullmatch(rest):
        return True
    m = re.match(r"^([\s\t_\.·…\-\u00a0\u2013\u2014]+)(.*)$", rest, re.DOTALL)
    if m and m.group(2).strip() == value:
        return True
    return False


def _collapse_colon_gap(text: str, value: str) -> str:
    """
    Turn 'Label:                     ANSWER' into 'Label: ANSWER'.
    Never truncate when a shorter profile value appears inside a longer answer.
    """
    if not value or ":" not in text or "{{" in text:
        return text
    if _URL_LIKE.search(text) or _TIME_LIKE.search(text):
        return text

    idx = text.rfind(":")
    if idx < 0:
        return text
    label = text[: idx + 1]
    rest = text[idx + 1 :]

    if not _rest_is_answer(rest, value):
        return text

    return _format_value_after_colon(label, value)

def _build_replacements(profile: dict[str, str], mapping: dict[str, str]) -> dict[str, str]:
    reps: dict[str, str] = {}
    for key, val in profile.items():
        reps[f"{{{{{key}}}}}"] = val or ""
    for label, profile_key in mapping.items():
        reps[f"{{{{{label}}}}}"] = profile.get(profile_key, "")
    return reps


def _smart_fill_line(text: str, profile: dict[str, str]) -> str:
    if "{{" in text or ":" not in text:
        return text

    target_key = ""
    for pattern, field_key in LABEL_FIELD_PATTERNS:
        if pattern.search(text):
            target_key = field_key
            break
    if not target_key:
        return text

    value = (profile.get(target_key) or "").strip()
    if not value:
        return text

    updated = text
    for marker in SOURCE_HINT_MARKERS:
        if marker in updated:
            updated = updated.replace(marker, value)

    updated = _collapse_colon_gap(updated, value)

    if re.search(r":\s*$", updated):
        updated = _format_value_after_colon(updated.rstrip(), value)

    return updated


def _replace_in_paragraph(
    paragraph,
    replacements: dict[str, str],
    profile: dict[str, str],
    mapping: dict[str, str] | None = None,
) -> None:
    text = paragraph.text
    if "{{" in text:
        for old, new in replacements.items():
            if old in text and new:
                text = text.replace(old, new)
        if mapping:
            for _label, profile_key in mapping.items():
                val = (profile.get(profile_key) or "").strip()
                if val and val in text:
                    text = _collapse_colon_gap(text, val)
    text = _smart_fill_line(text, profile)
    if text != paragraph.text:
        paragraph.text = text


def _fill_docx_template(template_path: Path, out_path: Path, profile: dict[str, str], mapping: dict[str, str]) -> None:
    doc = DocxDocument(str(template_path))
    replacements = _build_replacements(profile, mapping)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements, profile, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements, profile, mapping)

    doc.save(str(out_path))


async def ensure_default_templates(db: AsyncSession) -> None:
    for code, mapping in FORM_MAPPINGS.items():
        result = await db.execute(select(FormTemplate).where(FormTemplate.code == code))
        existing = result.scalar_one_or_none()
        disk_path = settings.templates_path / f"{code}.docx"
        template_path = str(disk_path) if disk_path.exists() else None

        if existing:
            if template_path and not existing.template_path:
                existing.template_path = template_path
            continue

        name = code.replace("_", " ").upper()
        db.add(
            FormTemplate(
                code=code,
                name=name,
                description=f"Worksheet for {name}",
                mapping_config=json.dumps(mapping),
                template_path=template_path,
            )
        )
    await db.flush()

    # Đăng ký file .docx có trong thư mục nhưng chưa có trong DB
    for docx in settings.templates_path.glob("*.docx"):
        code = docx.stem
        result = await db.execute(select(FormTemplate).where(FormTemplate.code == code))
        if result.scalar_one_or_none():
            continue
        db.add(
            FormTemplate(
                code=code,
                name=code.replace("_", " ").title(),
                description="Custom template from templates/forms/",
                mapping_config=json.dumps({}),
                template_path=str(docx),
            )
        )
    await db.flush()


def _profile_dict(fields: list[ProfileField]) -> dict[str, str]:
    return {f.field_key: f.field_value or "" for f in fields}


def _generate_table_export(applicant: Applicant, template: FormTemplate, profile: dict[str, str], out_path: Path) -> None:
    mapping = json.loads(template.mapping_config)
    doc = DocxDocument()
    doc.add_heading(f"{template.name}", level=0)
    doc.add_paragraph(f"Applicant: {applicant.display_name}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    for label, profile_key in mapping.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = profile.get(profile_key, "")

    doc.add_paragraph("")
    p = doc.add_paragraph(
        "DISCLAIMER: AI-assisted draft. Verify all fields against original documents before submission."
    )
    p.runs[0].font.size = Pt(9)

    doc.save(str(out_path))


def generate_word_export(
    applicant: Applicant,
    template: FormTemplate,
    profile: dict[str, str],
) -> Path:
    export_dir = settings.export_path / str(applicant.id)
    export_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    out_path = export_dir / f"{template.code}_{timestamp}.docx"

    mapping = json.loads(template.mapping_config) if template.mapping_config else {}

    if template.template_path:
        src = Path(template.template_path)
        if src.exists():
            _fill_docx_template(src, out_path, profile, mapping)
            return out_path

    _generate_table_export(applicant, template, profile, out_path)
    return out_path


async def register_template_file(
    db: AsyncSession,
    code: str,
    name: str,
    file_path: Path,
    mapping: dict[str, str] | None = None,
) -> FormTemplate:
    code = re.sub(r"[^a-zA-Z0-9_-]", "_", code).lower()
    result = await db.execute(select(FormTemplate).where(FormTemplate.code == code))
    template = result.scalar_one_or_none()
    mapping_json = json.dumps(mapping or FORM_MAPPINGS.get(code, {}))

    if template:
        template.name = name
        template.template_path = str(file_path)
        template.mapping_config = mapping_json
    else:
        template = FormTemplate(
            code=code,
            name=name,
            description="Custom uploaded form template",
            mapping_config=mapping_json,
            template_path=str(file_path),
        )
        db.add(template)

    await db.flush()
    return template


async def delete_form_template(db: AsyncSession, template_id) -> FormTemplate:
    """Xóa template upload — gỡ file .docx trên disk và bản ghi DB."""
    from fastapi import HTTPException

    tpl = await db.get(FormTemplate, template_id)
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    if tpl.code in PROTECTED_TEMPLATE_CODES:
        raise HTTPException(status_code=400, detail="Không thể xóa template hệ thống mặc định")

    export_count = await db.scalar(
        select(func.count()).select_from(Export).where(Export.template_id == template_id)
    )
    if export_count and export_count > 0:
        raise HTTPException(
            status_code=400,
            detail=f"Template đã dùng trong {export_count} lần xuất file. Hãy 'Tắt' thay vì xóa.",
        )

    if tpl.template_path:
        path = Path(tpl.template_path)
        try:
            if path.is_file() and path.resolve().parent == settings.templates_path.resolve():
                path.unlink(missing_ok=True)
        except OSError:
            pass

    await db.delete(tpl)
    await db.flush()
    return tpl


async def create_export(
    db: AsyncSession,
    applicant: Applicant,
    fields: list[ProfileField],
    template_code: str,
) -> Export:
    result = await db.execute(select(FormTemplate).where(FormTemplate.code == template_code))
    template = result.scalar_one_or_none()
    if not template:
        raise ValueError(f"Unknown template: {template_code}")

    profile = _profile_dict(fields)
    out_path = generate_word_export(applicant, template, profile)

    export = Export(
        applicant_id=applicant.id,
        template_id=template.id,
        file_path=str(out_path),
    )
    db.add(export)
    applicant.status = ApplicantStatus.exported
    await db.flush()
    return export
