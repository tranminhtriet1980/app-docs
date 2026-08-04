import base64
import json
import re
from pathlib import Path

from openai import APIStatusError, RateLimitError
from sqlalchemy import delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.entities import Applicant, Document, DocumentStatus, ExtractedField
from app.services.document_registry import (
    RECORDABLE_DOC_TYPES,
    RECORDABLE_REGISTRY_BY_CODE,
    parse_document_filename,
)
from app.services.ds260_mapping import get_extract_keys_for_doc_type
from app.services.field_mapping import DOCUMENT_TYPES
from app.services.llm_client import get_ocr_client, get_openai_client, is_openai_configured
from app.services.llm_usage import UsageContext, chat_completion

QUOTA_WARNING = (
    "OpenAI hết quota/credits (lỗi 429). Đã dùng chế độ demo từ tên file — "
    "nạp tiền tại https://platform.openai.com/account/billing rồi upload lại."
)

IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tif", ".tiff"}

CLASSIFICATION_PROMPT = """Bạn là hệ thống phân loại hồ sơ DS260 / định cư Mỹ.

Bước 1: Xác định loại tài liệu từ nội dung (ảnh/PDF). Các loại chuẩn:
- passport
- judicial_certificate
- divorce
- birth_certificate
- death_certificate
- marriage_certificate
- birth_certificate_child
- military_discharge

Các loại khác (nếu không thuộc 8 loại trên): visa, i20, i94, diploma_transcript, financial, employment_letter, application_form, address_document, ds260_customer_form, photo, other

application_form = đơn xin việc / Job Application — có mục "Applicant's Information" (địa chỉ cá nhân, số điện thoại, email, chức danh, lịch sử việc làm, học vấn). KHÁC employment_letter (chỉ xác nhận việc làm hiện tại). Đây là giấy tờ chính thức để đối chiếu địa chỉ trong DS-260 worksheet.

ds260_customer_form = bản DS-260 khách tự khai (worksheet khách điền tay/đánh máy). Gồm cả:
(a) bản CŨ ImmiPath nhiều mục (cá nhân, hộ chiếu, ĐỊA CHỈ, LIÊN LẠC, MXH, cha/mẹ/phối ngẫu/con,
    công việc/học vấn…) — thường có nhãn tiếng Việt trong ngoặc và khối "From/To" cho địa chỉ;
(b) bản MỚI trùng mẫu DS-260 export.
Chọn loại này cho MỌI worksheet khách tự khai, kể cả khi có nhiều mục — KHÔNG hạ xuống "other" hay
"address_document". address_document chỉ dành cho hoá đơn/hợp đồng thuê/giấy xác nhận cư trú riêng lẻ.

QUY TẮC CHỐNG BỊA:
- KHÔNG ĐƯỢC đoán/suy luận loại tài liệu khi không rõ ràng từ nội dung.
- Nếu ảnh không đủ sáng, mờ, bị cắt, hoặc nội dung không đọc được → đừng cố đoán.
- Nếu không chắc chắn → ghi "other" với confidence thấp (< 0.5) và giải thích lý do trong "reason".
- LUÔN ưu tiên chính xác hơn tìm kiếm sự hoàn hảo — nếu không rõ thì nói không rõ, không bịa loại.

Trả về ONLY valid JSON:
{"document_type": "<code>", "confidence": 0.0-1.0, "reason": "brief"}"""

EXTRACTION_PROMPT_TEMPLATE = """Bạn trích xuất dữ liệu có cấu trúc cho form DS260.

Loại tài liệu đã xác định: {doc_type}

CHỈ trích xuất các field thuộc loại tài liệu này — KHÔNG gộp field trùng tên từ loại khác.
Ví dụ passport có full_name; birth_certificate cũng có full_name — chỉ trích field của {doc_type}.

Trả về ONLY valid JSON:
{{
  "document_type": "{doc_type}",
  "image_quality": {{"readable": true/false, "issue": "<'' nếu ổn, hoặc 'blurry'/'glare'/'too_dark'/'cut_off'/'handwriting_illegible'>"}},
  "fields": {{
    "<field_key>": {{"value": "<string or null>", "confidence": 0.0-1.0, "source_page": "page_1"}}
  }},
  "warnings": ["<danh sách các vấn đề/không rõ ràng>"]
}}

QUY TẮC CHỐNG BỊA DỮ LIỆU:
- CHỈ extract những gì thực sự thấy RÕ RÀNG trên tài liệu. KHÔNG được đoán/suy luận/fill gaps.
- KHÔNG infer giá trị dựa trên ngữ cảnh hay logic. Ví dụ: không cố "guess" ngày sinh từ tuổi, không điền địa chỉ từ giả định.
- Nếu field bị mờ/loá/thiếu sáng/bị cắt/chữ viết tay khó đọc → để null trong value và LUÔN thêm warning chi tiết.
- Nếu có phần nào không đọc được → set image_quality.readable = false, dù phần khác vẫn trích được. Không "hy vọng" phần còn lại sẽ bù được.
- KHÔNG bao giờ bịa dữ liệu để field "có vẻ đầy đủ hay hợp lý". Thà để null còn hơn bịa.
- Nếu phát hiện lỗi (ảnh xoay, định dạng kỳ lạ, dữ liệu mâu thuẫn...) → thêm warning rõ ràng, KHÔNG bỏ qua.

QUY TẮC ĐỊNH DẠNG:
- ISO dates YYYY-MM-DD khi có thể
- UPPERCASE cho tên như in trên giấy
- null chỉ khi thật sự không đọc được
- image_quality.readable = false khi ảnh mờ/loá/thiếu sáng/bị cắt khiến BẤT KỲ phần nội dung nào không đọc được rõ ràng.
- CHỈ dùng các key sau (không thêm key ngoài schema):
{expected_keys}"""


def _is_quota_error(exc: Exception) -> bool:
    if isinstance(exc, RateLimitError):
        return True
    if isinstance(exc, APIStatusError) and exc.status_code in (429, 402):
        return True
    msg = str(exc).lower()
    return "insufficient_quota" in msg or "exceeded your current quota" in msg


def _encode_image(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    mime = "image/jpeg"
    if suffix == ".png":
        mime = "image/png"
    elif suffix == ".webp":
        mime = "image/webp"
    elif suffix == ".gif":
        mime = "image/gif"
    data = base64.standard_b64encode(path.read_bytes()).decode("utf-8")
    return mime, data


# Cạnh dài tối đa (px) của ảnh render gửi Vision API. Một số PDF scan có trang khổ
# siêu lớn (mediabox phóng đại) → get_pixmap(dpi) sinh ảnh hàng chục nghìn px, payload
# base64 lên tới hàng trăm MB khiến cổng Vision API trả 502 Bad Gateway. Kẹp cạnh dài +
# nén JPEG để payload chỉ còn vài MB/trang mà vẫn đủ nét cho OCR.
_MAX_RENDER_PX = 2600
_RENDER_JPEG_QUALITY = 85


# Cạnh dài chuẩn hoá trước khi đo độ nét: phương sai Laplacian tăng theo độ phân giải
# (ảnh gốc càng to càng nhiều chi tiết cạnh) nên phải resize về cùng kích thước mới so
# sánh được với MỘT ngưỡng cố định — nếu không, ảnh chụp 12MP tự nhiên "nét" hơn ảnh
# chụp cùng cảnh ở độ phân giải thấp dù độ mờ thật giống nhau.
_BLUR_MEASURE_PX = 1000

# Kernel Laplacian 3x3 chuẩn — làm nổi cạnh; ảnh mờ có ít cạnh sắc → phương sai thấp.
_LAPLACIAN_KERNEL = (0, 1, 0, 1, -4, 1, 0, 1, 0)


def _measure_blur_variance(path: Path) -> float | None:
    """Phương sai Laplacian trên ảnh grayscale — thấp = mờ/out nét. None nếu đọc lỗi."""
    try:
        from PIL import Image, ImageFilter, ImageStat

        with Image.open(path) as img:
            gray = img.convert("L")
            if max(gray.size) > _BLUR_MEASURE_PX:
                scale = _BLUR_MEASURE_PX / max(gray.size)
                gray = gray.resize(
                    (max(1, round(gray.width * scale)), max(1, round(gray.height * scale)))
                )
            edges = gray.filter(ImageFilter.Kernel((3, 3), _LAPLACIAN_KERNEL, scale=1))
            return ImageStat.Stat(edges).var[0]
    except Exception:
        return None


def _pdf_pages_to_images(path: Path, max_pages: int = 3, dpi: int = 180) -> list[Path]:
    """Render PDF pages to JPEG for vision API. PyMuPDF first (no Poppler), then pdf2image.

    Cạnh dài mỗi ảnh bị kẹp ≤ _MAX_RENDER_PX bất kể kích thước trang PDF, tránh payload
    khổng lồ (→ 502 Bad Gateway) với PDF có trang khổ lớn.
    """
    images: list[Path] = []

    try:
        import fitz

        doc = fitz.open(str(path))
        for i in range(min(len(doc), max_pages)):
            # DPI + max_px trong tên cache → đổi tham số sẽ render lại, không dùng nhầm ảnh cũ.
            out = path.parent / f"{path.stem}_page{i + 1}_dpi{dpi}_px{_MAX_RENDER_PX}.jpg"
            if not out.exists() or out.stat().st_mtime < path.stat().st_mtime:
                page = doc[i]
                zoom = dpi / 72.0
                longest = max(page.rect.width, page.rect.height) * zoom
                if longest > _MAX_RENDER_PX:
                    zoom *= _MAX_RENDER_PX / longest
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                out.write_bytes(pix.tobytes(output="jpg", jpg_quality=_RENDER_JPEG_QUALITY))
            images.append(out)
        doc.close()
        if images:
            return images
    except Exception:
        pass

    try:
        from pdf2image import convert_from_path

        rendered = convert_from_path(str(path), first_page=1, last_page=max_pages, dpi=dpi)
        for i, img in enumerate(rendered):
            out = path.parent / f"{path.stem}_page{i + 1}_dpi{dpi}_px{_MAX_RENDER_PX}.jpg"
            if max(img.size) > _MAX_RENDER_PX:
                scale = _MAX_RENDER_PX / max(img.size)
                img = img.resize((max(1, round(img.width * scale)), max(1, round(img.height * scale))))
            img.convert("RGB").save(str(out), "JPEG", quality=_RENDER_JPEG_QUALITY)
            images.append(out)
    except Exception:
        pass

    return images


# Vision API (detail=high) thu ảnh vào khung 2048x2048 rồi co/giãn sao cho CẠNH NGẮN
# bằng 768px. Nên render DPI cao KHÔNG tự động làm chữ nét hơn: cả trang A4 dựng đứng
# luôn về ~768x1086, chữ tay chỉ còn quãng nửa số điểm ảnh ban đầu.
#
# Cắt trang thành băng NGANG thì cạnh ngắn của mỗi băng là CHIỀU CAO, và chiều cao được
# kéo lên 768 thay vì bị nén xuống — theo cơ chế trên thì phần chữ giữ được nhiều điểm
# ảnh gốc hơn. CẢNH BÁO: đây là suy ra từ quy tắc co giãn, CHƯA đối chứng bằng kết quả
# OCR thật, và ràng buộc thắt cổ chai trên thực tế không phải độ nét mà là trần
# tokens-per-minute (xem ocr_max_images_per_request). Đừng tăng rows dựa vào comment này
# — hãy đo trên hồ sơ thật trước.
#
# Chồng lấn rộng ở trục ngang vì form là cặp NHÃN (trái) – CÂU TRẢ LỜI viết tay (phải),
# cắt dọc giữa trang dễ tách đôi cặp này.
_TILE_OVERLAP_X = 0.18
_TILE_OVERLAP_Y = 0.08


def _split_image_into_tiles(path: Path, rows: int, cols: int) -> list[Path]:
    """Cắt ảnh thành lưới rows x cols chồng lấn. Lỗi/không cần cắt → trả [path]."""
    if rows < 2 and cols < 2:
        return [path]
    try:
        from PIL import Image

        with Image.open(path) as img:
            width, height = img.size
            tile_w, tile_h = width / cols, height / rows
            ov_x, ov_y = tile_w * _TILE_OVERLAP_X, tile_h * _TILE_OVERLAP_Y
            out: list[Path] = []
            for r in range(rows):
                for c in range(cols):
                    left = max(0, int(c * tile_w - ov_x))
                    right = min(width, int((c + 1) * tile_w + ov_x))
                    top = max(0, int(r * tile_h - ov_y))
                    bottom = min(height, int((r + 1) * tile_h + ov_y))
                    dest = path.parent / f"{path.stem}_r{r + 1}c{c + 1}of{rows}x{cols}.jpg"
                    if not dest.exists() or dest.stat().st_mtime < path.stat().st_mtime:
                        img.crop((left, top, right, bottom)).convert("RGB").save(
                            str(dest), "JPEG", quality=_RENDER_JPEG_QUALITY
                        )
                    out.append(dest)
            return out
    except Exception:
        return [path]


def _split_images_into_batches(
    image_paths: list[Path], image_labels: list[str] | None, max_per_batch: int
) -> list[tuple[list[Path], list[str] | None]]:
    """Chia ảnh thành nhiều lô để mỗi request nằm dưới trần tokens-per-minute.

    OpenAI áp trần TPM cho TỪNG request: một request vượt trần thì hỏng vĩnh viễn,
    retry bao nhiêu lần cũng vô ích (sự cố prod 30/07/2026 — worksheet 20 trang cắt 4x1
    thành 80 ảnh ≈ 116k token, đập vào trần 30k). Chia nhỏ rồi gộp kết quả thì mỗi
    request đều lọt.
    """
    if max_per_batch <= 0 or len(image_paths) <= max_per_batch:
        return [(image_paths, image_labels)]

    batches: list[tuple[list[Path], list[str] | None]] = []
    for start in range(0, len(image_paths), max_per_batch):
        stop = start + max_per_batch
        batches.append(
            (
                image_paths[start:stop],
                image_labels[start:stop] if image_labels else None,
            )
        )
    return batches


def _is_blank_for_merge(value: object) -> bool:
    """Giá trị coi như model không đọc được — để lô sau có quyền điền vào.

    Tên phải khác `_is_empty_value` (đã có ở dưới, dùng cho việc khác): trùng tên thì
    định nghĩa sau ghi đè định nghĩa trước và hàm gộp lặng lẽ dùng nhầm luật.
    """
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == ""
    if isinstance(value, (list, dict)):
        return len(value) == 0
    return False


def _merge_batch_results(results: list[dict]) -> dict:
    """Gộp kết quả nhiều lô: giá trị KHÁC RỖNG đầu tiên thắng.

    Mỗi lô chỉ nhìn thấy một phần tài liệu nên phần lớn field sẽ rỗng; lô nào thực sự
    đọc được trang chứa field đó mới có giá trị. Không ghi đè giá trị đã có để một lô
    đoán mò (bịa ra dù không nhìn thấy trang tương ứng) không đè mất số liệu đọc thật.
    """
    if not results:
        return {}
    if len(results) == 1:
        return results[0]

    merged: dict = {}
    # image_quality không theo luật "khác rỗng đầu tiên thắng": một lô bất kỳ báo ảnh khó
    # đọc thì cả tài liệu coi như có vấn đề, dù lô đầu tiên báo "readable": true.
    issues = []
    unreadable = False
    all_warnings: list[str] = []
    for result in results:
        if not isinstance(result, dict):
            continue
        q = result.get("image_quality")
        if isinstance(q, dict):
            if q.get("readable") is False:
                unreadable = True
            issue = str(q.get("issue") or "").strip()
            if issue and issue not in issues:
                issues.append(issue)
        warnings = result.get("warnings")
        if isinstance(warnings, list):
            for w in warnings:
                w_str = str(w).strip()
                if w_str and w_str not in all_warnings:
                    all_warnings.append(w_str)
    if unreadable or issues:
        merged["image_quality"] = {"readable": not unreadable, "issue": ", ".join(issues)}
    if all_warnings:
        merged["warnings"] = all_warnings

    for result in results:
        if not isinstance(result, dict):
            continue
        for key, value in result.items():
            if key in ("image_quality", "warnings"):
                continue  # đã gộp riêng ở trên bằng luật OR, không phải "khác rỗng đầu tiên thắng"
            if key not in merged or _is_blank_for_merge(merged[key]):
                merged[key] = value
    return merged


def _pdf_text_excerpt(path: Path, max_chars: int = 12000) -> str:
    """Extract plain text from PDF pages (fallback when vision conversion unavailable)."""
    chunks: list[str] = []

    try:
        import fitz

        doc = fitz.open(str(path))
        for i in range(min(len(doc), 8)):
            text = (doc[i].get_text() or "").strip()
            if text:
                chunks.append(text)
            if sum(len(c) for c in chunks) >= max_chars:
                break
        doc.close()
    except Exception:
        pass

    if not chunks:
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            for page in reader.pages[:8]:
                text = (page.extract_text() or "").strip()
                if text:
                    chunks.append(text)
                if sum(len(c) for c in chunks) >= max_chars:
                    break
        except Exception:
            pass

    merged = "\n\n".join(chunks).strip()
    if len(merged) > max_chars:
        merged = merged[:max_chars]
    return merged


# DS-260 worksheet (.docx) is long (full form ~60k chars) — cap generously so later sections
# (cha/mẹ, công việc/học vấn, du lịch, quân sự) không bị cắt khỏi prompt trích xuất.
_DOCX_MAX_CHARS = 80000


def _docx_text_excerpt(path: Path, max_chars: int = _DOCX_MAX_CHARS) -> str:
    """Extract text from a Word worksheet — paragraphs + table cells (DS-260 form is a label/value table)."""
    try:
        from docx import Document as _DocxDocument
    except ImportError:
        return ""
    try:
        doc = _DocxDocument(str(path))
    except Exception:
        return ""
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    merged = "\n".join(parts).strip()
    return merged[:max_chars]


def _resolve_document_images(file_path: Path, max_pages: int = 3, dpi: int = 180) -> tuple[list[Path], str]:
    """
    Return image paths suitable for vision API and a hint about source.
    hint: direct_image | pdf_pages | pdf_text_only | filename_only
    """
    suffix = file_path.suffix.lower()
    if suffix in IMAGE_SUFFIXES:
        return [file_path], "direct_image"
    if suffix == ".pdf":
        pages = _pdf_pages_to_images(file_path, max_pages=max_pages, dpi=dpi)
        if pages:
            return pages, "pdf_pages"
        text = _pdf_text_excerpt(file_path)
        if text:
            return [], "pdf_text_only"
        return [], "filename_only"
    if suffix == ".txt":
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
            if text:
                return [], "plain_text"
        except OSError:
            pass
        return [], "filename_only"
    if suffix in {".xlsx", ".xls"}:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            cells: list[str] = []
            for sheet in wb.worksheets[:3]:
                for row in sheet.iter_rows(max_row=80, values_only=True):
                    cells.extend(str(c) for c in row if c is not None)
            wb.close()
            text = " ".join(cells).strip()
            if text:
                return [], "plain_text"
        except Exception:
            pass
        return [], "filename_only"
    if suffix == ".docx":
        if _docx_text_excerpt(file_path):
            return [], "plain_text"
        return [], "filename_only"
    return [], "filename_only"


def _mock_classification(filename: str) -> dict:
    registry_type, _is_exc = parse_document_filename(filename)
    if registry_type:
        defn = registry_type
        return {"document_type": defn, "confidence": 0.85, "reason": "filename registry (standard template)"}

    name = filename.lower()
    rules = [
        ("passport", "passport"),
        ("visa", "visa"),
        ("i-20", "i20"),
        ("i20", "i20"),
        ("i-94", "i94"),
        ("i94", "i94"),
        ("birth", "birth_certificate"),
        ("marriage", "marriage_certificate"),
        ("divorce", "divorce"),
        ("judicial", "judicial_certificate"),
        ("death", "death_certificate"),
        ("military", "military_discharge"),
        ("transcript", "diploma_transcript"),
        ("diploma", "diploma_transcript"),
        ("job application", "application_form"),
        ("application form", "application_form"),
        ("employment", "employment_letter"),
        ("financial", "financial"),
        ("address", "address_document"),
        ("utility", "address_document"),
        ("lease", "address_document"),
        ("residence", "address_document"),
        ("ds260", "ds260_customer_form"),
        ("ds-260", "ds260_customer_form"),
        ("photo", "photo"),
    ]
    for token, doc_type in rules:
        if token in name:
            return {"document_type": doc_type, "confidence": 0.75, "reason": "filename heuristic (demo)"}
    return {"document_type": "other", "confidence": 0.5, "reason": "unknown (demo)"}


def _mock_extraction(doc_type: str, filename: str) -> dict:
    keys = get_extract_keys_for_doc_type(doc_type)
    if not keys:
        keys = ["notes"]
    person = _person_name_from_filename(filename)
    fields = {}
    for key in keys:
        if key in {"family_name", "surname", "last_name"}:
            val = person["family_name"] or "NGUYEN"
        elif key in {"given_names", "given_name", "first_name"}:
            val = person["given_names"] or "VAN A"
        elif key in {"birth_city", "city_of_birth"}:
            val = "HANOI"
        elif key in {"birth_state", "state_of_birth"}:
            val = ""
        elif key in {"birth_country", "country_of_birth"}:
            val = "VIETNAM"
        elif key in {"full_name", "name", "child_full_name"}:
            val = person["full_name"] or "NGUYEN VAN A"
        elif key in {"date_of_birth", "child_date_of_birth"}:
            val = "1990-01-15"
        elif key == "passport_number":
            val = "B1234567"
        elif key == "nationality":
            val = "VIETNAM"
        elif key in {"gender", "sex", "child_gender"}:
            val = "MALE"
        elif key in {"service_from_date", "service_to_date"}:
            val = "2010-01-01" if "from" in key else "2012-01-01"
        else:
            val = f"[demo — {doc_type}]"
        fields[key] = {"value": val, "confidence": 0.5, "source_page": "demo"}
    return {"document_type": doc_type, "fields": fields}


_BIRTH_CERT_KEY_REMAP: dict[str, str] = {
    "mother_full_name": "mother_name",
    "mother_residence": "mother_address",
    "mother_nationality": "mother_country",
    "mother_year_of_birth": "mother_date_of_birth",
    "mother_birth_date": "mother_date_of_birth",
    "mother_dob": "mother_date_of_birth",
    "mother_birth_place": "mother_place_of_birth",
    "mother_city_of_birth": "mother_birth_city",
    "father_full_name": "father_name",
    "father_residence": "father_address",
    "father_nationality": "father_country",
    "father_year_of_birth": "father_date_of_birth",
    "father_birth_date": "father_date_of_birth",
    "father_dob": "father_date_of_birth",
    "father_birth_place": "father_place_of_birth",
    "father_city_of_birth": "father_birth_city",
}


def _field_meta_value(meta: object) -> str:
    if isinstance(meta, dict):
        val = meta.get("value")
        return "" if val is None else str(val).strip()
    return "" if meta is None else str(meta).strip()


def _coerce_birth_certificate_extraction(extraction: dict) -> dict:
    """Remap common LLM keys on English/VN birth certs before schema filter."""
    fields = extraction.get("fields")
    if not isinstance(fields, dict):
        return extraction

    remapped: dict[str, dict] = {}
    for key, meta in fields.items():
        if not isinstance(meta, dict):
            continue
        target = _BIRTH_CERT_KEY_REMAP.get(key, key)
        existing = remapped.get(target)
        if existing and _field_meta_value(existing) and not _field_meta_value(meta):
            continue
        if existing and _field_meta_value(existing) and _field_meta_value(meta):
            if len(_field_meta_value(meta)) <= len(_field_meta_value(existing)):
                continue
        remapped[target] = meta

    extraction["fields"] = remapped
    return extraction


def _filter_extraction_to_schema(doc_type: str, extraction: dict) -> dict:
    """Chỉ giữ field thuộc schema loại tài liệu — tránh merge field trùng tên."""
    allowed = set(get_extract_keys_for_doc_type(doc_type))
    if not allowed:
        return extraction
    fields = extraction.get("fields")
    if not isinstance(fields, dict):
        return extraction
    extraction["fields"] = {k: v for k, v in fields.items() if k in allowed}
    extraction["document_type"] = doc_type
    return extraction


def _resolve_document_type(ai_doc_type: str, filename: str) -> tuple[str, bool]:
    """
    AI phân loại trước; tên file chỉ fallback khi AI không nhận ra loại chuẩn.
    _new → variant exception, cùng doc_type (alias normalization).
    """
    registry_type, is_exception = parse_document_filename(filename or "")
    ai = (ai_doc_type or "other").strip().lower()

    if ai in RECORDABLE_DOC_TYPES:
        return ai, is_exception
    if registry_type:
        return registry_type, is_exception
    if ai in DOCUMENT_TYPES:
        return ai, is_exception
    return "other", is_exception


def _person_name_from_filename(filename: str) -> dict[str, str]:
    """Extract uppercase person name hints from filename."""
    stem = Path(filename).stem
    text = re.sub(r"[_\-]+", " ", stem)
    text = re.sub(r"[()]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()

    # If filename has "A & B", take the first person as primary subject.
    if "&" in text:
        text = text.split("&", 1)[0].strip()

    stop_words = {
        "PASSPORT", "PHOTO", "BIRTH", "CERTIFICATE", "JUDICIAL", "DIVORCE",
        "DECREE", "MARRIAGE", "VISA", "I", "I20", "I94", "COPY", "SCAN", "ORIGINAL",
        "DEATH", "MILITARY", "DISCHARGE", "CHILD", "NEW",
    }
    parts = []
    for p in text.upper().split(" "):
        if not p or p in stop_words:
            continue
        if p.isdigit():
            continue
        parts.append(p)

    if not parts:
        return {"full_name": "", "family_name": "", "given_names": ""}
    full_name = " ".join(parts)
    family_name = parts[0]
    given_names = " ".join(parts[1:]) if len(parts) > 1 else ""
    return {
        "full_name": full_name,
        "family_name": family_name,
        "given_names": given_names,
    }


def _is_empty_value(value: object) -> bool:
    if value is None:
        return True
    if isinstance(value, str) and not value.strip():
        return True
    return False


def _fallback_enrich_extraction(doc_type: str, filename: str, extraction: dict) -> dict:
    """If OCR returns many nulls, enrich from filename."""
    fields = extraction.get("fields")
    if not isinstance(fields, dict):
        return extraction

    person = _person_name_from_filename(filename)

    def set_if_empty(key: str, value: str, conf: float = 0.35) -> None:
        if not value:
            return
        if key not in fields:
            fields[key] = {"value": value, "confidence": conf, "source_page": "filename"}
            return
        meta = fields.get(key) or {}
        if not isinstance(meta, dict):
            return
        if _is_empty_value(meta.get("value")):
            meta["value"] = value
            meta["confidence"] = max(float(meta.get("confidence") or 0), conf)
            meta["source_page"] = meta.get("source_page") or "filename"
            fields[key] = meta

    set_if_empty("full_name", person["full_name"])
    set_if_empty("name", person["full_name"])
    set_if_empty("surname", person["family_name"])
    set_if_empty("family_name", person["family_name"])
    set_if_empty("given_name", person["given_names"].split(" ")[0] if person["given_names"] else "")
    set_if_empty("given_names", person["given_names"])

    extraction["fields"] = fields
    return extraction


def _parse_json_response(text: str) -> dict:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    return json.loads(text)


def _build_llm_user_content(
    *,
    filename: str,
    image_paths: list[Path],
    hint: str,
    file_path: Path,
    task: str,
    image_labels: list[str] | None = None,
    batch_position: tuple[int, int] | None = None,
) -> list[dict] | str:
    """Build OpenAI user message content (multimodal or text-only)."""
    if image_paths:
        parts: list[dict] = []
        intro = f"{task} Filename: {filename}."
        if batch_position:
            # Tài liệu bị chia làm nhiều request cho khỏi vượt trần TPM. Nói rõ để model
            # KHÔNG đoán field nằm ở phần nó không nhìn thấy — kết quả các phần được gộp
            # theo luật "giá trị khác rỗng đầu tiên thắng", nên một giá trị bịa sẽ khoá
            # mất giá trị đọc thật ở phần sau.
            part_no, part_total = batch_position
            intro += (
                f" IMPORTANT: this is PART {part_no} of {part_total} of the document —"
                " you are seeing only some of its pages. Extract only what is VISIBLE in"
                " these images. Leave every other field empty; do NOT guess or infer"
                " fields whose page is not shown here."
            )
        if image_labels:
            intro += (
                f" Document has {len(image_labels)} image(s): each page is split into"
                " OVERLAPPING tiles so handwriting is legible. Tiles of the same page are"
                " parts of ONE page — reassemble them (left tile holds the printed label,"
                " right tile usually holds the handwritten answer for the SAME row)."
                " Because tiles overlap, the same text may appear twice — report it once."
            )
        elif hint == "pdf_pages":
            intro += f" Document has {len(image_paths)} page image(s) — read all pages."
        parts.append({"type": "text", "text": intro})
        for i, img_path in enumerate(image_paths):
            mime, data = _encode_image(img_path)
            label = (
                image_labels[i]
                if image_labels and i < len(image_labels)
                else f"Page {i + 1}"
            )
            parts.append({"type": "text", "text": f"--- {label} ---"})
            parts.append(
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{data}", "detail": "high"}}
            )
        return parts

    if hint == "pdf_text_only":
        pdf_text = _pdf_text_excerpt(file_path)
        return (
            f"{task} Filename: {filename}. PDF scanned without image renderer — use OCR text below:\n\n"
            f"{pdf_text}"
        )

    if hint == "plain_text":
        cap = 12000
        try:
            if file_path.suffix.lower() in {".xlsx", ".xls"}:
                from openpyxl import load_workbook

                wb = load_workbook(file_path, read_only=True, data_only=True)
                cells: list[str] = []
                for sheet in wb.worksheets[:3]:
                    for row in sheet.iter_rows(max_row=80, values_only=True):
                        cells.extend(str(c) for c in row if c is not None)
                wb.close()
                text = "\n".join(cells)
            elif file_path.suffix.lower() == ".docx":
                text = _docx_text_excerpt(file_path)
                cap = _DOCX_MAX_CHARS
            else:
                text = file_path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            text = ""
        return f"{task} Filename: {filename}. Document text:\n\n{text[:cap]}"

    return f"{task} Filename: {filename}. No image available — infer document type from filename only."


async def _openai_classify(
    file_path: Path, filename: str, db: AsyncSession, ctx: UsageContext
) -> dict:
    image_paths, hint = _resolve_document_images(file_path)
    content = _build_llm_user_content(
        filename=filename,
        image_paths=image_paths,
        hint=hint,
        file_path=file_path,
        task="Classify this document.",
    )
    messages: list[dict] = [
        {"role": "system", "content": CLASSIFICATION_PROMPT},
        {"role": "user", "content": content},
    ]

    response = await chat_completion(
        db,
        operation="document.classify",
        context=ctx,
        messages=messages,
        client=get_ocr_client(),
        model=settings.ocr_model,
        response_format={"type": "json_object"},
        temperature=0,
    )
    return _parse_json_response(response.choices[0].message.content or "{}")


ADDRESS_EXTRACT_HINT = """
For ADDRESS / CONTACT documents (utility bill, lease, DS-160 address or contact section), extract:

Address: current_address, city, state/province, postal_code, country, address_from_date,
other_addresses_since_16 (YES/NO), other_addresses_history (prior residences with date ranges).

Phone: primary_phone_number, secondary_phone_number, work_phone_number,
other_phones_used (YES/NO), other_phones_history (numbers used in last 5 years if YES).

Email: email_address, other_emails_used (YES/NO), other_emails_history (other emails in last 5 years if YES).

Social media: social_media_platform (e.g. Facebook – Display Name), social_media_identifier (profile URL or handle),
other_social_media_used (YES/NO), other_social_history (other platforms in last 5 years if YES, e.g. Zalo – Name).
"""

SOCIAL_EXTRACT_HINT = """
For DS-160 SOCIAL MEDIA section, extract:
social_media_platform, social_media_identifier (URL or short link),
other_social_media_used (YES/NO), other_social_history (platform + identifier if YES).
"""

DS260_REFERENCE_EXTRACT_HINT = """
This is a CUSTOMER-REFERENCE upload (_new) — often a filled DS-260 worksheet scan.
Extract ALL DS-260 sections visible on the document, including:

ADDRESS (section 3): current_address, address_city, address_state, postal_code, address_country,
address_from_date, other_addresses_since_16 (Yes/No), other_addresses_history (prior addresses with dates).
The legacy/old-version worksheet prints A.3 as repeated labeled blocks (Vietnamese + English), e.g.:
  (Street)(Đường) ... (City)(Thành phố) ... (State/Province)(Tỉnh/Bang) ...
  (Country)(Quốc gia) ... (Postal code)(Mã bưu điện) ...
  From (Từ tháng/năm)(MM/YYYY) ... To (Đến tháng/năm)(MM/YYYY) ...
Read EVERY block. The block whose To date is the most recent / blank / "present" is the CURRENT
address → current_address (Street), address_city (City), address_state (State/Province),
address_country (Country), postal_code, address_from_date (that block's From). Put each earlier
block on its own line in other_addresses_history as "Street, City, State, Country (MM/YYYY–MM/YYYY)"
and set other_addresses_since_16 = Yes. Disregard impossible To dates (To before From — OCR/typo)
when choosing the current block.

CONTACT (section 4): primary_phone_number, secondary_phone_number, work_phone_number,
other_phones_used (Yes/No), other_phones_history,
email_address, other_emails_used (Yes/No), other_emails_history.

SOCIAL MEDIA (section 5): social_media_platform, social_media_identifier,
other_social_media_used (Yes/No), other_social_history.

RESIDENCE-HISTORY rules (the form's "Lưu ý") — these add ADDRESS entries to
other_addresses_history, they are NOT child/birth data:
  - time spent studying or working in another locality → the real address there;
  - (male) military service → the address(es) while serving;
  - (female) gave birth in another locality → that locality's address.
Capture each such period as its own line in other_addresses_history
("Street, City, State, Country (MM/YYYY–MM/YYYY)") and set other_addresses_since_16 = Yes.

Use Yes/No for yes-no questions. Dates as YYYY-MM-DD or dd/mm/yyyy as printed.
"""

PASSPORT_EXTRACT_HINT = """
For PASSPORT documents (especially Vietnamese passport), extract ALL visible fields:

Personal: full_name, family_name, given_names, date_of_birth (YYYY-MM-DD),
place_of_birth (full text as printed — used for State/Province and Country of Birth on DS-260),
birth_city if shown separately, gender/sex (M/F or MALE/FEMALE), nationality,
id_card_number (Số CMND/CCCD / ID card N°).

Do NOT copy nationality into birth_country — country of birth is derived from place_of_birth separately.
Do NOT extract birth_state or birth_country as separate fields unless explicitly printed on the passport.

Passport document: passport_type or type (e.g. P), country_code (e.g. VNM),
passport_number, issue_date, expiration_date/expiry_date,
place_of_issue / issuing_authority (e.g. Cục Quản lý xuất nhập cảnh),
issuing_country if shown separately from place of issue.

Use ISO dates YYYY-MM-DD when possible. Names as printed on passport (uppercase OK).
"""

JUDICIAL_EXTRACT_HINT = """
For JUDICIAL CERTIFICATE (lý lịch tư pháp), extract:
full_name, date_of_birth, nationality, father_name, mother_name,
document_number, issue_date, document_type.
"""

BIRTH_CERT_EXTRACT_HINT = """
For BIRTH CERTIFICATE (giấy khai sinh chủ hồ sơ — applicant's own birth cert), extract ALL visible fields.

Applicant: full_name, date_of_birth (YYYY-MM-DD), place_of_birth, gender, registration_number.

English translated Vietnamese birth certificates (header "BIRTH CERTIFICATE", "Socialist Republic of Vietnam")
use labeled blocks — map them exactly:
  Mother block:
    "Mother's full name" → mother_name (also split mother_surname + mother_given_names if possible)
    "Year of birth" (often full date like 01 January 1954) → mother_date_of_birth as YYYY-MM-DD
    "Nationality" under mother → mother_country (e.g. Vietnamese)
    "Residence" under mother → mother_address AND mother_place_of_birth (full address)
    City in residence (e.g. Da Nang City) → mother_city
  Father block (same pattern):
    father_name, father_date_of_birth, father_country, father_address/father_place_of_birth, father_city

Vietnamese originals: mother_surname/mother_given_names, mother_date_of_birth, mother_birth_city,
mother_place_of_birth, mother_country.

If the document has NO father or mother information, leave all corresponding parent fields empty.
Use ISO dates YYYY-MM-DD when the FULL date (day+month+year) is printed. If only a YEAR is legible
(vd. 1945), output just the year "1945" — do NOT invent day/month ("01 January 1945"). Same for
month+year only → "MM/YYYY". Do not use applicant nationality for parent country fields.
Do NOT skip parent date/address/nationality when they are printed on the certificate.
"""

DIVORCE_EXTRACT_HINT = """
For DIVORCE DECREE, extract BOTH parties:
husband_full_name, wife_full_name, husband_date_of_birth, wife_date_of_birth,
marriage_date, divorce_date, document_number, document_type.
If date of birth is printed for either party, use ISO YYYY-MM-DD.
Names UPPERCASE as printed (strip titles like MR./MRS. only in values, not keys).
"""

MARRIAGE_CERT_EXTRACT_HINT = """
For MARRIAGE CERTIFICATE (giấy kết hôn / giấy chứng nhận kết hôn), extract BOTH parties:

Husband (chồng): husband_full_name, husband_surname, husband_given_names,
husband_date_of_birth, husband_birth_city, husband_place_of_birth, husband_birth_country,
husband_address, husband_occupation.

Wife (vợ): wife_full_name, wife_surname, wife_given_names,
wife_date_of_birth, wife_birth_city, wife_place_of_birth, wife_birth_country,
wife_address, wife_occupation.

Marriage: marriage_date (YYYY-MM-DD), marriage_place (full), marriage_city, marriage_state, marriage_country.
document_number / registration_number if shown.
Use ISO dates. Names UPPERCASE as printed.
"""

DEATH_EXTRACT_HINT = """
For DEATH CERTIFICATE, extract:
deceased_full_name, date_of_death, place_of_death, relationship_to_applicant, document_number.
"""

CHILD_BIRTH_EXTRACT_HINT = """
For BIRTH CERTIFICATE CHILD (giấy khai sinh con), extract:
child_full_name, child_date_of_birth, child_place_of_birth,
child_birth_city, child_birth_state, child_birth_country, child_gender,
father_name, mother_name, registration_number.
Use ISO dates YYYY-MM-DD. Names UPPERCASE as printed.
If place of birth is one line, also split into city/state/country when possible.
"""

MILITARY_EXTRACT_HINT = """
For MILITARY DISCHARGE, extract:
full_name, military_country, military_branch, military_rank, military_specialty,
service_from_date, service_to_date, document_number.
"""

DS260_CUSTOMER_FORM_EXTRACT_HINT = """
This is a FULL CUSTOMER DS-260 WORKSHEET (ImmiPath form — all sections). Extract EVERY filled field.

This is often a SCANNED, HANDWRITTEN form — read each page carefully, including later pages
(security questions, Social Security). Read the customer's answer written after each label.
Name handling — read EXACTLY as written, do not guess or merge:
  - "Surnames (Họ)" = family name ONLY (e.g. LE, HUYNH, LAM); "Given Names (Tên)" = the rest
    (e.g. VAN TOT, THI KIM PHUC, THI MUOI). Keep father/mother/spouse surname vs given separate.
  - NEVER put the form header/title into a name field: ignore "DS-260", "DS 260", "DS260",
    "KHACH KHAI", "BANG CAU HOI" — these are NOT the applicant's name.
Free-text answers (10-year job history, military service, prior addresses, travel) — transcribe the
WHOLE text verbatim, do not shorten to "Yes". If a long answer is hard to read, capture what is legible.

PERSONAL (section 1): applicant_name, applicant_name_native, other_name_used, other_names,
gender/sex, current_marital_status, date_of_birth, birth_city, birth_state, birth_country,
nationality, id_card_number.

PASSPORT (section 2): passport_number, passport_type, country_code, passport_issue_date,
passport_expiration_date, passport_place_of_issue, passport_issuing_country,
other_nationality_used, other_nationality_history.

ADDRESS (section 3): current_address, address_city/current_city, address_state/current_state,
postal_code, address_country/current_country, address_from_date,
other_addresses_since_16 (Yes/No), other_addresses_history.
The legacy/old-version worksheet prints A.3 as repeated labeled blocks (Vietnamese + English), e.g.:
  (Street)(Đường) ... (City)(Thành phố) ... (State/Province)(Tỉnh/Bang) ...
  (Country)(Quốc gia) ... (Postal code)(Mã bưu điện) ...
  From (Từ tháng/năm)(MM/YYYY) ... To (Đến tháng/năm)(MM/YYYY) ...
Read EVERY block. The block whose To date is the most recent / blank / "present" is the CURRENT
address → current_address (Street), address_city (City), address_state (State/Province),
address_country (Country), postal_code, address_from_date (that block's From). Put each earlier
block on its own line in other_addresses_history as "Street, City, State, Country (MM/YYYY–MM/YYYY)"
and set other_addresses_since_16 = Yes. Disregard impossible To dates (To before From — OCR/typo)
when choosing the current block.
RESIDENCE-HISTORY rules (the form's "Lưu ý") add ADDRESS entries to other_addresses_history — they
are NOT child/birth data: time studying/working in another locality → that real address; (male)
military service → the address while serving; (female) gave birth in another locality → that
locality's address. List each as its own line and set other_addresses_since_16 = Yes.

CONTACT (section 4): primary_phone_number, secondary_phone_number, work_phone_number,
other_phones_used (Yes/No), other_phones_history,
email_address, other_emails_used (Yes/No), other_emails_history.

SOCIAL (section 5): social_media_platform, social_media_identifier,
other_social_media_used (Yes/No), other_social_history.

FATHER: father_surname, father_given_names, father_date_of_birth, father_birth_city,
father_birth_state, father_birth_country, father_is_living, father_death_year,
father_address/current_address (in father section), father_city, father_state, father_country.

MOTHER: mother_surname, mother_given_names, mother_date_of_birth, mother_birth_city,
mother_birth_state, mother_birth_country, mother_is_living, mother_death_year,
mother_address, mother_city, mother_state, mother_country.

SPOUSE: spouse_surname, spouse_given_names, spouse_full_name, spouse_date_of_birth,
spouse_birth_city, spouse_birth_state, spouse_birth_country, spouse_address,
spouse_occupation, spouse_marriage_date, spouse_marriage_city, spouse_marriage_state,
spouse_marriage_country, marriage_husband_name, marriage_wife_name.

PREVIOUS SPOUSE / DIVORCE: previous_spouses_used, previous_spouse_full_name,
previous_spouse_date_of_birth, previous_divorce_date, previous_marriage_date.

CHILDREN: children_used (Yes/No), children_count,
child_1_full_name, child_1_date_of_birth, child_1_birth_city, child_1_birth_state, child_1_birth_country,
child_2_full_name, child_2_date_of_birth, child_3_full_name, child_3_date_of_birth.

WORK / EDUCATION (section D) — Vietnamese header "CÔNG VIỆC / HỌC VẤN", "Work /Education /Training":
primary_occupation (NGHỀ NGHIỆP CHÍNH, e.g. Owner/Manager/Student),
occupation_other_specify (NGÀNH NGHỀ ghi rõ, e.g. Tailor),
present_employer (CÔNG TY hiện tại / trường học hiện tại, e.g. Self-employed),
employer_address, employer_city, employer_state, employer_postal_code, employer_country,
job_title, employment_start_date,
other_occupation_used (Yes/No), other_occupation_detail,
prior_jobs_10_years_used (Yes/No),
prior_jobs_history = FULL narrative of past 10 years employment exactly as written, keep every line
  (e.g. "From 01 January 2009 to 31 December 2023 / Occupation: Manager / Company name: ... /
   Company address: ... / Supervisor name: ... / Supervisor phone number: ...").
EDUCATION (cấp 2/cấp 3/đại học) — for each level capture name, address and period "from ... to ...":
middle_school_name (Cấp 2 = Trung học cơ sở / THCS / Secondary), middle_school_address, middle_school_period,
high_school_name (Cấp 3 = Trung học phổ thông / THPT / Highschool), high_school_address, high_school_period,
college_name (Cao đẳng/Đại học), college_address, college_major, college_period.
Map by level: "Trung học cơ sở"/"THCS" → middle_school_*, "Trung học phổ thông"/"THPT" → high_school_*.

MILITARY: military_country, military_branch, military_rank, military_specialty,
military_service_start, military_service_end.

QUY TẮC BẮT BUỘC (chống bịa/nhầm dữ liệu):
- NGÀY: chỉ ghi đúng những gì đọc được. Nếu chỉ có NĂM (vd. 1945) → ghi đúng "1945", TUYỆT ĐỐI KHÔNG
  tự thêm ngày/tháng (không được ghi "01 January 1945", "01/01/1945"). Nếu có tháng+năm → "MM/YYYY".
- QUỐC GIA NƠI SINH: người mang hộ chiếu Việt Nam sinh tại một địa danh Việt Nam (tỉnh/thành VN) thì
  birth_country / *_birth_country = "Vietnam". KHÔNG suy ra hay bịa nước khác (Canada, USA…) khi không có
  bằng chứng rõ ràng trên giấy tờ.
- NGHỀ vs CÔNG TY: primary_occupation = CHỨC DANH nghề, DỊCH sang tiếng Anh (vd. "Nhân viên bán hàng" →
  "Sales Staff", "Thợ may" → "Tailor"); present_employer = TÊN công ty/cửa hàng/trường (vd. "Dung Grocery
  Store"). KHÔNG đặt chức danh vào ô công ty và ngược lại; không để trống occupation nếu đọc được chức danh.

C — LỊCH SỬ ĐẾN MỸ: been_in_us (Yes/No), issued_us_visa (Yes/No), refused_us_visa (Yes/No),
us_travel_history (mỗi đợt: ngày đến – ngày đi – loại visa), us_visa_history.

E.2 — THÔNG TIN KHÁC: other_languages_used (Yes/No) và other_languages_list — LIỆT KÊ đúng
ngôn ngữ khách ghi (vd. "Tiếng Anh" → "English"). ĐỪNG để trống rồi mặc định "No" khi khách
CÓ ghi ngôn ngữ. traveled_countries_5yr_used (Yes/No) + traveled_countries_history.

F — AN NINH & LÝ LỊCH: đọc HẾT các câu Yes/No ở những trang cuối (bệnh lây nhiễm, giấy chích
ngừa, ma túy, tiền án, mại dâm, rửa tiền, buôn người, khủng bố, trục xuất…). Mỗi câu trả lời
đúng như khách đánh dấu. Nếu khách BỎ TRỐNG một câu thì để null — KHÔNG tự suy thành Yes/No.
Đặc biệt has_vaccination_docs: ghi đúng lựa chọn của khách.

G — SỐ AN SINH XÃ HỘI: applied_ssn_before (Yes/No), want_ssn_issued (Yes/No),
authorize_ssn_disclosure (Yes/No), ssn_number nếu có.

Use mapping keys above (English snake_case). Yes/No for yes-no questions.
Dates as printed (dd/mm/yyyy or month/year). Vietnamese section headers: THÔNG TIN CÁ NHÂN, HỘ CHIẾU, ĐỊA CHỈ, CÔNG VIỆC/HỌC VẤN, THÔNG TIN LIÊN LẠC, MẠNG XÃ HỘI, THÔNG TIN CỦA CHA/MẸ/PHỐI NGẪU/CON, AN NINH, SỐ AN SINH XÃ HỘI.
"""

APPLICATION_FORM_EXTRACT_HINT = """
IMMIGRATION APPLICATION FORM — DS-260 Part D (Work / Education / Training).

PERSONAL ADDRESS (nếu form có mục địa chỉ RIÊNG của người khai, KHÁC employer_address —
đây là nơi người khai đang ở, dùng để đối chiếu với worksheet DS-260, KHÔNG phải địa chỉ
công ty): current_address (street), address_city, address_state, postal_code, address_country.

WORK: primary_occupation, occupation_other_specify, present_employer, employer_name,
employer_address, employer_city, employer_state, employer_postal_code, employer_country,
job_title, employment_start_date, prior_jobs_history.

EDUCATION: middle_school_name (Cấp 2 / Trung học cơ sở / THCS),
middle_school_address, middle_school_period,
high_school_name (Cấp 3 / Trung học phổ thông / THPT), high_school_address, high_school_period,
college_name, college_address, college_major, college_period.

Yes/No for yes-no questions. Dates as printed.
"""

DOC_TYPE_EXTRACT_HINTS: dict[str, str] = {
    "passport": PASSPORT_EXTRACT_HINT,
    "birth_certificate": BIRTH_CERT_EXTRACT_HINT,
    "judicial_certificate": JUDICIAL_EXTRACT_HINT,
    "divorce": DIVORCE_EXTRACT_HINT,
    "marriage_certificate": MARRIAGE_CERT_EXTRACT_HINT,
    "death_certificate": DEATH_EXTRACT_HINT,
    "birth_certificate_child": CHILD_BIRTH_EXTRACT_HINT,
    "military_discharge": MILITARY_EXTRACT_HINT,
    "ds260_customer_form": DS260_CUSTOMER_FORM_EXTRACT_HINT,
    "application_form": APPLICATION_FORM_EXTRACT_HINT,
}


async def _openai_extract(
    file_path: Path, doc_type: str, filename: str, db: AsyncSession, ctx: UsageContext
) -> dict:
    from app.services.document_registry import parse_document_filename
    from app.services.ds260_conflicts import LUONG1_DOC_TYPES
    from app.services.field_mapping import CONTACT_AND_SOCIAL_MAP

    expected_keys = get_extract_keys_for_doc_type(doc_type)
    if not expected_keys:
        from app.services.field_mapping import FIELD_MAP

        expected_keys = list(FIELD_MAP.get(doc_type, {}).keys())
    extra = DOC_TYPE_EXTRACT_HINTS.get(doc_type, "")
    if doc_type == "ds260_customer_form":
        extra = DS260_CUSTOMER_FORM_EXTRACT_HINT + (extra or "")
    elif doc_type not in RECORDABLE_REGISTRY_BY_CODE and doc_type in {
        "address_document",
        "financial",
        "employment_letter",
        "application_form",
    }:
        extra = ADDRESS_EXTRACT_HINT + SOCIAL_EXTRACT_HINT

    _, is_exception = parse_document_filename(filename)
    if is_exception and doc_type in LUONG1_DOC_TYPES:
        extra = (extra or "") + DS260_REFERENCE_EXTRACT_HINT
        expected_keys = list(
            dict.fromkeys([*expected_keys, *CONTACT_AND_SOCIAL_MAP.keys()])
        )
    if doc_type == "ds260_customer_form":
        # Worksheet có ~500 key nếu đổ phẳng (gồm cả alias và tên field của LOẠI GIẤY KHÁC)
        # → prompt loãng, model điền chéo. Chỉ nêu key thuộc worksheet, nhóm theo mục.
        from app.services.ds260_customer_keys import render_ds260_customer_prompt_keys

        rendered_keys = render_ds260_customer_prompt_keys()
    else:
        rendered_keys = ", ".join(expected_keys)
    prompt = EXTRACTION_PROMPT_TEMPLATE.format(
        doc_type=doc_type,
        expected_keys=rendered_keys or "any relevant fields for this document type",
    ) + extra
    # DS-260 worksheet khách khai thường là bản SCAN/VIẾT TAY nhiều trang (≈18 trang) →
    # đọc hết trang + render DPI cao hơn cho rõ nét chữ tay.
    is_ds260_ws = doc_type == "ds260_customer_form"
    image_paths, hint = _resolve_document_images(
        file_path,
        max_pages=3
        if doc_type == "passport"
        else (
            4
            if doc_type in {"birth_certificate", "marriage_certificate"}
            else (20 if is_ds260_ws else 2)
        ),
        dpi=300 if is_ds260_ws else 180,
    )
    image_labels: list[str] | None = None
    rows, cols = settings.ocr_worksheet_rows, settings.ocr_worksheet_cols
    if is_ds260_ws and image_paths and (rows > 1 or cols > 1):
        # Worksheet khách khai là bản SCAN VIẾT TAY → cắt lưới để chữ đủ lớn cho Vision API.
        tiled: list[Path] = []
        labels: list[str] = []
        for page_no, page_img in enumerate(image_paths, start=1):
            crops = _split_image_into_tiles(page_img, rows, cols)
            if len(crops) == 1:
                tiled.append(crops[0])
                labels.append(f"Page {page_no}")
                continue
            for i, crop in enumerate(crops):
                r, c = divmod(i, cols)
                tiled.append(crop)
                labels.append(
                    f"Page {page_no} — tile row {r + 1}/{rows}, col {c + 1}/{cols}"
                    f" ({'top' if r == 0 else 'bottom' if r == rows - 1 else 'middle'}"
                    f" {'left' if c == 0 else 'right' if c == cols - 1 else 'centre'})"
                )
        image_paths, image_labels = tiled, labels

    usage_ctx = UsageContext(
        user_id=ctx.user_id,
        applicant_id=ctx.applicant_id,
        document_id=ctx.document_id,
        filename=filename,
        doc_type=doc_type,
    )
    client = get_ocr_client()
    batches = _split_images_into_batches(
        image_paths, image_labels, settings.ocr_max_images_per_request
    )

    results: list[dict] = []
    for batch_no, (batch_paths, batch_labels) in enumerate(batches, start=1):
        content = _build_llm_user_content(
            filename=filename,
            image_paths=batch_paths,
            hint=hint,
            file_path=file_path,
            task="Extract all fields from this document.",
            image_labels=batch_labels,
            batch_position=(batch_no, len(batches)) if len(batches) > 1 else None,
        )
        response = await chat_completion(
            db,
            operation="document.extract",
            context=usage_ctx,
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": content},
            ],
            client=client,
            model=settings.ocr_model,
            response_format={"type": "json_object"},
            temperature=0,
        )
        results.append(_parse_json_response(response.choices[0].message.content or "{}"))

    return _merge_batch_results(results)


async def classify_document(
    file_path: Path, filename: str, db: AsyncSession, ctx: UsageContext
) -> tuple[dict, str | None]:
    if not is_openai_configured():
        return _mock_classification(filename), None
    try:
        return await _openai_classify(file_path, filename, db, ctx), None
    except Exception as exc:
        if _is_quota_error(exc):
            return _mock_classification(filename), QUOTA_WARNING
        raise


_IMAGE_QUALITY_ISSUE_VI: dict[str, str] = {
    "blurry": "ảnh mờ/out nét",
    "glare": "ảnh bị loá sáng",
    "too_dark": "ảnh thiếu sáng",
    "cut_off": "ảnh bị cắt thiếu nội dung",
    "handwriting_illegible": "chữ viết tay không đọc được",
}


def _image_quality_warning(extraction: dict) -> str | None:
    """Model tự báo ảnh khó đọc dù không dưới ngưỡng blur cứng (vd. loá sáng, thiếu sáng)."""
    quality = extraction.get("image_quality")
    if not isinstance(quality, dict) or quality.get("readable") is not False:
        return None
    issue = str(quality.get("issue") or "").strip()
    reason = _IMAGE_QUALITY_ISSUE_VI.get(issue, issue or "ảnh khó đọc")
    return (
        f"AI báo {reason} — một số trường có thể thiếu hoặc sai. "
        "Vui lòng kiểm tra kỹ, hoặc chụp/scan lại ảnh rõ hơn rồi xử lý lại."
    )


def _low_confidence_field_warning(extraction: dict) -> str | None:
    """Field CÓ giá trị nhưng model không chắc → chữ mực nhạt/thiếu sáng cục bộ mà
    _measure_blur_variance (đo trên toàn ảnh) không bắt được. Field null (không đọc được)
    không tính — đó là trường hợp model đã trung thực báo trống, không phải đoán mò.
    """
    threshold = settings.ocr_low_confidence_threshold
    min_ratio = settings.ocr_low_confidence_field_ratio
    if threshold <= 0:
        return None
    fields = extraction.get("fields")
    if not isinstance(fields, dict):
        return None

    filled_count = 0
    low_keys: list[str] = []
    for key, meta in fields.items():
        if not isinstance(meta, dict) or _is_empty_value(meta.get("value")):
            continue
        filled_count += 1
        try:
            confidence = float(meta.get("confidence"))
        except (TypeError, ValueError):
            continue
        if confidence < threshold:
            low_keys.append(key)

    if not low_keys or filled_count == 0 or len(low_keys) / filled_count < min_ratio:
        return None

    shown = ", ".join(low_keys[:8])
    if len(low_keys) > 8:
        shown += f" và {len(low_keys) - 8} trường khác"
    return (
        f"{len(low_keys)}/{filled_count} trường có độ tin cậy thấp (có thể do chữ mờ/nhạt màu, "
        f"không phải do ảnh out nét): {shown}. Vui lòng đối chiếu kỹ các trường này với bản gốc."
    )


def _extraction_warnings(extraction: dict) -> str | None:
    """Collect warnings reported by the AI model during extraction."""
    warnings = extraction.get("warnings")
    if not isinstance(warnings, list) or not warnings:
        return None
    formatted = [str(w).strip() for w in warnings if str(w).strip()]
    return " | ".join(formatted) if formatted else None


def _combine_warnings(*warnings: str | None) -> str | None:
    parts = [w for w in warnings if w]
    return " | ".join(parts) if parts else None


async def extract_document(
    file_path: Path, doc_type: str, filename: str, db: AsyncSession, ctx: UsageContext
) -> tuple[dict, str | None]:
    if not is_openai_configured():
        raw = _fallback_enrich_extraction(doc_type, filename, _mock_extraction(doc_type, filename))
        return _filter_extraction_to_schema(doc_type, raw), None
    try:
        extracted = await _openai_extract(file_path, doc_type, filename, db, ctx)
        enriched = _fallback_enrich_extraction(doc_type, filename, extracted)
        if doc_type == "birth_certificate":
            enriched = _coerce_birth_certificate_extraction(enriched)
        if doc_type == "ds260_customer_form":
            from app.services.ds260_customer_keys import coerce_ds260_customer_extraction

            enriched = coerce_ds260_customer_extraction(enriched)
        filtered = _filter_extraction_to_schema(doc_type, enriched)
        warning = _combine_warnings(
            _extraction_warnings(extracted),
            _image_quality_warning(enriched),
            _low_confidence_field_warning(filtered),
        )
        return filtered, warning
    except Exception as exc:
        if _is_quota_error(exc):
            raw = _fallback_enrich_extraction(doc_type, filename, _mock_extraction(doc_type, filename))
            return _filter_extraction_to_schema(doc_type, raw), QUOTA_WARNING
        raise


def _coerce_confidence_to_float(value: object) -> float | None:
    """Cùng lớp lỗi với _coerce_field_value_to_str: cột FLOAT insert thẳng string ("high"…) cũng
    ném DataError. Model thi thoảng trả confidence dạng chuỗi thay vì số dù prompt yêu cầu số."""
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _coerce_field_value_to_str(value: object) -> str | None:
    """`ExtractedField.field_value` là cột VARCHAR/TEXT — asyncpg insert thẳng list/dict vào đó
    ném DataError và ROLLBACK CẢ TRANSACTION (mất luôn mọi field khác đã đọc đúng của tài liệu).

    Prompt đã yêu cầu "value" luôn là string, nhưng model vẫn có lúc trả về object/array lồng
    nhau cho một field (sự cố thực tế 2026-08-04, OCR "01_7 JOB APPLICATION...docx" — model trả
    value kiểu [{"surname": "NGUYEN", "given_names": "..."}] thay vì chuỗi). Không thể chỉ tin
    prompt — phải chặn ở biên DB.
    """
    if value is None or isinstance(value, str):
        return value
    if isinstance(value, (list, dict)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


async def save_extracted_fields(db: AsyncSession, document_id, extraction: dict) -> list[ExtractedField]:
    await db.execute(delete(ExtractedField).where(ExtractedField.document_id == document_id))
    fields_data = extraction.get("fields", {})
    saved: list[ExtractedField] = []
    for key, meta in fields_data.items():
        if not isinstance(meta, dict):
            continue
        ef = ExtractedField(
            document_id=document_id,
            field_key=key,
            field_value=_coerce_field_value_to_str(meta.get("value")),
            confidence=_coerce_confidence_to_float(meta.get("confidence")),
            source_page=_coerce_field_value_to_str(meta.get("source_page")),
        )
        db.add(ef)
        saved.append(ef)
    await db.flush()
    return saved


async def process_document(db: AsyncSession, document: Document) -> Document:
    file_path = Path(document.file_path)
    document.status = DocumentStatus.processing
    await db.flush()

    # Chỉ kiểm ảnh CHỤP trực tiếp (jpg/png/...) — PDF thường là bản scan/xuất từ máy,
    # hiếm khi out nét kiểu chụp tay run, và _pdf_pages_to_images đã render riêng.
    if (
        settings.ocr_blur_variance_threshold > 0
        and file_path.suffix.lower() in IMAGE_SUFFIXES
        and file_path.exists()
    ):
        variance = _measure_blur_variance(file_path)
        if variance is not None and variance < settings.ocr_blur_variance_threshold:
            document.status = DocumentStatus.failed
            document.error_message = (
                f"Ảnh bị mờ/out nét (độ nét đo được {variance:.0f}, cần ≥ "
                f"{settings.ocr_blur_variance_threshold:.0f}) — hệ thống chưa trích xuất để tránh "
                "đọc sai dữ liệu. Vui lòng chụp lại ảnh rõ nét, đủ sáng, giữ máy ổn định rồi tải lên lại."
            )
            await db.flush()
            return document

    applicant = await db.get(Applicant, document.applicant_id)
    usage_ctx = UsageContext(
        user_id=applicant.user_id if applicant else None,
        applicant_id=document.applicant_id,
        document_id=document.id,
        filename=document.original_filename,
    )

    warnings: list[str] = []

    try:
        classification, w1 = await classify_document(file_path, document.original_filename, db, usage_ctx)
        if w1:
            warnings.append(w1)

        ai_type = classification.get("document_type", "other")
        doc_type, is_exception = _resolve_document_type(ai_type, document.original_filename or "")
        document.document_type = doc_type
        document.registry_doc_type = doc_type if doc_type in RECORDABLE_DOC_TYPES else None
        document.is_exception = is_exception
        document.classification_confidence = float(classification.get("confidence", 0))

        extraction, w2 = await extract_document(file_path, doc_type, document.original_filename, db, usage_ctx)
        if w2 and w2 not in warnings:
            warnings.append(w2)

        await save_extracted_fields(db, document.id, extraction)

        document.status = DocumentStatus.extracted
        from datetime import datetime, timezone

        document.processed_at = datetime.now(timezone.utc)
        document.error_message = warnings[0] if warnings else None
    except Exception as exc:
        if _is_quota_error(exc):
            classification = _mock_classification(document.original_filename)
            ai_type = classification.get("document_type", "other")
            doc_type, is_exception = _resolve_document_type(ai_type, document.original_filename or "")
            document.document_type = doc_type
            document.registry_doc_type = doc_type if doc_type in RECORDABLE_DOC_TYPES else None
            document.is_exception = is_exception
            document.classification_confidence = float(classification.get("confidence", 0))
            extraction = _filter_extraction_to_schema(
                doc_type,
                _fallback_enrich_extraction(
                    doc_type, document.original_filename, _mock_extraction(doc_type, document.original_filename)
                ),
            )
            await save_extracted_fields(db, document.id, extraction)
            document.status = DocumentStatus.extracted
            from datetime import datetime, timezone

            document.processed_at = datetime.now(timezone.utc)
            document.error_message = QUOTA_WARNING
        else:
            document.status = DocumentStatus.failed
            document.error_message = str(exc)[:2000]

    await db.flush()
    return document
